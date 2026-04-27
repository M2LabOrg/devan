"""
MCP Client UI - Web-based interface for managing MCP servers and local LLMs

This application provides a simple UI to:
- Select and configure local LLM providers (Ollama, Azure Foundry)
- Enable/disable MCP servers from this project
- Chat with the selected LLM using enabled MCP tools
"""

from flask import Flask, render_template, request, jsonify, session
from flask_socketio import SocketIO, emit
import asyncio
import json
import os
import re
import sys
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass, asdict
import requests
from dotenv import load_dotenv
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client
import threading
import uuid
import shutil
from contextlib import AsyncExitStack

# ── Audit logger for sandbox mode ──
audit_logger = logging.getLogger('sandbox_audit')
audit_logger.setLevel(logging.INFO)
_audit_log_dir = Path(__file__).parent / 'logs'
_audit_log_dir.mkdir(exist_ok=True)
_audit_handler = logging.FileHandler(_audit_log_dir / 'audit.log')
_audit_handler.setFormatter(logging.Formatter('%(asctime)s | %(message)s'))
audit_logger.addHandler(_audit_handler)

# Load .env from project root
load_dotenv(Path(__file__).parent.parent / '.env')

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# In-memory session storage for chat history (per-session, not persisted)
chat_histories: Dict[str, List[Dict]] = {}
uploaded_files: Dict[str, List[Dict]] = {}

# Project root directory
PROJECT_ROOT = Path(__file__).parent.parent

@dataclass
class MCPServer:
    """Configuration for an MCP server"""
    id: str
    name: str
    description: str
    path: str
    command: str
    args: List[str]
    enabled: bool = False

@dataclass
class LLMProvider:
    """Configuration for a local LLM provider"""
    id: str
    name: str
    type: str  # 'ollama', 'lmstudio', 'openai-compatible'
    base_url: str
    models: List[str] = None
    enabled: bool = False

# Available MCP servers in this project
MCP_SERVERS = [
    MCPServer(
        id="document_mcp",
        name="Document MCP",
        description="Unified document processing — Excel, PDF, and Word extraction, table detection, OCR, and RAG chunking",
        path=str(PROJECT_ROOT / "servers" / "document" / "mcp_project"),
        command="uv",
        args=["run", "document_server.py"]
    ),
    MCPServer(
        id="prompt_mcp",
        name="Prompt Engineering MCP",
        description="Prompt engineering with library management, PDF processing, and structured prompts",
        path=str(PROJECT_ROOT / "servers" / "prompt-engineering" / "mcp_project"),
        command="uv",
        args=["run", "prompt_server.py"]
    ),
    MCPServer(
        id="guardrail_mcp",
        name="Guardrail MCP",
        description="Content moderation and safety guardrails for LLM outputs",
        path=str(PROJECT_ROOT / "servers" / "guardrail" / "mcp_project"),
        command="uv",
        args=["run", "guardrail_server.py"]
    ),
    MCPServer(
        id="webdesign_mcp",
        name="Web Design MCP",
        description="Web design and HTML/CSS generation tools",
        path=str(PROJECT_ROOT / "servers" / "webdesign" / "mcp_project"),
        command="uv",
        args=["run", "webdesign_server.py"]
    ),
    MCPServer(
        id="webscraper_mcp",
        name="Web Scraper MCP",
        description="Intelligent web scraping with Firecrawl (premium) and BeautifulSoup (free) engines",
        path=str(PROJECT_ROOT / "servers" / "webscraper" / "mcp_project"),
        command="uv",
        args=["run", "webscraper_server.py"]
    ),
    MCPServer(
        id="excel_pipeline_mcp",
        name="Excel Pipeline MCP",
        description="Robust Excel data extraction, canonical modelling, validation, lineage tracking, and export",
        path=str(PROJECT_ROOT / "servers" / "excel-pipeline" / "mcp_project"),
        command="uv",
        args=["run", "excel_pipeline_server.py"]
    ),
    MCPServer(
        id="data_modelling_mcp",
        name="Data Modelling MCP",
        description="Ingest structured or semi-structured data, infer a relational data model, and export to SQLite, Arrow, Feather, or JSON",
        path=str(PROJECT_ROOT / "servers" / "data-modelling" / "mcp_project"),
        command="uv",
        args=["run", "data_modelling_server.py"]
    ),
    MCPServer(
        id="indexer_mcp",
        name="Document Indexer MCP",
        description="SQLite-backed RAG index: store chunks from any document type and answer questions with cited sources using the query tool",
        path=str(PROJECT_ROOT / "servers" / "indexer" / "mcp_project"),
        command="uv",
        args=["run", "indexer_server.py"]
    ),
]

# Default LLM providers
DEFAULT_LLM_PROVIDERS = [
    LLMProvider(
        id="ollama",
        name="Ollama",
        type="ollama",
        base_url=os.environ.get("OLLAMA_HOST", "http://localhost:11434"),
        models=[]
    ),
    LLMProvider(
        id="azure_foundry",
        name="Azure Foundry",
        type="openai-compatible",
        base_url="",
        models=[]
    ),
    LLMProvider(
        id="github_copilot",
        name="GitHub Copilot",
        type="copilot",
        base_url="https://models.inference.ai.azure.com",
        models=[]
    )
]

# Global state
active_sessions: Dict[str, Dict] = {}
config_file = PROJECT_ROOT / "client" / "config.json"
# Tracks active indexer session per Socket.IO session (sid → indexer session_id)
index_sessions: Dict[str, str] = {}

# Persistent KB session registry — survives container restarts alongside the SQLite index
KB_DATA_DIR = PROJECT_ROOT / "client" / "kb_data"
KB_SESSIONS_FILE = KB_DATA_DIR / "sessions.json"


def _load_kb_sessions() -> list:
    try:
        return json.loads(KB_SESSIONS_FILE.read_text()) if KB_SESSIONS_FILE.exists() else []
    except Exception:
        return []


def _save_kb_session(folder_path: str, session_id: str, files_indexed: int, total_chunks: int) -> None:
    # Deduplicate by both session_id and folder_path so re-indexing replaces the old entry
    sessions = [s for s in _load_kb_sessions()
                if s.get('session_id') != session_id and s.get('folder_path') != folder_path]
    sessions.insert(0, {
        'folder_path': folder_path,
        'host_path': _docker_to_host_path(folder_path),
        'session_id': session_id,
        'files_indexed': files_indexed,
        'total_chunks': total_chunks,
        'indexed_at': __import__('datetime').datetime.now().isoformat(timespec='seconds'),
    })
    KB_DATA_DIR.mkdir(parents=True, exist_ok=True)
    KB_SESSIONS_FILE.write_text(json.dumps(sessions[:20], indent=2))

# Supported extensions for the knowledge-base indexing pipeline
_INDEXABLE_EXTENSIONS = {
    '.pdf', '.xlsx', '.xls', '.xlsm',
    '.csv', '.txt', '.md',
    '.docx', '.doc',
    '.parquet',
}


# Sandbox-approved providers — sandbox mode restricts to local/approved
# providers with audit logging and guardrail enforcement.
SANDBOX_APPROVED_IDS = {'ollama', 'azure_foundry', 'github_copilot'}

# Quick PII patterns used for input scanning in sandbox mode.
# Mirrors a subset of PII_PATTERNS in guardrail_server.py.
_PII_QUICK_SCAN = {
    "email":       r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
    "ssn":         r'\b\d{3}-\d{2}-\d{4}\b',
    "credit_card": r'\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13})\b',
    "phone":       r'\b(?:\+1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',
    "iban":        r'\b[A-Z]{2}\d{2}[A-Z0-9]{4}\d{7}(?:[A-Z0-9]{0,16})?\b',
    "passport":    r'(?i)(?:passport[\s_-](?:number|no))[\s:=]+[A-Z]{1,2}[0-9]{6,9}',
    "medical_id":  r'(?i)(?:mrn|patient[\s_-]id|medical[\s_-]record)[\s:=]+[A-Z0-9]{5,15}',
    "national_id": r'(?i)(?:national[\s_-]id|nin|nric)[\s:=]+[A-Z0-9]{6,15}',
}


def scan_message_for_pii(message: str) -> List[str]:
    """Return a list of PII type names detected in *message* (fast, no line detail)."""
    return [
        pii_type
        for pii_type, pattern in _PII_QUICK_SCAN.items()
        if re.search(pattern, message)
    ]


def load_config():
    """Load configuration from file"""
    if config_file.exists():
        with open(config_file, 'r') as f:
            cfg = json.load(f)
            # Ensure compliance_mode field exists
            if 'compliance_mode' not in cfg:
                cfg['compliance_mode'] = False
            # Load GitHub Copilot token: .env takes priority over config.json
            for p in cfg.get('llm_providers', []):
                if p['id'] == 'github_copilot':
                    env_token = os.environ.get('github_copilot_token', '')
                    if env_token:
                        p['api_key'] = env_token
            return cfg
    return {
        "mcp_servers": [asdict(s) for s in MCP_SERVERS],
        "llm_providers": [asdict(p) for p in DEFAULT_LLM_PROVIDERS],
        "selected_llm": None,
        "selected_model": None,
        "compliance_mode": False
    }


def save_config(config):
    """Save configuration to file"""
    config_file.parent.mkdir(exist_ok=True)
    with open(config_file, 'w') as f:
        json.dump(config, f, indent=2)


def check_ollama_available():
    """Check if Ollama is running and get available models"""
    try:
        base_url = os.environ.get("OLLAMA_HOST", "http://localhost:11434").rstrip("/")
        response = requests.get(f"{base_url}/api/tags", timeout=2)
        if response.status_code == 200:
            models = [m['name'] for m in response.json().get('models', [])]
            return True, models
    except Exception:
        pass
    return False, []


def check_azure_foundry_available(provider: dict):
    """Check if Azure Foundry endpoint is configured and reachable"""
    base_url = provider.get('base_url', '')
    api_key = provider.get('api_key', '')
    if not base_url or not api_key:
        return False, []
    try:
        headers = {"api-key": api_key}
        # Azure OpenAI list-models endpoint
        url = base_url.rstrip('/') + '/models'
        response = requests.get(url, headers=headers, timeout=5)
        if response.status_code == 200:
            models = [m['id'] for m in response.json().get('data', [])]
            return (True, models) if models else (True, [provider.get('deployment_name', 'default')])
    except Exception:
        pass
    # If endpoint is configured with a deployment name, mark as available
    if base_url and api_key and provider.get('deployment_name'):
        return True, [provider['deployment_name']]
    return False, []


def check_github_copilot_available(provider: dict):
    """Check if GitHub Copilot token is configured and valid via GitHub Models API"""
    token = provider.get('api_key', '')
    if not token:
        return False, []
    try:
        headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json"
        }
        response = requests.get(
            "https://models.inference.ai.azure.com/models",
            headers=headers, timeout=8
        )
        if response.status_code == 200:
            data = response.json()
            models = []
            items = data if isinstance(data, list) else data.get('data', [])
            for m in items:
                mid = m.get('id', '')
                # Extract short name from azureml://registries/.../models/NAME/versions/N
                if '/models/' in mid:
                    name = mid.split('/models/')[1].split('/versions/')[0]
                else:
                    name = mid
                # Only include chat-capable models (skip embeddings, whisper, dall-e)
                if not any(x in name.lower() for x in ['embed', 'whisper', 'dall', 'safety']):
                    models.append(name)
            if models:
                return True, models
    except Exception:
        pass
    # Fallback: if token is set, assume available with known models
    if token:
        return True, ['gpt-4o', 'gpt-4o-mini', 'Meta-Llama-3.1-405B-Instruct', 'Meta-Llama-3.1-8B-Instruct']
    return False, []


async def call_llm(provider: dict, model: str, messages: List[dict], tools: List[dict] = None):
    """Call local LLM with OpenAI-compatible API"""
    if provider['type'] == 'ollama':
        # Ollama API
        url = f"{provider['base_url']}/api/chat"
        payload = {
            "model": model,
            "messages": messages,
            "stream": False
        }
        if tools:
            payload["tools"] = tools
        
        response = requests.post(url, json=payload)
        return response.json()
    
    else:  # OpenAI-compatible (includes Azure Foundry and GitHub Copilot)
        url = f"{provider['base_url'].rstrip('/')}/chat/completions"
        headers = {"Content-Type": "application/json"}
        # GitHub Copilot uses Bearer token
        if provider.get('type') == 'copilot' and provider.get('api_key'):
            headers["Authorization"] = f"Bearer {provider['api_key']}"
        # Azure Foundry uses api-key header
        elif provider.get('api_key'):
            headers["api-key"] = provider['api_key']
        payload = {
            "model": model,
            "messages": messages,
            "max_tokens": 2048
        }
        if tools:
            payload["tools"] = tools
        
        response = requests.post(url, json=payload, headers=headers)
        return response.json()


def _is_setup_needed():
    """Check if first-run setup is needed (no venv or no servers directory)."""
    client_dir = Path(__file__).parent
    venv_dir = client_dir / 'venv'
    servers_dir = PROJECT_ROOT / 'servers'
    # Setup is needed if there's no venv or no servers directory
    venv_ok = venv_dir.exists() and (venv_dir / 'bin' / 'python').exists()
    if sys.platform == 'win32':
        venv_ok = venv_dir.exists() and (venv_dir / 'Scripts' / 'python.exe').exists()
    servers_ok = servers_dir.exists() and any(servers_dir.iterdir())
    return not (venv_ok and servers_ok)


@app.route('/')
def index():
    """Main page — redirects to setup if first run."""
    if _is_setup_needed():
        return render_template('setup.html')
    return render_template('index.html')


@app.route('/setup')
def setup_page():
    """Setup page — always accessible for re-running setup."""
    return render_template('setup.html')


# ── First-Run Setup API ──────────────────────────────────────────────────────

@app.route('/api/setup/check-python', methods=['POST'])
def setup_check_python():
    """Check if Python 3.10+ is available."""
    import subprocess as _sp
    for cmd in ['python3', 'python']:
        try:
            out = _sp.run([cmd, '--version'], capture_output=True, text=True, timeout=5)
            if out.returncode == 0:
                version = out.stdout.strip() or out.stderr.strip()
                path = _sp.run(['which' if sys.platform != 'win32' else 'where', cmd],
                               capture_output=True, text=True, timeout=5).stdout.strip()
                return jsonify({'ok': True, 'version': version, 'path': path})
        except Exception:
            continue
    return jsonify({'ok': False, 'version': None, 'path': None})


@app.route('/api/setup/install-python', methods=['POST'])
def setup_install_python():
    """Auto-install Python using Homebrew (macOS), apt (Linux), or winget/Microsoft Store (Windows)."""
    import subprocess as _sp
    import platform
    
    system = platform.system()
    
    try:
        if system == 'Darwin':  # macOS
            # Check if Homebrew is available
            brew_check = _sp.run(['which', 'brew'], capture_output=True, text=True)
            if brew_check.returncode != 0:
                return jsonify({'ok': False, 'error': 'Homebrew not found. Please install Homebrew first from https://brew.sh'})
            
            # Install Python via Homebrew
            cmd = ['brew', 'install', 'python@3.11']
            result = _sp.run(cmd, capture_output=True, text=True, timeout=300)
            
        elif system == 'Linux':
            # Try apt (Ubuntu/Debian)
            cmd = ['sh', '-c', 'sudo apt update && sudo apt install -y python3 python3-pip']
            result = _sp.run(cmd, capture_output=True, text=True, timeout=300)
            
        elif system == 'Windows':
            # Try winget first (Windows Package Manager)
            winget_check = _sp.run(['winget', '--version'], capture_output=True, text=True)
            if winget_check.returncode == 0:
                # Use winget to install Python
                cmd = ['winget', 'install', '-e', '--id', 'Python.Python.3.11', '--accept-source-agreements', '--accept-package-agreements']
                result = _sp.run(cmd, capture_output=True, text=True, timeout=300)
            else:
                # Fall back to Microsoft Store Python
                return jsonify({
                    'ok': False, 
                    'error': 'Auto-install not available. Please install Python from the Microsoft Store (search "Python 3.11") or download from python.org',
                    'fallback_url': 'https://apps.microsoft.com/detail/9pjpw5ldxlzc'
                })
        else:
            return jsonify({'ok': False, 'error': f'Unsupported platform: {system}'})
        
        if result.returncode == 0:
            # Verify installation
            python_cmd = 'python' if system == 'Windows' else 'python3'
            try:
                verify = _sp.run([python_cmd, '--version'], capture_output=True, text=True, timeout=5)
                if verify.returncode == 0:
                    version = verify.stdout.strip() or verify.stderr.strip()
                    where_cmd = 'where' if system == 'Windows' else 'which'
                    path = _sp.run([where_cmd, python_cmd], capture_output=True, text=True, timeout=5).stdout.strip().split('\n')[0]
                    return jsonify({'ok': True, 'version': version, 'path': path, 'output': result.stdout})
            except Exception as e:
                return jsonify({'ok': False, 'error': f'Installation appeared to succeed but Python not found: {str(e)}', 'output': result.stdout})
        else:
            return jsonify({'ok': False, 'error': result.stderr or 'Installation failed', 'output': result.stdout})
    except _sp.TimeoutExpired:
        return jsonify({'ok': False, 'error': 'Installation timed out (took longer than 5 minutes)'})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})


@app.route('/api/setup/check-uv', methods=['POST'])
def setup_check_uv():
    """Check if uv is available (in PATH or common install locations)."""
    import subprocess as _sp
    import os
    
    # First try PATH
    try:
        out = _sp.run(['uv', '--version'], capture_output=True, text=True, timeout=5)
        if out.returncode == 0:
            version = out.stdout.strip()
            path = _sp.run(['which' if sys.platform != 'win32' else 'where', 'uv'],
                           capture_output=True, text=True, timeout=5).stdout.strip()
            return jsonify({'ok': True, 'version': version, 'path': path})
    except Exception:
        pass
    
    # Check common install locations
    possible_paths = [
        os.path.expanduser('~/.local/bin/uv'),
        os.path.expanduser('~/.cargo/bin/uv'),
        '/usr/local/bin/uv',
    ]
    
    for uv_path in possible_paths:
        if os.path.exists(uv_path):
            try:
                out = _sp.run([uv_path, '--version'], capture_output=True, text=True, timeout=5)
                if out.returncode == 0:
                    version = out.stdout.strip()
                    return jsonify({'ok': True, 'version': version, 'path': uv_path})
            except Exception:
                pass
    
    return jsonify({'ok': False, 'version': None, 'path': None})


@app.route('/api/setup/install-uv', methods=['POST'])
def setup_install_uv():
    """Auto-install uv using the official installer."""
    import subprocess as _sp
    import platform
    
    system = platform.system()
    
    try:
        if system == 'Darwin':  # macOS
            # Use the official curl installer
            cmd = ['sh', '-c', 'curl -LsSf https://astral.sh/uv/install.sh | sh']
        elif system == 'Linux':
            cmd = ['sh', '-c', 'curl -LsSf https://astral.sh/uv/install.sh | sh']
        elif system == 'Windows':
            # Windows PowerShell install
            cmd = ['powershell', '-Command', 
                   'irm https://astral.sh/uv/install.ps1 | iex']
        else:
            return jsonify({'ok': False, 'error': f'Unsupported platform: {system}'})
        
        # Run the installer
        result = _sp.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode == 0:
            # Verify installation - check common install locations first
            uv_path = None
            possible_paths = [
                os.path.expanduser('~/.local/bin/uv'),
                os.path.expanduser('~/.cargo/bin/uv'),
                '/usr/local/bin/uv',
            ]
            for p in possible_paths:
                if os.path.exists(p):
                    uv_path = p
                    break
            
            # If not found in common locations, try PATH
            if not uv_path:
                uv_path = 'uv'
            
            try:
                verify = _sp.run([uv_path, '--version'], capture_output=True, text=True, timeout=5)
                if verify.returncode == 0:
                    version = verify.stdout.strip()
                    # Get actual path if we used a full path
                    if uv_path != 'uv':
                        path = uv_path
                    else:
                        path = _sp.run(['which' if sys.platform != 'win32' else 'where', 'uv'],
                                       capture_output=True, text=True, timeout=5).stdout.strip()
                    return jsonify({'ok': True, 'version': version, 'path': path, 'output': result.stdout})
                else:
                    return jsonify({'ok': False, 'error': f'uv found but returned error: {verify.stderr}', 'output': result.stdout})
            except Exception as e:
                return jsonify({'ok': False, 'error': f'Installation appeared to succeed but uv not found: {str(e)}', 'output': result.stdout})
        else:
            return jsonify({'ok': False, 'error': result.stderr or 'Installation failed', 'output': result.stdout})
    except _sp.TimeoutExpired:
        return jsonify({'ok': False, 'error': 'Installation timed out (took longer than 2 minutes)'})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})


@app.route('/api/setup/check-git', methods=['POST'])
def setup_check_git():
    """Check if Git is available."""
    import subprocess as _sp
    try:
        out = _sp.run(['git', '--version'], capture_output=True, text=True, timeout=5)
        if out.returncode == 0:
            version = out.stdout.strip()
            path = _sp.run(['which' if sys.platform != 'win32' else 'where', 'git'],
                           capture_output=True, text=True, timeout=5).stdout.strip()
            return jsonify({'ok': True, 'version': version, 'path': path})
    except Exception:
        pass
    return jsonify({'ok': False, 'version': None, 'path': None})


@app.route('/api/setup/check-repo', methods=['POST'])
def setup_check_repo():
    """Check if MCP servers repo exists locally."""
    servers_dir = PROJECT_ROOT / 'servers'
    if servers_dir.exists() and any(servers_dir.iterdir()):
        return jsonify({'ok': True, 'path': str(servers_dir)})
    return jsonify({'ok': False, 'path': None})


@app.route('/api/setup/clone-repo', methods=['POST'])
def setup_clone_repo():
    """Clone the MCP servers repository."""
    import subprocess as _sp
    repo_url = 'https://github.com/M2LabOrg/mcp-design-deploy.git'
    clone_target = PROJECT_ROOT
    try:
        # If PROJECT_ROOT doesn't have servers/, clone the repo there
        if not (clone_target / 'servers').exists():
            out = _sp.run(
                ['git', 'clone', '--depth', '1', repo_url, str(clone_target / '_mcp_repo')],
                capture_output=True, text=True, timeout=300
            )
            if out.returncode == 0:
                # Move servers directory into place
                cloned_servers = clone_target / '_mcp_repo' / 'servers'
                target_servers = clone_target / 'servers'
                if cloned_servers.exists():
                    shutil.move(str(cloned_servers), str(target_servers))
                # Clean up the clone
                shutil.rmtree(str(clone_target / '_mcp_repo'), ignore_errors=True)

                # Update config.json with correct server paths
                _update_server_paths(target_servers)

                return jsonify({'ok': True, 'path': str(target_servers)})
            else:
                return jsonify({'ok': False, 'error': out.stderr.strip()})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})
    return jsonify({'ok': False, 'error': 'Could not clone repository'})


def _update_server_paths(servers_dir: Path):
    """Update config.json MCP server paths to point to the actual location."""
    config = load_config()
    server_map = {
        'document_mcp': 'document/mcp_project',
        'prompt_mcp': 'prompt-engineering/mcp_project',
        'guardrail_mcp': 'guardrail/mcp_project',
        'webdesign_mcp': 'webdesign/mcp_project',
    }
    for srv in config.get('mcp_servers', []):
        relative = server_map.get(srv['id'])
        if relative:
            srv['path'] = str(servers_dir / relative)
    save_config(config)


@app.route('/api/setup/setup-venv', methods=['POST'])
def setup_create_venv():
    """Create the client virtual environment and install dependencies."""
    import subprocess as _sp
    client_dir = Path(__file__).parent
    venv_dir = client_dir / 'venv'
    requirements = client_dir / 'requirements.txt'
    try:
        # Create venv
        _sp.run([sys.executable, '-m', 'venv', str(venv_dir)],
                capture_output=True, text=True, timeout=60, check=True)
        # Install requirements
        pip = str(venv_dir / ('Scripts' if sys.platform == 'win32' else 'bin') / 'pip')
        _sp.run([pip, 'install', '-q', '-r', str(requirements)],
                capture_output=True, text=True, timeout=300, check=True)
        return jsonify({'ok': True})
    except _sp.CalledProcessError as e:
        return jsonify({'ok': False, 'error': e.stderr.strip() if e.stderr else str(e)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})


@app.route('/api/setup/check-ollama', methods=['POST'])
def setup_check_ollama():
    """Check if Ollama is running and list models."""
    available, models = check_ollama_available()
    return jsonify({'ok': available, 'models': models})


@app.route('/api/setup/pull-model', methods=['POST'])
def setup_pull_model():
    """Pull the default Ollama model (llama3.2)."""
    import subprocess as _sp
    try:
        out = _sp.run(['ollama', 'pull', 'llama3.2'],
                      capture_output=True, text=True, timeout=600)
        if out.returncode == 0:
            return jsonify({'ok': True})
        return jsonify({'ok': False, 'error': out.stderr.strip()})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})


@app.route('/docs/<path:filename>')
def serve_docs(filename):
    """Serve documentation files"""
    from flask import send_from_directory
    docs_dir = PROJECT_ROOT / 'docs'
    return send_from_directory(str(docs_dir), filename)


@app.route('/api/config', methods=['GET'])
def get_config():
    """Get current configuration"""
    config = load_config()
    
    # Check LLM availability
    for provider in config['llm_providers']:
        if provider['id'] == 'ollama':
            available, models = check_ollama_available()
            provider['available'] = available
            provider['models'] = models
        elif provider['id'] == 'azure_foundry':
            available, models = check_azure_foundry_available(provider)
            provider['available'] = available
            if models:
                provider['models'] = models
        elif provider['id'] == 'github_copilot':
            available, models = check_github_copilot_available(provider)
            provider['available'] = available
            if models:
                provider['models'] = models

    # Mark which providers are compliant (localhost-only)
    for provider in config['llm_providers']:
        provider['compliant'] = provider['id'] in SANDBOX_APPROVED_IDS

    return jsonify(config)


@app.route('/api/compliance', methods=['POST'])
def toggle_compliance():
    """Toggle sandbox mode — restricts to local providers, enables guardrails and audit logging"""
    data = request.json
    enabled = data.get('enabled', False)

    config = load_config()
    config['compliance_mode'] = enabled

    # If turning on sandbox mode, deselect any non-local provider
    if enabled and config.get('selected_llm') and config['selected_llm'] not in SANDBOX_APPROVED_IDS:
        audit_logger.info(f"SANDBOX | Deselected non-local provider '{config['selected_llm']}' on activation")
        config['selected_llm'] = None
        config['selected_model'] = None

    # Auto-enable guardrail MCP when sandbox is activated
    if enabled:
        for server in config.get('mcp_servers', []):
            if server['id'] == 'guardrail_mcp':
                server['enabled'] = True

    save_config(config)
    audit_logger.info(f"SANDBOX | Sandbox mode {'ENABLED' if enabled else 'DISABLED'}")
    return jsonify({"status": "success", "compliance_mode": enabled})


@app.route('/api/sandbox/status')
def sandbox_status():
    """Return sandbox health and statistics (provider restrictions, guardrail state, audit counts)."""
    config = load_config()
    is_active = config.get('compliance_mode', False)
    guardrail_enabled = any(
        s['id'] == 'guardrail_mcp' and s.get('enabled', False)
        for s in config.get('mcp_servers', [])
    )
    log_file = Path(__file__).parent / 'logs' / 'audit.log'
    total_entries = 0
    log_size_bytes = 0
    if log_file.exists():
        log_size_bytes = log_file.stat().st_size
        with open(log_file, 'r') as f:
            total_entries = sum(1 for line in f if line.strip())
    return jsonify({
        "active": is_active,
        "approved_providers": sorted(SANDBOX_APPROVED_IDS),
        "guardrail_enforced": guardrail_enabled,
        "pii_scanning": is_active,          # input PII scan runs whenever sandbox is on
        "audit_log_entries": total_entries,
        "audit_log_size_bytes": log_size_bytes,
        "selected_llm": config.get('selected_llm'),
        "selected_model": config.get('selected_model'),
    })


@app.route('/api/sandbox/audit')
def get_audit_log():
    """Return recent sandbox audit log entries as structured JSON."""
    limit = request.args.get('limit', 100, type=int)
    limit = min(limit, 500)  # cap to prevent large payloads
    log_file = Path(__file__).parent / 'logs' / 'audit.log'
    if not log_file.exists():
        return jsonify({"entries": [], "total": 0})
    try:
        with open(log_file, 'r') as f:
            lines = [l.rstrip('\n') for l in f if l.strip()]
        recent = lines[-limit:]
        entries = []
        for raw in recent:
            # Format: "2025-01-01 12:00:00,123 | EVENT_TYPE | detail text"
            parts = raw.split(' | ', 2)
            entries.append({
                "timestamp": parts[0] if len(parts) > 0 else "",
                "event":     parts[1] if len(parts) > 1 else raw,
                "detail":    parts[2] if len(parts) > 2 else "",
                "raw":       raw,
            })
        return jsonify({"entries": entries, "total": len(lines)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/config', methods=['POST'])
def update_config():
    """Update configuration — enforces sandbox invariants"""
    config = request.json

    # Sandbox enforcement: if sandbox is on, enforce invariants
    if config.get('compliance_mode'):
        # Prevent selecting non-approved providers
        if config.get('selected_llm') and config['selected_llm'] not in SANDBOX_APPROVED_IDS:
            audit_logger.warning(f"BLOCKED | Config write tried to set non-approved provider '{config['selected_llm']}' in sandbox mode")
            config['selected_llm'] = None
            config['selected_model'] = None
        # Guardrail MCP must stay enabled in sandbox mode
        for server in config.get('mcp_servers', []):
            if server['id'] == 'guardrail_mcp':
                server['enabled'] = True

    save_config(config)
    return jsonify({"status": "success"})


@app.route('/api/mcp/toggle', methods=['POST'])
def toggle_mcp():
    """Toggle MCP server on/off"""
    data = request.json
    server_id = data.get('server_id')
    enabled = data.get('enabled')

    config = load_config()

    # Sandbox enforcement: cannot disable Guardrail MCP while sandbox is active
    if config.get('compliance_mode') and server_id == 'guardrail_mcp' and not enabled:
        audit_logger.warning("BLOCKED | Attempted to disable Guardrail MCP while sandbox is active")
        return jsonify({"status": "error", "message": "Cannot disable Guardrail MCP while Sandbox Mode is active"}), 403

    for server in config['mcp_servers']:
        if server['id'] == server_id:
            server['enabled'] = enabled
            break

    save_config(config)
    return jsonify({"status": "success"})


@app.route('/api/llm/configure', methods=['POST'])
def configure_llm():
    """Configure LLM provider connection details (Azure Foundry / GitHub Copilot)"""
    data = request.json
    provider_id = data.get('provider_id')
    base_url = data.get('base_url', '').strip()
    api_key = data.get('api_key', '').strip()
    deployment_name = data.get('deployment_name', '').strip()

    config = load_config()
    for provider in config['llm_providers']:
        if provider['id'] == provider_id:
            if base_url:
                provider['base_url'] = base_url
            if api_key:
                provider['api_key'] = api_key
            if deployment_name:
                provider['deployment_name'] = deployment_name
                provider['models'] = [deployment_name]
            break

    save_config(config)

    # Verify connectivity
    updated_provider = next((p for p in config['llm_providers'] if p['id'] == provider_id), None)
    if updated_provider:
        if provider_id == 'github_copilot':
            available, models = check_github_copilot_available(updated_provider)
        else:
            available, models = check_azure_foundry_available(updated_provider)
        return jsonify({"status": "success", "available": available, "models": models})
    return jsonify({"status": "error", "message": "Provider not found"}), 404


@app.route('/api/llm/select', methods=['POST'])
def select_llm():
    """Select LLM provider and model"""
    data = request.json
    provider_id = data.get('provider_id')
    model = data.get('model')

    config = load_config()

    # Sandbox enforcement: block non-approved provider selection
    if config.get('compliance_mode') and provider_id not in SANDBOX_APPROVED_IDS:
        audit_logger.warning(f"BLOCKED | Attempted to select non-approved provider '{provider_id}' in sandbox mode")
        return jsonify({"status": "error", "message": "Provider not allowed in Sandbox Mode"}), 403

    config['selected_llm'] = provider_id
    config['selected_model'] = model
    save_config(config)

    return jsonify({"status": "success"})


@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    sid = request.sid
    # Initialize empty chat history for this session keyed by Socket.IO SID
    chat_histories[sid] = []
    uploaded_files[sid] = []
    emit('connected', {'session_id': sid})


@socketio.on('start_chat')
def handle_start_chat(data):
    """Initialize chat session with MCP servers"""
    from flask import request
    sid = request.sid
    print(f"DEBUG: start_chat received from SID: {sid}")
    
    config = load_config()
    enabled_servers = [s for s in config['mcp_servers'] if s.get('enabled', False)]
    
    # Send status immediately using socketio.emit for maximum reliability
    socketio.emit('chat_status', {
        'status': 'info',
        'message': f'Initializing with {len(enabled_servers)} MCP server(s)...'
    }, room=sid)
    
    # Activate chat immediately
    socketio.emit('chat_ready', {
        'tools_count': 0,
        'servers': [s['name'] for s in enabled_servers]
    }, room=sid)
    
    print(f"DEBUG: Chat initialized and ready emitted for session {sid}")


def initialize_mcp_sessions(sid: str, servers: List[dict], config: dict):
    """Background task for MCP connections (currently simplified)"""
    pass


@socketio.on('send_message')
def handle_message(data):
    """Handle chat message with @resource and /command syntax support"""
    from flask import request
    sid = request.sid
    message = data.get('message')
    request_id = data.get('requestId')
    
    config = load_config()
    
    # Check for @resource syntax
    if message.strip().startswith('@'):
        threading.Thread(
            target=handle_resource_request,
            args=(sid, message, config)
        ).start()
        return
    
    # Check for /command syntax
    if message.strip().startswith('/'):
        threading.Thread(
            target=handle_command_request,
            args=(sid, message, config)
        ).start()
        return
    
    # Regular chat message
    if not config.get('selected_llm') or not config.get('selected_model'):
        emit('chat_response', {
            'error': 'Please select an LLM provider and model first'
        })
        return
    
    # Process message in background with history
    threading.Thread(
        target=process_chat_message,
        args=(sid, message, config, request_id)
    ).start()


def handle_resource_request(sid: str, message: str, config: dict):
    """Handle @resource syntax requests"""
    try:
        # Remove @ sign and parse
        resource_path = message.strip()[1:].strip()
        
        # Map common resource patterns
        resource_map = {
            'pdf://files': 'List available PDF files',
            'prompts://library': 'Browse prompt library',
            'folders': 'List available folders (Document MCP)',
        }
        
        # Check if it's a known resource
        if resource_path in resource_map or resource_path.startswith('prompts://') or resource_path.startswith('pdf://'):
            socketio.emit('chat_response', {
                'message': f'📁 Resource Request: `@{resource_path}`\n\nThis would fetch the resource from enabled MCP servers.\n\n**Available resources:**\n- `@pdf://files` - List PDF files\n- `@prompts://library` - Browse prompts\n- `@prompts://<name>` - View specific prompt\n\nNote: Full MCP resource integration requires active server connections.',
                'model': 'system'
            }, room=sid)
        else:
            socketio.emit('chat_response', {
                'message': f'❓ Unknown resource: `@{resource_path}`\n\nTry:\n- `@pdf://files`\n- `@prompts://library`\n- `@prompts://risk_assessment`',
                'model': 'system'
            }, room=sid)
    
    except Exception as e:
        socketio.emit('chat_response', {
            'error': f'Error handling resource: {str(e)}'
        }, room=sid)


def handle_command_request(sid: str, message: str, config: dict):
    """Handle /command syntax requests"""
    try:
        # Parse command
        parts = message.strip()[1:].split()
        command = parts[0].lower() if parts else ''
        args = parts[1:] if len(parts) > 1 else []
        
        # Handle different commands
        if command == 'help':
            help_text = """
📖 **Available Commands:**

**Resources (@):**
- `@pdf://files` - List available PDFs
- `@prompts://library` - Browse prompt library
- `@prompts://<name>` - View specific prompt

**Commands (/):**
- `/help` - Show this help
- `/prompts` - List available prompts
- `/tools` - List available MCP tools
- `/servers` - Show enabled servers
- `/clear` - Clear chat history

**Examples:**
- `@pdf://files` then ask "Analyze report.pdf"
- `/prompts` to see available prompts
- Regular chat for LLM conversations
"""
            socketio.emit('chat_response', {
                'message': help_text,
                'model': 'system'
            }, room=sid)
        
        elif command == 'prompts':
            enabled_servers = [s for s in config['mcp_servers'] if s.get('enabled')]
            has_prompt_server = any('prompt' in s['id'].lower() for s in enabled_servers)
            
            if has_prompt_server:
                socketio.emit('chat_response', {
                    'message': '📚 **Available Prompts:**\n\n- `executive_summary`\n- `introduction`\n- `technical_summary`\n- `custom_prompt`\n\nUse `@prompts://library` to browse the full library.',
                    'model': 'system'
                }, room=sid)
            else:
                socketio.emit('chat_response', {
                    'message': '⚠️ Prompt Engineering MCP server is not enabled. Enable it to access prompts.',
                    'model': 'system'
                }, room=sid)
        
        elif command == 'tools':
            enabled_servers = [s for s in config['mcp_servers'] if s.get('enabled')]
            if enabled_servers:
                tools_info = f"🔧 **Enabled MCP Servers:** {len(enabled_servers)}\n\n"
                for server in enabled_servers:
                    tools_info += f"- **{server['name']}**\n  {server['description']}\n\n"
                socketio.emit('chat_response', {
                    'message': tools_info,
                    'model': 'system'
                }, room=sid)
            else:
                socketio.emit('chat_response', {
                    'message': '⚠️ No MCP servers enabled. Click the command bar and toggle servers on.',
                    'model': 'system'
                }, room=sid)
        
        elif command == 'servers':
            enabled = [s['name'] for s in config['mcp_servers'] if s.get('enabled')]
            if enabled:
                socketio.emit('chat_response', {
                    'message': f"✅ **Enabled Servers:**\n\n" + "\n".join(f"- {s}" for s in enabled),
                    'model': 'system'
                }, room=sid)
            else:
                socketio.emit('chat_response', {
                    'message': '⚠️ No servers enabled.',
                    'model': 'system'
                }, room=sid)
        
        elif command == 'clear' or command == 'new':
            # sid is already available as the function parameter — do NOT use
            # request.sid here; this function runs in a background thread where
            # Flask's request context is unavailable.
            if sid in chat_histories:
                chat_histories[sid] = []
            if sid in uploaded_files:
                uploaded_files[sid] = []
            socketio.emit('clear_chat', {}, room=sid)
            socketio.emit('chat_response', {
                'message': '🗑️ New chat started. History cleared.',
                'model': 'system'
            }, room=sid)
        
        else:
            socketio.emit('chat_response', {
                'message': f'❓ Unknown command: `/{command}`\n\nType `/help` for available commands.',
                'model': 'system'
            }, room=sid)
    
    except Exception as e:
        socketio.emit('chat_response', {
            'error': f'Error handling command: {str(e)}'
        }, room=sid)


# ── Agentic loop helpers ──────────────────────────────────────────────────────

def _uv_path() -> str:
    """Resolve the uv binary, checking common install locations."""
    uv = shutil.which('uv')
    if uv:
        return uv
    local_uv = os.path.expanduser('~/.local/bin/uv')
    if os.path.exists(local_uv):
        return local_uv
    return 'uv'


def _subprocess_env() -> dict:
    """Build env dict for MCP subprocesses, ensuring ~/.local/bin is on PATH."""
    env = dict(os.environ)
    local_bin = os.path.expanduser('~/.local/bin')
    if local_bin not in env.get('PATH', ''):
        env['PATH'] = local_bin + ':' + env.get('PATH', '')
    return env


def _call_llm_sync(provider: dict, model: str, messages: List[dict], tools=None) -> dict:
    """Synchronous LLM call — safe to run in asyncio.to_thread."""
    if provider['type'] == 'ollama':
        url = f"{provider['base_url']}/api/chat"
        payload = {"model": model, "messages": messages, "stream": False}
        if tools:
            payload["tools"] = tools
        resp = requests.post(url, json=payload, timeout=300)
        return resp.json()
    else:
        url = f"{provider['base_url'].rstrip('/')}/chat/completions"
        headers = {"Content-Type": "application/json"}
        if provider.get('type') == 'copilot' and provider.get('api_key'):
            headers["Authorization"] = f"Bearer {provider['api_key']}"
        elif provider.get('api_key'):
            headers["api-key"] = provider['api_key']
        payload = {"model": model, "messages": messages, "max_tokens": 2048}
        if tools:
            payload["tools"] = tools
        resp = requests.post(url, json=payload, headers=headers, timeout=300)
        return resp.json()


def _extract_tool_calls(response: dict, provider_type: str) -> List[dict]:
    """Return list of {id, name, args} dicts from an LLM response."""
    calls = []
    if provider_type == 'ollama':
        for tc in response.get('message', {}).get('tool_calls', []):
            fn = tc.get('function', {})
            args = fn.get('arguments', {})
            if isinstance(args, str):
                try:
                    args = json.loads(args)
                except Exception:
                    args = {}
            calls.append({'id': None, 'name': fn.get('name', ''), 'args': args})
    else:
        choices = response.get('choices', [])
        if choices:
            for tc in choices[0].get('message', {}).get('tool_calls', []):
                fn = tc.get('function', {})
                args = fn.get('arguments', '{}')
                if isinstance(args, str):
                    try:
                        args = json.loads(args)
                    except Exception:
                        args = {}
                calls.append({'id': tc.get('id'), 'name': fn.get('name', ''), 'args': args})
    return calls


def _docker_to_host_path(docker_path: str) -> str:
    """Translate a container-internal path to the equivalent host path."""
    host_home = os.environ.get('HOST_HOME', '')
    if docker_path.startswith('/host/home/') and host_home:
        return host_home + '/' + docker_path[len('/host/home/'):]
    if docker_path.startswith('/host/volumes/'):
        return '/Volumes/' + docker_path[len('/host/volumes/'):]
    return docker_path


def _get_response_text(response: dict, provider_type: str) -> str:
    """Extract the final text content from an LLM response."""
    if provider_type == 'ollama':
        return response.get('message', {}).get('content', '') or ''
    choices = response.get('choices', [])
    if choices:
        return choices[0].get('message', {}).get('content', '') or ''
    return ''


async def _collect_mcp_tools(servers: List[dict]):
    """
    Connect to all enabled MCP servers and discover their tools.

    Returns (openai_tools, tool_registry, exit_stack).
    The caller must call `await exit_stack.aclose()` when done.
    tool_registry maps tool_name -> {'session': ClientSession, 'server_name': str}.
    """
    exit_stack = AsyncExitStack()
    openai_tools: List[dict] = []
    tool_registry: Dict[str, dict] = {}
    uv = _uv_path()
    env = _subprocess_env()

    for server in servers:
        try:
            command = uv if server.get('command') == 'uv' else server['command']
            raw_args = server.get('args', [])
            if server.get('command') == 'uv' and raw_args and raw_args[0] == 'run':
                # Use --directory so uv activates the project's own .venv,
                # rather than falling back to the system Python.
                # Result: uv --directory /path/to/mcp_project run script.py
                resolved_args = ['--directory', server['path']] + raw_args
            else:
                resolved_args = raw_args
            params = StdioServerParameters(
                command=command,
                args=resolved_args,
                env=env,
            )
            read, write = await exit_stack.enter_async_context(stdio_client(params))
            session = await exit_stack.enter_async_context(ClientSession(read, write))
            await session.initialize()

            result = await session.list_tools()
            for tool in result.tools:
                schema = dict(tool.inputSchema) if tool.inputSchema else {
                    "type": "object", "properties": {}
                }
                openai_tools.append({
                    "type": "function",
                    "function": {
                        "name": tool.name,
                        "description": tool.description or "",
                        "parameters": schema,
                    }
                })
                tool_registry[tool.name] = {
                    'session': session,
                    'server_name': server['name'],
                }
            logging.info(f"MCP connected: {server['name']} — {len(result.tools)} tools")
        except Exception as e:
            logging.warning(f"MCP connect failed for {server['name']}: {e}")

    return openai_tools, tool_registry, exit_stack


async def _agentic_loop(
    provider: dict,
    model: str,
    messages: List[dict],
    servers: List[dict],
    sid: str,
    request_id: Optional[str],
) -> tuple:
    """
    Run the LLM + MCP tool-calling loop until the model returns a final answer
    or MAX_ITERATIONS is reached.

    Returns (final LLM response dict, list of tool names actually called).
    """
    MAX_ITERATIONS = 8

    # No servers — single LLM call, no tools
    if not servers:
        return await asyncio.to_thread(_call_llm_sync, provider, model, messages), [], []

    openai_tools, tool_registry, exit_stack = await _collect_mcp_tools(servers)

    async with exit_stack:
        tools_param = openai_tools if openai_tools else None
        tools_actually_called = []
        kb_citations: List[dict] = []

        for iteration in range(MAX_ITERATIONS):
            # Call LLM (runs in thread so it doesn't block the event loop)
            response = await asyncio.to_thread(
                _call_llm_sync, provider, model, messages, tools_param
            )

            if 'error' in response:
                return response, [], []

            tool_calls = _extract_tool_calls(response, provider['type'])

            if not tool_calls:
                # No tool calls — this is the final answer
                return response, tools_actually_called, kb_citations

            # ── Tell the UI which tools are being used ──────────────────────
            tool_names = [tc['name'] for tc in tool_calls]
            tools_actually_called.extend(tool_names)
            socketio.emit('chat_status', {
                'status': 'tool_use',
                'message': f"Using: {', '.join(tool_names)}",
                'tools': tool_names,
            }, room=sid)

            # ── Append assistant's tool-call turn to history ─────────────────
            if provider['type'] == 'ollama':
                messages.append(response.get('message', {'role': 'assistant', 'content': ''}))
            else:
                choices = response.get('choices', [{}])
                if choices:
                    messages.append(choices[0].get('message', {'role': 'assistant', 'content': ''}))

            # ── Execute each tool call via MCP ───────────────────────────────
            for tc in tool_calls:
                tool_name = tc['name']
                tool_args = tc['args']
                tool_id = tc.get('id')

                if tool_name in tool_registry:
                    try:
                        mcp_session = tool_registry[tool_name]['session']
                        mcp_result = await mcp_session.call_tool(tool_name, tool_args)
                        texts = [
                            block.text for block in (mcp_result.content or [])
                            if hasattr(block, 'text')
                        ]
                        result_text = '\n'.join(texts) if texts else 'Tool returned no content.'
                        if tool_name == 'query':
                            try:
                                parsed = json.loads(result_text)
                                for c in parsed.get('citations', []):
                                    c['host_path'] = _docker_to_host_path(c.get('file_path', ''))
                                kb_citations.extend(parsed.get('citations', []))
                            except (json.JSONDecodeError, AttributeError):
                                pass
                    except Exception as e:
                        result_text = f'Tool execution error: {e}'
                else:
                    result_text = f"Tool '{tool_name}' not available."

                # Add tool result to message history
                if provider['type'] == 'ollama':
                    messages.append({'role': 'tool', 'content': result_text})
                else:
                    tool_msg: dict = {'role': 'tool', 'content': result_text}
                    if tool_id:
                        tool_msg['tool_call_id'] = tool_id
                    messages.append(tool_msg)

        # Max iterations reached — ask the model to summarise without tools
        socketio.emit('chat_status', {
            'status': 'thinking',
            'message': 'Synthesising final answer...',
        }, room=sid)
        return await asyncio.to_thread(_call_llm_sync, provider, model, messages), tools_actually_called, kb_citations


def process_chat_message(sid: str, message: str, config: dict, request_id: str = None):
    """Process chat message with LLM including conversation history (runs in background thread)"""
    try:
        # Validate config
        if not config.get('selected_llm'):
            socketio.emit('chat_response', {
                'error': 'No LLM provider selected. Please select a provider first.',
                'requestId': request_id
            }, room=sid)
            return
            
        if not config.get('selected_model'):
            socketio.emit('chat_response', {
                'error': 'No model selected. Please select a model first.',
                'requestId': request_id
            }, room=sid)
            return
        
        # Find provider
        try:
            provider = next(p for p in config['llm_providers'] if p['id'] == config['selected_llm'])
        except StopIteration:
            socketio.emit('chat_response', {
                'error': f"Provider '{config['selected_llm']}' not found in configuration.",
                'requestId': request_id
            }, room=sid)
            return
            
        model = config['selected_model']
        
        # Get enabled MCP servers for this session
        enabled_servers = [s for s in config['mcp_servers'] if s.get('enabled', False)]
        
        # Emit thinking status with rotating verbs
        thinking_verbs = ['Brewing', 'Pondering', 'Contemplating', 'Reflecting', 'Weaving', 'Synthesizing', 'Distilling']
        current_verb = 0
        
        # Emit initial thinking status
        socketio.emit('chat_status', {
            'status': 'thinking',
            'message': f'{thinking_verbs[0]} your request...',
            'tools': [s['name'] for s in enabled_servers] if enabled_servers else None
        }, room=sid)
        
        # Build message list with history
        messages = []
        
        # Add system prompt with MCP context if servers are enabled
        if enabled_servers:
            mcp_context = (
                "You are a helpful assistant with access to tools from MCP servers.\n\n"
                "TOOL USE RULES — follow these strictly:\n"
                "1. NEVER call a tool for greetings, small talk, or general questions (e.g. 'Hello', 'How are you?', 'What can you do?'). Reply directly.\n"
                "2. Only call a tool when the user EXPLICITLY asks to analyse, read, search, or process a specific file or data source.\n"
                "3. Do NOT call tools speculatively or to 'check' something — only call them when the task clearly requires it.\n"
                "4. Do not fabricate tool results — always use a real tool call when you do need data.\n\n"
                "Available MCP servers:\n"
            )
            for server in enabled_servers:
                mcp_context += f"- **{server['name']}**: {server['description']}\n"
            messages.append({'role': 'system', 'content': mcp_context})
        
        # Add project context if a project is selected
        if current_project:
            project_path = Path(current_project['path'])
            project_context = f"You are working within the project **{current_project['name']}**."
            # List files in the project (excluding metadata, max 30 for brevity)
            try:
                project_files = [
                    f.relative_to(project_path)
                    for f in sorted(project_path.rglob('*'))
                    if f.is_file() and f.name != PROJECT_META_FILENAME
                ][:30]
                if project_files:
                    project_context += "\n\nProject files:\n" + "\n".join(f"- {pf}" for pf in project_files)
            except Exception:
                pass
            messages.append({'role': 'system', 'content': project_context})

        # Add system prompt if we have files uploaded
        files = uploaded_files.get(sid, [])
        if files:
            file_context = "Files available:\n" + "\n".join([f"- {f['name']}" for f in files])
            messages.append({'role': 'system', 'content': file_context})

        # Inject active knowledge-base session so the LLM knows to use the indexer
        if sid in index_sessions:
            kb_sid = index_sessions[sid]
            kb_context = (
                f"IMPORTANT: A knowledge base is active (session_id=\"{kb_sid}\"). "
                "You MUST call the `query` tool BEFORE answering any question about documents, "
                f"papers, files, or data. Always pass session_id=\"{kb_sid}\" to the query tool. "
                "Do NOT answer from memory — retrieve from the knowledge base first. "
                "Use `list_indexed_files` to see which files are indexed. "
                "When citing sources, use [1], [2], [3] notation inline in your answer — "
                "e.g. 'Revenue grew 7% [1].' The UI will render a clickable Sources footer automatically. "
                "Do NOT write a 'References:' section yourself. "
                "KB files are LOCAL — never fabricate web URLs."
            )
            messages.append({'role': 'system', 'content': kb_context})
        
        # Add conversation history (last 20 messages to keep context manageable)
        history = chat_histories.get(sid, [])
        for msg in history[-20:]:
            messages.append(msg)
        
        # Add current user message
        messages.append({'role': 'user', 'content': message})

        # Store user message in history (initialize if server restarted)
        if sid not in chat_histories:
            chat_histories[sid] = []
        chat_histories[sid].append({'role': 'user', 'content': message})

        # Sandbox PII guard: scan user input for sensitive data before it reaches the LLM
        if config.get('compliance_mode'):
            pii_found = scan_message_for_pii(message)
            if pii_found:
                audit_logger.warning(
                    f"PII_DETECTED | types={pii_found} provider={provider['id']} model={model} session={sid}"
                )
                socketio.emit('chat_status', {
                    'status': 'warning',
                    'message': (
                        f"Sandbox: PII detected in your message "
                        f"({', '.join(pii_found).upper()}). "
                        "The Guardrail MCP can redact it — ask it to use 'redact_sensitive_data' if needed."
                    )
                }, room=sid)

        # Sandbox guard: block non-local providers in sandbox mode
        if config.get('compliance_mode') and provider['id'] not in SANDBOX_APPROVED_IDS:
            audit_logger.warning(f"BLOCKED | Attempted call to non-local provider '{provider['id']}' in sandbox mode")
            socketio.emit('chat_response', {
                'error': 'Sandbox Mode is active. Only localhost providers (Ollama) are allowed. Guardrails are enforced.',
                'requestId': request_id
            }, room=sid)
            return

        # Audit log every LLM call in sandbox mode
        if config.get('compliance_mode'):
            audit_logger.info(f"LLM_CALL | provider={provider['id']} model={model} servers={[s['name'] for s in enabled_servers]}")

        # Run the agentic loop (LLM + real MCP tool calls)
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            response, tools_called, kb_citations = loop.run_until_complete(
                _agentic_loop(provider, model, messages, enabled_servers, sid, request_id)
            )
        finally:
            loop.close()

        # Check for API error responses
        if 'error' in response:
            err = response['error']
            err_msg = err.get('message', str(err)) if isinstance(err, dict) else str(err)
            if 'does not support tools' in err_msg.lower() or 'tool_use' in err_msg.lower():
                err_msg = (f'**{model}** does not support tool use. '
                           'Switch to **qwen3.5:4b** (fast, reliable tool calls) via the model picker.')
            elif 'models' in err_msg.lower() and 'permission' in err_msg.lower():
                err_msg = ('Your GitHub token needs the **models:read** permission. '
                           'Go to github.com/settings/tokens, edit your token, '
                           'and enable models:read under Account permissions.')
            socketio.emit('chat_response', {
                'error': err_msg,
                'requestId': request_id
            }, room=sid)
            return

        # Extract and store the final text response
        content = _get_response_text(response, provider['type'])

        if sid not in chat_histories:
            chat_histories[sid] = []
        chat_histories[sid].append({'role': 'assistant', 'content': content})

        socketio.emit('chat_response', {
            'message': content,
            'model': model,
            'requestId': request_id,
            'tools_used': tools_called if tools_called else None,
            'citations': kb_citations if kb_citations else None,
        }, room=sid)

    except BaseException as e:
        import traceback
        logging.error(f"Chat error: {traceback.format_exc()}")
        # Unwrap ExceptionGroup / BaseExceptionGroup (Python 3.11 TaskGroup / anyio)
        inner = e
        if hasattr(e, 'exceptions') and e.exceptions:
            inner = e.exceptions[0]
        # Surface a human-readable cause
        import requests as _req
        if isinstance(inner, _req.exceptions.ReadTimeout):
            err_display = "The LLM timed out — it took too long to respond. Try a smaller model."
        elif isinstance(inner, (_req.exceptions.ConnectionError, ConnectionRefusedError)):
            err_display = "Cannot reach the LLM. Make sure 'ollama serve' is running."
        else:
            err_display = str(inner) or str(e)
        socketio.emit('chat_response', {
            'error': f'Error: {err_display}',
            'requestId': request_id
        }, room=sid)


@socketio.on('stop_generation')
def handle_stop_generation(data):
    """Handle stop generation request"""
    sid = request.sid
    _request_id = data.get('requestId')  # noqa: F841 - reserved for future use
    # Note: Actual cancellation would require more complex async handling
    # For now, we just acknowledge the request
    socketio.emit('chat_status', {
        'status': 'info',
        'message': 'Stopping generation...'
    }, room=sid)


@socketio.on('clear_history')
def handle_clear_history():
    """Silently clear chat + file history for this session (no UI events emitted).

    Used by the 'New Chat' button so the welcome screen isn't clobbered by
    server-sent clear_chat / chat_response events.
    """
    sid = request.sid
    if sid in chat_histories:
        chat_histories[sid] = []
    if sid in uploaded_files:
        uploaded_files[sid] = []


@socketio.on('set_index_session')
def handle_set_index_session(data):
    """Associate (or clear) an indexer session_id with this Socket.IO session."""
    sid = request.sid
    session_id = (data or {}).get('session_id', '').strip()
    if session_id:
        index_sessions[sid] = session_id
    else:
        index_sessions.pop(sid, None)


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload for the session"""
    from flask import request
    sid = request.headers.get('X-Session-ID')

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Store file info (in production, save to temp location)
    file_info = {
        'name': file.filename,
        'size': len(file.read()),
        'type': file.content_type
    }
    file.seek(0)  # Reset file pointer

    # Initialize session if needed
    if sid not in uploaded_files:
        uploaded_files[sid] = []

    uploaded_files[sid].append(file_info)

    return jsonify({
        'status': 'success',
        'file': file_info,
        'message': f'📎 Uploaded: {file.filename}'
    })


@app.route('/api/files', methods=['GET'])
def get_files():
    """Get uploaded files for the session"""
    sid = request.headers.get('X-Session-ID')
    files = uploaded_files.get(sid, [])
    return jsonify({'files': files})


# ── VS Code MCP Integration ─────────────────────────────────────────────────

@app.route('/api/vscode/config', methods=['GET'])
def get_vscode_config():
    """Return the generated .vscode/mcp.json content using ${workspaceFolder}"""
    servers = {
        "document": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "document_server.py"],
            "cwd": "${workspaceFolder}/servers/document/mcp_project"
        },
        "prompt-engineering": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "prompt_server.py"],
            "cwd": "${workspaceFolder}/servers/prompt-engineering/mcp_project"
        },
        "guardrail": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "guardrail_server.py"],
            "cwd": "${workspaceFolder}/servers/guardrail/mcp_project"
        },
        "webdesign": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "webdesign_server.py"],
            "cwd": "${workspaceFolder}/servers/webdesign/mcp_project"
        },
        "webscraper": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "webscraper_server.py"],
            "cwd": "${workspaceFolder}/servers/webscraper/mcp_project"
        }
    }
    return jsonify({"servers": servers})


@app.route('/api/vscode/setup', methods=['POST'])
def setup_vscode():
    """Write .vscode/mcp.json to the project root for VS Code + GitHub Copilot"""
    vscode_dir = PROJECT_ROOT / '.vscode'
    vscode_dir.mkdir(exist_ok=True)
    mcp_json_path = vscode_dir / 'mcp.json'

    servers = {
        "document": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "document_server.py"],
            "cwd": "${workspaceFolder}/servers/document/mcp_project"
        },
        "prompt-engineering": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "prompt_server.py"],
            "cwd": "${workspaceFolder}/servers/prompt-engineering/mcp_project"
        },
        "guardrail": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "guardrail_server.py"],
            "cwd": "${workspaceFolder}/servers/guardrail/mcp_project"
        },
        "webdesign": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "webdesign_server.py"],
            "cwd": "${workspaceFolder}/servers/webdesign/mcp_project"
        },
        "webscraper": {
            "type": "stdio",
            "command": "uv",
            "args": ["run", "webscraper_server.py"],
            "cwd": "${workspaceFolder}/servers/webscraper/mcp_project"
        }
    }
    config = {"servers": servers}

    with open(mcp_json_path, 'w') as f:
        json.dump(config, f, indent=2)

    return jsonify({
        "status": "success",
        "path": str(mcp_json_path),
        "message": f"Written to {mcp_json_path}"
    })


# ── Projects ─────────────────────────────────────────────────────────────────

PROJECTS_DIR = PROJECT_ROOT / "client" / "projects"
PROJECTS_DIR.mkdir(exist_ok=True)

# Track current project in-memory (per-session could be added later)
current_project: Optional[dict] = None

PROJECT_META_FILENAME = 'project_meta.json'


def _project_meta_path(project_path: Path) -> Path:
    return project_path / PROJECT_META_FILENAME


def _ensure_project_metadata(project_path: Path, project_name: Optional[str] = None) -> dict:
    """Ensure a stable metadata file exists for a project."""
    now = datetime.now().isoformat()
    meta_path = _project_meta_path(project_path)
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding='utf-8'))
        except Exception:
            meta = {}
    else:
        meta = {}

    if not meta.get('project_id'):
        meta['project_id'] = str(uuid.uuid4())
        meta['created_at'] = meta.get('created_at', now)

    if project_name:
        meta['name'] = project_name
    else:
        meta['name'] = meta.get('name', project_path.name)
    meta['updated_at'] = now

    meta_path.write_text(json.dumps(meta, indent=2), encoding='utf-8')
    return meta


@app.route('/api/projects', methods=['GET'])
def list_projects():
    """List all projects in the projects folder"""
    projects = []
    try:
        for item in PROJECTS_DIR.iterdir():
            if item.is_dir() and not item.name.startswith('.'):
                # Count files in project
                file_count = len(list(item.glob('**/*')))
                projects.append({
                    'name': item.name,
                    'path': str(item),
                    'file_count': file_count,
                    'created': datetime.fromtimestamp(item.stat().st_ctime).isoformat()
                })
        projects.sort(key=lambda p: p['created'], reverse=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    return jsonify({'projects': projects})


@app.route('/api/projects', methods=['POST'])
def create_project():
    """Create a new project folder"""
    global current_project
    data = request.json or {}
    name = data.get('name', '').strip()

    # Auto-generate name if not provided
    if not name:
        date = datetime.now()
        name = f"project-{date.year}-{date.month:02d}-{date.day:02d}"
        # Ensure unique
        counter = 1
        original_name = name
        while (PROJECTS_DIR / name).exists():
            name = f"{original_name}-{counter}"
            counter += 1

    # Validate name
    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400

    project_path = PROJECTS_DIR / name

    # Check if exists
    if project_path.exists():
        # Return existing project info
        file_count = len(list(project_path.glob('**/*')))
        meta = _ensure_project_metadata(project_path, name)
        current_project = {
            'project_id': meta.get('project_id'),
            'name': name,
            'path': str(project_path),
            'file_count': file_count,
            'kb_session_id': meta.get('kb_session_id'),
            'kb_folder_path': meta.get('kb_folder_path'),
        }
        return jsonify({
            'status': 'exists',
            'project': current_project
        })

    # Create project folder
    try:
        project_path.mkdir(parents=True, exist_ok=True)
        # Create subfolders for organization
        (project_path / 'input').mkdir(exist_ok=True)
        (project_path / 'output').mkdir(exist_ok=True)
        (project_path / 'scripts').mkdir(exist_ok=True)
        meta = _ensure_project_metadata(project_path, name)

        current_project = {
            'project_id': meta.get('project_id'),
            'name': name,
            'path': str(project_path),
            'file_count': 0,
            'kb_session_id': meta.get('kb_session_id'),
            'kb_folder_path': meta.get('kb_folder_path'),
        }

        return jsonify({
            'status': 'created',
            'project': current_project
        })
    except Exception as e:
        return jsonify({'error': f'Failed to create project: {str(e)}'}), 500


@app.route('/api/projects/current', methods=['GET'])
def get_current_project():
    """Get the currently active project"""
    if current_project:
        return jsonify({'project': current_project})
    return jsonify({'project': None}), 404


@app.route('/api/projects/select', methods=['POST'])
def select_project():
    """Select an existing project as current"""
    global current_project
    data = request.json or {}
    name = data.get('name', '').strip()

    if not name:
        return jsonify({'error': 'Project name required'}), 400

    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    file_count = len(list(project_path.glob('**/*')))
    meta = _ensure_project_metadata(project_path, name)
    current_project = {
        'project_id': meta.get('project_id'),
        'name': name,
        'path': str(project_path),
        'file_count': file_count,
        'kb_session_id': meta.get('kb_session_id'),
        'kb_folder_path': meta.get('kb_folder_path'),
    }

    return jsonify({
        'status': 'selected',
        'project': current_project
    })


@app.route('/api/projects/<name>', methods=['PUT'])
def rename_project(name):
    """Rename an existing project"""
    global current_project
    data = request.json or {}
    new_name = data.get('new_name', '').strip()

    if not new_name:
        return jsonify({'error': 'new_name is required'}), 400
    if '/' in new_name or '\\' in new_name or '..' in new_name:
        return jsonify({'error': 'Invalid project name'}), 400

    old_path = PROJECTS_DIR / name
    new_path = PROJECTS_DIR / new_name

    if not old_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404
    if new_path.exists():
        return jsonify({'error': f'Project already exists: {new_name}'}), 409

    try:
        old_path.rename(new_path)
        meta = _ensure_project_metadata(new_path, new_name)
        file_count = len(list(new_path.glob('**/*')))
        project_info = {
            'project_id': meta.get('project_id'),
            'name': new_name,
            'path': str(new_path),
            'file_count': file_count
        }
        if current_project and current_project['name'] == name:
            current_project = project_info
        return jsonify({'status': 'renamed', 'project': project_info})
    except Exception as e:
        return jsonify({'error': f'Failed to rename: {str(e)}'}), 500


@app.route('/api/projects/<name>', methods=['DELETE'])
def delete_project(name):
    """Delete a project and all its contents"""
    global current_project

    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400

    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    try:
        import shutil
        shutil.rmtree(project_path)
        if current_project and current_project['name'] == name:
            current_project = None
        return jsonify({'status': 'deleted', 'name': name})
    except Exception as e:
        return jsonify({'error': f'Failed to delete: {str(e)}'}), 500


# ── Project Virtual Environment Management ───────────────────────────────────

import hashlib
import subprocess


def _get_project_venv_path(project_path: str) -> Path:
    """Get the .venv path for a project."""
    return Path(project_path) / '.venv'

REQUIRED_PROJECT_REQUIREMENTS = [
    # Core app/runtime
    'requests>=2.31.0',
    'python-dotenv>=1.0.0',
    # Web scraping stack
    'httpx>=0.27.0',
    'beautifulsoup4>=4.12',
    'lxml>=5.0',
    'firecrawl-py>=1.0.0',
    # PDF extraction — pypdf only (AGPL pymupdf/pymupdf4llm excluded for license compliance)
    'pypdf>=4.0.0',
    # Excel extraction stack
    'openpyxl>=3.1.0',
]


def _ensure_project_requirements(project_path: Path) -> Path:
    """Create/merge requirements.txt with all required tooling."""
    req_path = project_path / 'requirements.txt'
    existing_lines = []
    if req_path.exists():
        existing_lines = [line.strip() for line in req_path.read_text(encoding='utf-8').splitlines()]
    existing_nonempty = {line for line in existing_lines if line and not line.startswith('#')}

    merged = list(existing_lines)
    for dep in REQUIRED_PROJECT_REQUIREMENTS:
        dep_name = dep.split('>=')[0].split('==')[0].strip().lower()
        if not any(
            line.lower().split('>=')[0].split('==')[0].strip() == dep_name
            for line in existing_nonempty
        ):
            merged.append(dep)

    req_path.write_text('\n'.join([line for line in merged if line]).strip() + '\n', encoding='utf-8')
    return req_path


def _ensure_scrape_project_environment(job_name: str, project_path_hint: str = '') -> dict:
    """Ensure scrape project folder, requirements, and .venv are ready/up to date."""
    safe_name = (job_name or '').strip() or f"scrape-{datetime.now().strftime('%Y-%m-%d-%H%M%S')}"
    safe_name = ''.join(c if c.isalnum() or c in ('-', '_', ' ') else '-' for c in safe_name).strip()
    safe_name = safe_name.replace(' ', '-')
    if not safe_name:
        safe_name = f"scrape-{datetime.now().strftime('%Y-%m-%d-%H%M%S')}"

    project_path = Path(project_path_hint).expanduser().resolve() if project_path_hint else (PROJECTS_DIR / safe_name)
    project_path.mkdir(parents=True, exist_ok=True)
    (project_path / 'input').mkdir(exist_ok=True)
    (project_path / 'output').mkdir(exist_ok=True)
    (project_path / 'scripts').mkdir(exist_ok=True)
    meta = _ensure_project_metadata(project_path, project_path.name)

    req_path = _ensure_project_requirements(project_path)

    status_before = _get_venv_status(str(project_path))
    venv_path = _get_project_venv_path(str(project_path))
    pip_path = venv_path / ('Scripts' if os.name == 'nt' else 'bin') / 'pip'

    if not venv_path.exists():
        subprocess.run([sys.executable, '-m', 'venv', str(venv_path)], check=True, capture_output=True, timeout=120)

    # Install/update only when venv is new or requirements changed.
    if (not status_before['exists']) or status_before['needs_update']:
        subprocess.run([str(pip_path), 'install', '--upgrade', 'pip'], check=True, capture_output=True, timeout=90)
        subprocess.run([str(pip_path), 'install', '-r', str(req_path)], check=True, capture_output=True, text=True, timeout=480)
        current_hash = hashlib.md5(req_path.read_bytes()).hexdigest()
        (venv_path / '.requirements_hash').write_text(current_hash)

    return {
        'project_id': meta.get('project_id'),
        'project_name': project_path.name,
        'project_path': str(project_path),
        'requirements_path': str(req_path),
        'venv_path': str(venv_path),
    }


def _get_venv_status(project_path: str) -> dict:
    """Check the status of a project's virtual environment."""
    venv_path = _get_project_venv_path(project_path)
    req_path = Path(project_path) / 'requirements.txt'

    exists = venv_path.exists()
    has_requirements = req_path.exists()

    status = {
        'exists': exists,
        'path': str(venv_path),
        'requirements_exists': has_requirements,
        'requirements_path': str(req_path) if has_requirements else None,
        'needs_update': False,
        'installed_packages': [],
        'message': ''
    }

    if not exists:
        status['message'] = 'Virtual environment not found. Will create on first use.'
        return status

    # Check if requirements.txt has changed
    if has_requirements:
        req_hash_path = venv_path / '.requirements_hash'
        current_hash = hashlib.md5(req_path.read_bytes()).hexdigest()

        if req_hash_path.exists():
            stored_hash = req_hash_path.read_text().strip()
            if stored_hash != current_hash:
                status['needs_update'] = True
                status['message'] = 'requirements.txt has changed. Update needed.'
        else:
            status['needs_update'] = True
            status['message'] = 'Requirements hash not found. Update needed.'

    # List installed packages
    try:
        pip_path = venv_path / ('Scripts' if os.name == 'nt' else 'bin') / 'pip'
        result = subprocess.run(
            [str(pip_path), 'freeze'],
            capture_output=True,
            text=True,
            timeout=10
        )
        if result.returncode == 0:
            status['installed_packages'] = [
                line.strip() for line in result.stdout.split('\n') if line.strip()
            ]
    except Exception:
        pass

    if not status['message']:
        status['message'] = 'Virtual environment is ready.'

    return status


@app.route('/api/projects/<name>/venv', methods=['GET'])
def get_project_venv_status(name):
    """Get virtual environment status for a project."""
    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400

    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    return jsonify(_get_venv_status(str(project_path)))


@app.route('/api/projects/<name>/venv', methods=['POST'])
def create_or_update_project_venv(name):
    """Create or update a project's virtual environment."""
    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400

    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    data = request.json or {}
    auto_install = data.get('auto_install', True)

    venv_path = _get_project_venv_path(str(project_path))
    req_path = project_path / 'requirements.txt'

    # Create venv if it doesn't exist
    if not venv_path.exists():
        try:
            subprocess.run(
                [sys.executable, '-m', 'venv', str(venv_path)],
                check=True,
                capture_output=True,
                timeout=120
            )
        except subprocess.CalledProcessError as e:
            return jsonify({
                'error': f'Failed to create virtual environment: {e.stderr.decode() if e.stderr else str(e)}'
            }), 500
        except subprocess.TimeoutExpired:
            return jsonify({'error': 'Virtual environment creation timed out'}), 500

    # Install/update requirements if they exist
    if auto_install and req_path.exists():
        pip_path = venv_path / ('Scripts' if os.name == 'nt' else 'bin') / 'pip'
        try:
            # Upgrade pip first
            subprocess.run(
                [str(pip_path), 'install', '--upgrade', 'pip'],
                check=True,
                capture_output=True,
                timeout=60
            )

            # Install requirements
            result = subprocess.run(
                [str(pip_path), 'install', '-r', str(req_path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=300
            )

            # Store requirements hash
            current_hash = hashlib.md5(req_path.read_bytes()).hexdigest()
            (venv_path / '.requirements_hash').write_text(current_hash)

            return jsonify({
                'status': 'venv_ready',
                'message': 'Virtual environment created and dependencies installed',
                'path': str(venv_path),
                'install_output': result.stdout[-1000:] if result.stdout else ''  # Last 1000 chars
            })

        except subprocess.CalledProcessError as e:
            return jsonify({
                'error': f'Failed to install dependencies: {e.stderr if e.stderr else str(e)}',
                'partial': True,
                'path': str(venv_path)
            }), 500
        except subprocess.TimeoutExpired:
            return jsonify({
                'error': 'Dependency installation timed out',
                'partial': True,
                'path': str(venv_path)
            }), 500

    return jsonify({
        'status': 'venv_ready',
        'message': 'Virtual environment ready',
        'path': str(venv_path)
    })


@app.route('/api/projects/<name>/venv/run', methods=['POST'])
def run_in_project_venv(name):
    """Run a Python command in the project's virtual environment."""
    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400

    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    data = request.json or {}
    command = data.get('command', '').strip()
    working_dir = data.get('working_dir', str(project_path))

    if not command:
        return jsonify({'error': 'Command is required'}), 400

    venv_path = _get_project_venv_path(str(project_path))
    if not venv_path.exists():
        # Auto-create venv
        try:
            subprocess.run(
                [sys.executable, '-m', 'venv', str(venv_path)],
                check=True,
                capture_output=True,
                timeout=120
            )
        except Exception as e:
            return jsonify({'error': f'Failed to create virtual environment: {str(e)}'}), 500

    # Get Python executable from venv
    python_path = venv_path / ('Scripts' if os.name == 'nt' else 'bin') / 'python'

    try:
        result = subprocess.run(
            [str(python_path), '-c', command],
            capture_output=True,
            text=True,
            timeout=60,
            cwd=working_dir
        )
        return jsonify({
            'returncode': result.returncode,
            'stdout': result.stdout,
            'stderr': result.stderr
        })
    except subprocess.TimeoutExpired:
        return jsonify({'error': 'Command timed out'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Folder Browser ────────────────────────────────────────────────────────────

@app.route('/api/browse', methods=['GET'])
def browse_folder():
    """Browse directories on the local filesystem for the folder picker.
    
    Query params:
      path (optional)       – directory to list.  Defaults to the user home dir.
      show_hidden (optional) – '1' to include dotfiles / hidden dirs.
    Returns JSON with parent path and list of child directories.
    """
    raw = request.args.get('path', '').strip()
    show_hidden = request.args.get('show_hidden', '') == '1'

    if not raw:
        target = Path.home()
    else:
        target = Path(raw).expanduser().resolve()

    if not target.exists():
        return jsonify({'error': f'Path not found: {raw}'}), 404
    if not target.is_dir():
        return jsonify({'error': f'Not a directory: {raw}'}), 400

    # Security: prevent traversal outside real directories
    target = target.resolve()

    dirs = []
    try:
        for child in sorted(target.iterdir()):
            try:
                if not child.is_dir():
                    continue
            except (PermissionError, OSError):
                continue
            if not show_hidden and child.name.startswith('.'):
                continue
            dirs.append(child.name)
    except PermissionError:
        pass

    return jsonify({
        'current': str(target),
        'parent': str(target.parent) if target != target.parent else None,
        'directories': dirs
    })


@app.route('/api/browse/quickaccess', methods=['GET'])
def browse_quickaccess():
    """Return well-known quick-access directories for the current OS."""
    home = Path.home()
    places = []

    # Standard directories
    candidates = [
        ('🏠 Home', home),
        ('🖥️ Desktop', home / 'Desktop'),
        ('📄 Documents', home / 'Documents'),
        ('⬇️ Downloads', home / 'Downloads'),
    ]

    if sys.platform == 'darwin':
        # macOS-specific common folders
        candidates += [
            ('🖼️ Pictures', home / 'Pictures'),
            ('🎵 Music', home / 'Music'),
            ('🎬 Movies', home / 'Movies'),
            ('📦 Applications', Path('/Applications')),
        ]
        # iCloud Drive
        icloud = home / 'Library' / 'Mobile Documents' / 'com~apple~CloudDocs'
        if icloud.exists():
            candidates.append(('☁️ iCloud Drive', icloud))
        # Mounted volumes (external drives, network shares, etc.)
        volumes = Path('/Volumes')
        if volumes.exists():
            try:
                for vol in sorted(volumes.iterdir()):
                    if vol.is_dir() and not vol.name.startswith('.'):
                        candidates.append((f'💾 {vol.name}', vol))
            except PermissionError:
                pass

    # Detect OneDrive directories (Windows patterns)
    try:
        for child in home.iterdir():
            if child.is_dir() and 'onedrive' in child.name.lower():
                # Skip OneDrive folders that are effectively empty (only desktop.ini)
                try:
                    real_children = [c for c in child.iterdir() if c.name != 'desktop.ini']
                    if not real_children:
                        continue
                except PermissionError:
                    continue
                label = '☁️ ' + child.name
                candidates.append((label, child))
    except PermissionError:
        pass

    # Windows drives
    if os.name == 'nt':
        import string
        for letter in string.ascii_uppercase:
            drive = Path(f'{letter}:\\')
            if drive.exists():
                candidates.append((f'💾 {letter}:\\', drive))

    for label, p in candidates:
        if p.exists() and p.is_dir():
            places.append({'label': label, 'path': str(p)})

    return jsonify({'places': places})


# ── Batch Folder Processing ──────────────────────────────────────────────────

# Manifest file name stored inside each output folder
_MANIFEST_NAME = '.batch_manifest.json'


def _load_manifest(output_path: Path) -> dict:
    """Load the processing manifest from a previous run (if any)."""
    manifest_path = output_path / _MANIFEST_NAME
    if manifest_path.exists():
        with open(manifest_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def _save_manifest(output_path: Path, manifest: dict):
    """Persist the processing manifest."""
    manifest_path = output_path / _MANIFEST_NAME
    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)


def _file_mtime_iso(file_path: Path) -> str:
    """Return a file's last-modified time as an ISO string."""
    return datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()


@app.route('/api/batch_process', methods=['POST'])
def batch_process():
    """
    Start batch processing of a folder containing PDFs, Excel, and Word files.
    Emits socket.io 'batch_progress' events with job_id for real-time progress.
    Output: one JSON file per input file in output_folder.
    """
    data = request.json
    source_folder = data.get('source_folder', '').strip()
    output_folder = data.get('output_folder', '').strip()
    incremental = data.get('incremental', True)  # skip unchanged files by default
    prompt = data.get('prompt', '').strip()
    job_id = str(uuid.uuid4())[:8]

    if not source_folder:
        return jsonify({'error': 'source_folder is required'}), 400

    source_path = Path(source_folder).expanduser().resolve()
    if not source_path.exists():
        return jsonify({'error': f'Source folder not found: {source_folder}'}), 404
    if not source_path.is_dir():
        return jsonify({'error': f'Not a directory: {source_folder}'}), 400

    # Default output folder: <source_folder>_processed
    if not output_folder:
        output_path = source_path.parent / f"{source_path.name}_processed"
    else:
        output_path = Path(output_folder).expanduser().resolve()

    output_path.mkdir(parents=True, exist_ok=True)

    # Collect all supported document types
    files_to_process = (
        list(source_path.glob('**/*.pdf')) +
        list(source_path.glob('**/*.xlsx')) +
        list(source_path.glob('**/*.xls')) +
        list(source_path.glob('**/*.docx'))
    )

    if not files_to_process:
        return jsonify({'error': 'No matching files found in source folder (PDF, Excel, Word — searched recursively)'}), 404

    # Count subfolders containing matched files
    subfolders = {str(f.parent.relative_to(source_path)) for f in files_to_process if f.parent != source_path}

    # Launch background processing thread
    threading.Thread(
        target=_run_batch_processing,
        args=(job_id, files_to_process, source_path, output_path, incremental, prompt),
        daemon=True
    ).start()

    return jsonify({
        'status': 'started',
        'job_id': job_id,
        'total_files': len(files_to_process),
        'subfolders': len(subfolders),
        'output_folder': str(output_path)
    })


def _run_batch_processing(job_id: str, files: list, source_path: Path, output_path: Path, incremental: bool = True, prompt: str = ''):
    """Background thread: process files and emit progress events.
    
    When *incremental* is True the manifest from the previous run is
    consulted and files whose last-modified timestamp has not changed are
    skipped automatically.
    """
    manifest = _load_manifest(output_path) if incremental else {}
    total = len(files)
    done = 0
    skipped = 0
    errors = 0

    socketio.emit('batch_progress', {
        'job_id': job_id,
        'status': 'started',
        'total': total,
        'done': 0,
        'message': f'Processing {total} files...'
    })

    for file_path in files:
        file_path = Path(file_path)
        rel_key = str(file_path.relative_to(source_path))
        current_mtime = _file_mtime_iso(file_path)

        # Incremental: skip if file hasn't changed since last run
        if incremental and rel_key in manifest:
            prev = manifest[rel_key]
            if prev.get('mtime') == current_mtime:
                skipped += 1
                socketio.emit('batch_progress', {
                    'job_id': job_id,
                    'status': 'progress',
                    'total': total,
                    'done': done,
                    'skipped': skipped,
                    'errors': errors,
                    'current_file': rel_key,
                    'message': f'[skip] {rel_key} (unchanged)'
                })
                continue

        try:
            result = _extract_file(file_path)
            if prompt:
                result['processing_prompt'] = prompt
            # Preserve relative folder structure in output
            rel = file_path.relative_to(source_path)
            out_file = output_path / rel.with_suffix('.json')
            out_file.parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            done += 1
            scanned = result.get('scanned_pages', [])
            pages = result.get('pages', '?')
            ocr_note = f', {len(scanned)} scanned pages' if scanned else ''
            ocr_hint = result.get('ocr_hint')

            # Update manifest entry
            manifest[rel_key] = {
                'filename': file_path.name,
                'mtime': current_mtime,
                'processed_at': datetime.now().isoformat(),
                'output': str(out_file)
            }

            socketio.emit('batch_progress', {
                'job_id': job_id,
                'status': 'progress',
                'total': total,
                'done': done,
                'skipped': skipped,
                'errors': errors,
                'current_file': file_path.name,
                'scanned_pages': scanned,
                'ocr_hint': ocr_hint,
                'message': f'[{done}/{total}] {file_path.name} — {pages} pages{ocr_note}'
            })
            if ocr_hint and scanned:
                socketio.emit('batch_progress', {
                    'job_id': job_id,
                    'status': 'ocr_hint',
                    'total': total,
                    'done': done,
                    'errors': errors,
                    'message': f'  ↳ Tip: {ocr_hint}'
                })
        except Exception as e:
            errors += 1
            socketio.emit('batch_progress', {
                'job_id': job_id,
                'status': 'error',
                'total': total,
                'done': done,
                'errors': errors,
                'current_file': rel_key,
                'message': f'Error: {rel_key} — {str(e)}'
            })

    # Persist updated manifest
    _save_manifest(output_path, manifest)

    skip_msg = f', {skipped} skipped (unchanged)' if skipped else ''
    socketio.emit('batch_progress', {
        'job_id': job_id,
        'status': 'complete',
        'total': total,
        'done': done,
        'skipped': skipped,
        'errors': errors,
        'output_folder': str(output_path),
        'message': f'Done. {done} processed, {errors} errors{skip_msg}. Output: {output_path}'
    })


# ── Knowledge-Base Indexing Pipeline ─────────────────────────────────────────

def _simple_word_chunks(text: str, chunk_size: int = 512, overlap: int = 50) -> list[dict]:
    """Split plain text into overlapping word-window chunks."""
    words = text.split()
    chunks = []
    step = max(1, chunk_size - overlap)
    for start in range(0, max(1, len(words)), step):
        batch = words[start: start + chunk_size]
        if not batch:
            break
        chunks.append({
            "content": " ".join(batch),
            "source_ref": f"words {start + 1}–{start + len(batch)}",
        })
    return chunks


async def _extract_chunks_via_mcp(file_path: Path, tool_registry: dict) -> list[dict]:
    """Call the document MCP server to extract RAG chunks for a single file."""
    ext = file_path.suffix.lower()
    fp = str(file_path)

    async def _call(tool: str, args: dict) -> dict:
        reg = tool_registry.get(tool)
        if not reg:
            return {}
        result = await reg['session'].call_tool(tool, args)
        texts = [b.text for b in (result.content or []) if hasattr(b, 'text')]
        if not texts:
            return {}
        try:
            return json.loads(texts[0])
        except Exception:
            return {}

    chunks: list[dict] = []

    if ext == '.pdf':
        data = await _call('chunk_pdf_for_rag', {'file_path': fp})
        for c in data.get('chunks', []):
            text = c.get('text', '').strip()
            if text:
                chunks.append({
                    'content': text,
                    'source_ref': (
                        f"pages {c['page_start']}–{c['page_end']}"
                        if 'page_start' in c
                        else f"chunk {c.get('chunk_index', 0) + 1}"
                    ),
                })

    elif ext in {'.xlsx', '.xls', '.xlsm'}:
        data = await _call('chunk_excel_content', {'file_path': fp})
        for c in data.get('chunks', []):
            text = c.get('text', '').strip()
            if text:
                page = (c.get('metadata') or {}).get('page_number')
                ref = f"page {page}" if page else f"chunk {c.get('chunk_number', 1)}"
                chunks.append({'content': text, 'source_ref': ref})

    elif ext == '.csv':
        data = await _call('chunk_csv_for_rag', {'file_path': fp})
        for c in data.get('chunks', []):
            if c.get('content', '').strip():
                chunks.append({'content': c['content'], 'source_ref': c.get('source_ref', '')})

    elif ext in {'.txt', '.md'}:
        data = await _call('chunk_text_for_rag', {'file_path': fp})
        for c in data.get('chunks', []):
            if c.get('content', '').strip():
                chunks.append({'content': c['content'], 'source_ref': c.get('source_ref', '')})

    elif ext in {'.docx', '.doc'}:
        data = await _call('word_to_markdown', {'file_path': fp})
        markdown = data.get('markdown', '').strip()
        if markdown:
            chunks = _simple_word_chunks(markdown)

    elif ext == '.parquet':
        data = await _call('read_parquet', {'file_path': fp, 'row_limit': 2000})
        # read_parquet returns a markdown table or JSON; treat as one text block
        content = data.get('content', '') or data.get('markdown', '') or str(data)
        content = content.strip()
        if content:
            chunks = _simple_word_chunks(content)

    return chunks


async def _index_folder_async(folder_path: str, session_id: str, sid: str) -> None:
    """Connect to document + indexer MCP servers, walk folder, index all files."""
    folder = Path(folder_path)
    files = []
    for f in sorted(folder.rglob('*')):
        try:
            if (f.is_file()
                    and f.suffix.lower() in _INDEXABLE_EXTENSIONS
                    and not f.name.startswith('.')
                    and '.venv' not in f.parts
                    and '__pycache__' not in f.parts):
                files.append(f)
        except (PermissionError, OSError):
            continue

    if not files:
        socketio.emit('index_complete', {
            'session_id': session_id,
            'files_indexed': 0,
            'total_chunks': 0,
            'error': 'No supported files found in the selected folder.',
        }, room=sid)
        return

    doc_server = next((s for s in MCP_SERVERS if s.id == 'document_mcp'), None)
    idx_server = next((s for s in MCP_SERVERS if s.id == 'indexer_mcp'), None)
    if not doc_server or not idx_server:
        socketio.emit('index_complete', {
            'session_id': session_id, 'error': 'document_mcp or indexer_mcp not configured.',
        }, room=sid)
        return

    _, tool_registry, exit_stack = await _collect_mcp_tools([
        asdict(doc_server), asdict(idx_server)
    ])

    async with exit_stack:
        # Create indexer session
        cs = tool_registry.get('create_session')
        if not cs:
            socketio.emit('index_complete', {
                'session_id': session_id, 'error': 'Indexer MCP not connected.',
            }, room=sid)
            return

        await cs['session'].call_tool('create_session', {
            'session_id': session_id,
            'metadata': json.dumps({'folder': str(folder)}),
        })

        total_chunks = 0
        files_done = 0
        errors = []

        for file_path in files:
            socketio.emit('index_progress', {
                'session_id': session_id,
                'file': file_path.name,
                'done': files_done,
                'total': len(files),
                'chunks_so_far': total_chunks,
            }, room=sid)

            try:
                chunks = await _extract_chunks_via_mcp(file_path, tool_registry)
                if chunks:
                    ic = tool_registry.get('index_chunks')
                    if ic:
                        result = await ic['session'].call_tool('index_chunks', {
                            'session_id': session_id,
                            'file_path': str(file_path),
                            'chunks': chunks,
                        })
                        texts = [b.text for b in (result.content or []) if hasattr(b, 'text')]
                        data = json.loads(texts[0]) if texts else {}
                        total_chunks += data.get('indexed', 0)
            except Exception as exc:
                errors.append(f"{file_path.name}: {exc}")
                logging.warning(f"Index error for {file_path.name}: {exc}")

            files_done += 1

    _save_kb_session(folder_path, session_id, files_done, total_chunks)
    socketio.emit('index_complete', {
        'session_id': session_id,
        'files_indexed': files_done,
        'total_chunks': total_chunks,
        'errors': errors,
    }, room=sid)


def _index_folder_bg(folder_path: str, session_id: str, sid: str) -> None:
    """Thread entry point: run the async indexing pipeline."""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        loop.run_until_complete(_index_folder_async(folder_path, session_id, sid))
    except Exception as exc:
        socketio.emit('index_complete', {
            'session_id': session_id,
            'error': str(exc),
        }, room=sid)
    finally:
        loop.close()


@app.route('/api/index/start', methods=['POST'])
def index_start():
    """Start background indexing of a folder into the knowledge base.

    Body JSON: { "folder_path": "/abs/path", "sid": "<socket-io-sid>" }
    Returns: { "session_id": "<uuid>", "status": "started", "total_files": N }
    """
    data = request.json or {}
    folder_path = data.get('folder_path', '').strip()
    sid = data.get('sid', '').strip()

    if not folder_path:
        return jsonify({'error': 'folder_path is required'}), 400

    folder = Path(folder_path).expanduser().resolve()
    if not folder.exists() or not folder.is_dir():
        return jsonify({'error': f'Not a directory: {folder_path}'}), 400

    # Count files for the immediate response
    files = []
    for f in folder.rglob('*'):
        try:
            if (f.is_file()
                    and f.suffix.lower() in _INDEXABLE_EXTENSIONS
                    and not f.name.startswith('.')
                    and '.venv' not in f.parts
                    and '__pycache__' not in f.parts):
                files.append(f)
        except (PermissionError, OSError):
            continue

    session_id = str(uuid.uuid4())

    threading.Thread(
        target=_index_folder_bg,
        args=(str(folder), session_id, sid),
        daemon=True,
    ).start()

    return jsonify({
        'session_id': session_id,
        'status': 'started',
        'total_files': len(files),
        'folder': str(folder),
    })


@app.route('/api/index/status', methods=['GET'])
def index_status():
    """Return the indexer session currently active for a given Socket.IO sid."""
    sid = request.args.get('sid', '').strip()
    session_id = index_sessions.get(sid, '')
    return jsonify({'session_id': session_id, 'active': bool(session_id)})


@app.route('/api/kb/sessions', methods=['GET'])
def kb_sessions_list():
    """Return all persisted KB sessions (folder → session_id mappings)."""
    return jsonify(_load_kb_sessions())


@app.route('/api/kb/sessions/<session_id>', methods=['DELETE'])
def kb_sessions_delete(session_id):
    """Remove a KB session from the persistent registry."""
    sessions = [s for s in _load_kb_sessions() if s.get('session_id') != session_id]
    KB_SESSIONS_FILE.write_text(json.dumps(sessions, indent=2))
    return jsonify({'ok': True})


@app.route('/api/projects/<name>/kb', methods=['POST', 'DELETE'])
def project_kb_link(name):
    """Link or unlink a KB session from a project's metadata."""
    if '/' in name or '\\' in name or '..' in name:
        return jsonify({'error': 'Invalid project name'}), 400
    project_path = PROJECTS_DIR / name
    if not project_path.exists():
        return jsonify({'error': f'Project not found: {name}'}), 404

    meta_path = _project_meta_path(project_path)
    try:
        meta = json.loads(meta_path.read_text(encoding='utf-8')) if meta_path.exists() else {}
    except Exception:
        meta = {}

    if request.method == 'POST':
        data = request.json or {}
        session_id = data.get('session_id', '').strip()
        folder_path = data.get('folder_path', '').strip()
        if not session_id:
            return jsonify({'error': 'session_id required'}), 400
        meta['kb_session_id'] = session_id
        meta['kb_folder_path'] = folder_path
    else:  # DELETE
        meta.pop('kb_session_id', None)
        meta.pop('kb_folder_path', None)

    meta['updated_at'] = datetime.now().isoformat()
    meta_path.write_text(json.dumps(meta, indent=2), encoding='utf-8')

    # Keep current_project in sync
    global current_project
    if current_project and current_project.get('name') == name:
        current_project['kb_session_id'] = meta.get('kb_session_id')
        current_project['kb_folder_path'] = meta.get('kb_folder_path')

    return jsonify({'ok': True, 'kb_session_id': meta.get('kb_session_id')})


@app.route('/api/prepare_vscode', methods=['POST'])
def prepare_vscode():
    """
    Prepare a Python script for batch processing in VS Code.
    Returns the path to the generated script and configuration.
    """
    data = request.json
    source_folder = data.get('source_folder', '').strip()
    output_folder = data.get('output_folder', '').strip()
    project_path_str = data.get('project_path', '').strip()
    prompt = data.get('prompt', '').strip()

    if not source_folder:
        return jsonify({'error': 'source_folder is required'}), 400

    source_path = Path(source_folder).expanduser().resolve()
    if not source_path.exists():
        return jsonify({'error': f'Source folder not found: {source_folder}'}), 404
    if not source_path.is_dir():
        return jsonify({'error': f'Not a directory: {source_folder}'}), 400

    # Determine output and scripts paths based on project structure
    if project_path_str:
        project_path = Path(project_path_str).expanduser().resolve()
        output_path = project_path / 'output'
        scripts_path = project_path / 'scripts'
    elif output_folder:
        output_path = Path(output_folder).expanduser().resolve()
        scripts_path = output_path
    else:
        output_path = source_path.parent / f"{source_path.name}_processed"
        scripts_path = output_path

    output_path.mkdir(parents=True, exist_ok=True)
    scripts_path.mkdir(parents=True, exist_ok=True)

    # Collect all supported document types
    files_to_process = (
        list(source_path.glob('**/*.pdf')) +
        list(source_path.glob('**/*.xlsx')) +
        list(source_path.glob('**/*.xls')) +
        list(source_path.glob('**/*.docx'))
    )

    if not files_to_process:
        return jsonify({'error': 'No matching files found in source folder (PDF, Excel, Word — searched recursively)'}), 404

    # Count subfolders containing matched files
    subfolders = {str(f.parent.relative_to(source_path)) for f in files_to_process if f.parent != source_path}

    # Generate Python script for VS Code
    script_content = _generate_vscode_script(source_path, output_path, files_to_process, scripts_path, prompt)
    script_path = scripts_path / 'batch_process.py'
    with open(script_path, 'w', encoding='utf-8') as f:
        f.write(script_content)

    # Generate configuration file
    config = {
        'source_folder': str(source_path),
        'output_folder': str(output_path),
        'total_files': len(files_to_process),
        'subfolders': sorted(subfolders) if subfolders else [],
        'files': [str(f.relative_to(source_path)) for f in files_to_process],
        'generated_at': datetime.now().isoformat(),
        'instructions': f'Run: python "{script_path}"'
    }
    if prompt:
        config['processing_prompt'] = prompt
    config_path = scripts_path / 'batch_config.json'
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2)

    subfolder_msg = f' across {len(subfolders)} subfolder(s)' if subfolders else ''
    return jsonify({
        'status': 'ready',
        'script_path': str(script_path),
        'config_path': str(config_path),
        'output_folder': str(output_path),
        'scripts_folder': str(scripts_path),
        'project_path': project_path_str or '',
        'total_files': len(files_to_process),
        'subfolders': len(subfolders),
        'message': f'Ready for VS Code: {len(files_to_process)} files prepared{subfolder_msg}'
    })


def _generate_vscode_script(source_path: Path, output_path: Path, files: list, scripts_path: Path = None, prompt: str = '') -> str:
    """Generate a standalone Python script for batch processing in VS Code."""
    # Show relative paths (including subfolder) for each file in the preview list
    file_list_str = ',\n        '.join(
        f'"{f.relative_to(source_path)}"' for f in files[:20]
    )
    if len(files) > 20:
        file_list_str += f',\n        # ... and {len(files) - 20} more files'

    # pymupdf/pymupdf4llm are AGPL v3 (not Apache 2.0 compatible) — use pypdf instead
    requirements = ['pypdf', 'openpyxl', 'python-docx']

    # Write requirements.txt next to the script
    req_path = (scripts_path or output_path) / 'requirements.txt'
    with open(req_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(requirements) + '\n')

    prompt_line = f'PROCESSING_PROMPT = {repr(prompt)}' if prompt else 'PROCESSING_PROMPT = ""  # No prompt specified'
    _datetime_now = datetime.now().isoformat()

    script = '''#!/usr/bin/env python3
"""
Batch Processing Script — Generated by MCP Client UI
=====================================================
Process PDF, Excel, and Word files for LLM/RAG pipelines.

This script was auto-generated for VS Code processing.
Run it directly or step through with the VS Code debugger.

Generated: __DATETIME_NOW__
"""

import json
import subprocess
import sys
from pathlib import Path
from typing import Dict, List, Any

# Configuration (auto-generated)
SOURCE_FOLDER = Path(r"__SOURCE_PATH__")
OUTPUT_FOLDER = Path(r"__OUTPUT_PATH__")
SCRIPT_DIR = Path(__file__).parent
__PROMPT_LINE__

# Files to process
FILES_TO_PROCESS: List[Path] = [
    __FILE_LIST_STR__
]
# Note: Full file list is loaded dynamically below

import shutil as _shutil

def _check_tesseract() -> bool:
    return _shutil.which('tesseract') is not None

def extract_pdf(file_path: Path) -> Dict[str, Any]:
    """PDF extraction using pypdf (BSD 3-Clause — Apache 2.0 compatible)."""
    from pypdf import PdfReader
    import datetime

    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    meta = reader.metadata or {}

    page_texts = []
    page_info = []
    scanned_pages = []

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or '').strip()
        word_count = len(text.split())
        needs_ocr = word_count < 15
        if needs_ocr:
            scanned_pages.append(i + 1)
        page_info.append({'page': i + 1, 'words': word_count, 'needs_ocr': needs_ocr})
        if text:
            page_texts.append(f'\n\n---\n\n*Page {i + 1}*\n\n{text}')
        elif needs_ocr:
            page_texts.append(f'\n\n---\n\n*Page {i + 1} [SCANNED — no selectable text]*\n')

    ocr_ratio = len(scanned_pages) / total_pages if total_pages > 0 else 0

    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'pdf',
        'pages': total_pages,
        'engine': 'pypdf',
        'scanned_pages': scanned_pages,
        'scanned_ratio': round(ocr_ratio, 2),
        'ocr_available': False,
        'ocr_hint': ('Some pages appear scanned (no selectable text).' if scanned_pages else None),
        'content': '\n'.join(page_texts),
        'page_stats': page_info,
        'metadata': {
            'title': meta.get('/Title', '') or '',
            'author': meta.get('/Author', '') or '',
            'subject': meta.get('/Subject', '') or '',
        },
        'processed_at': datetime.datetime.now().isoformat(),
    }


def extract_excel(file_path: Path) -> Dict[str, Any]:
    """Extract Excel content as markdown tables."""
    import openpyxl
    import datetime

    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    sheets_content = []
    sheet_names = wb.sheetnames

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue

        md = f'## {sheet_name}\n\n'
        header = [str(c) if c is not None else '' for c in rows[0]]
        md += '| ' + ' | '.join(header) + ' |\n'
        md += '| ' + ' | '.join(['---'] * len(header)) + ' |\n'
        for row in rows[1:]:
            cells = [str(c) if c is not None else '' for c in row]
            md += '| ' + ' | '.join(cells) + ' |\n'
        sheets_content.append(md)

    wb.close()
    content = '\n\n'.join(sheets_content)

    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'excel',
        'sheets': sheet_names,
        'sheet_count': len(sheet_names),
        'content': content,
        'metadata': {},
        'processed_at': datetime.datetime.now().isoformat()
    }


def extract_word(file_path: Path) -> Dict[str, Any]:
    """Extract Word (.docx) content as markdown using python-docx."""
    from docx import Document
    import datetime

    doc = Document(str(file_path))
    sections = []

    for block in doc.element.body:
        tag = block.tag.split('}')[-1]
        if tag == 'p':
            from docx.oxml.ns import qn
            style = block.find(qn('w:pPr'))
            text = ''.join(n.text or '' for n in block.iter() if n.tag.split('}')[-1] in ('t', 'delText'))
            if text.strip():
                sections.append(text)
        elif tag == 'tbl':
            from docx.table import Table as DocxTable
            tbl = DocxTable(block, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            if rows:
                md = '| ' + ' | '.join(rows[0]) + ' |\n'
                md += '| ' + ' | '.join(['---'] * len(rows[0])) + ' |\n'
                for row in rows[1:]:
                    md += '| ' + ' | '.join(row) + ' |\n'
                sections.append(md)

    content = '\n\n'.join(sections)
    core = doc.core_properties
    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'word',
        'paragraphs': len(doc.paragraphs),
        'content': content,
        'metadata': {'title': core.title or '', 'author': core.author or ''},
        'processed_at': datetime.datetime.now().isoformat()
    }


def extract_file(file_path: Path) -> Dict[str, Any]:
    """Route to appropriate extractor based on file type."""
    suffix = file_path.suffix.lower()

    if suffix == '.pdf':
        return extract_pdf(file_path)
    elif suffix in ('.xlsx', '.xls'):
        return extract_excel(file_path)
    elif suffix == '.docx':
        return extract_word(file_path)
    else:
        raise ValueError(f'Unsupported file type: {{suffix}}')


def run_batch_processing():
    """Main batch processing routine — run this in VS Code with F5 to debug."""
    # Discover all supported document types
    files_to_process: List[Path] = (
        list(SOURCE_FOLDER.glob('**/*.pdf')) +
        list(SOURCE_FOLDER.glob('**/*.xlsx')) +
        list(SOURCE_FOLDER.glob('**/*.xls')) +
        list(SOURCE_FOLDER.glob('**/*.docx'))
    )

    total = len(files_to_process)
    subfolders = {str(f.parent.relative_to(SOURCE_FOLDER)) for f in files_to_process if f.parent != SOURCE_FOLDER}
    subfolder_msg = f'   Subfolders: {len(subfolders)} ({", ".join(sorted(subfolders))[:120]})' if subfolders else ''
    print(f"\n📁 Batch Processing: {total} files")
    print(f"   Source: {SOURCE_FOLDER}")
    print(f"   Output: {OUTPUT_FOLDER}")
    if PROCESSING_PROMPT:
        print(f"   Prompt: {PROCESSING_PROMPT[:80]}{'...' if len(PROCESSING_PROMPT) > 80 else ''}")
    if subfolder_msg:
        print(subfolder_msg)
    print()

    if not files_to_process:
        print("❌ No files found to process!")
        return

    OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)

    done = 0
    errors = 0

    for file_path in files_to_process:
        try:
            result = extract_file(file_path)
            if PROCESSING_PROMPT:
                result['processing_prompt'] = PROCESSING_PROMPT
            rel = file_path.relative_to(SOURCE_FOLDER)
            out_file = OUTPUT_FOLDER / rel.with_suffix('.json')
            out_file.parent.mkdir(parents=True, exist_ok=True)

            with open(out_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)

            done += 1
            print(f"  ✓ [{{done:03d}}/{{total:03d}}] {{rel}}")
        except Exception as e:
            errors += 1
            rel = file_path.relative_to(SOURCE_FOLDER)
            print(f"  ✗ [{{done+1:03d}}/{{total:03d}}] {{rel}} — ERROR: {{e}}")

    print(f"\\n✅ Done! {{done}} files processed, {{errors}} errors.")
    print(f"   Output: {{OUTPUT_FOLDER}}\\n")

    # Summary
    summary = {{
        'total': total,
        'processed': done,
        'errors': errors,
        'source_folder': str(SOURCE_FOLDER),
        'output_folder': str(OUTPUT_FOLDER),
        'processing_prompt': PROCESSING_PROMPT or None,
        'files': [str(f.relative_to(SOURCE_FOLDER)) for f in files_to_process]
    }}
    summary_path = OUTPUT_FOLDER / 'batch_summary.json'
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2)
    print(f"📊 Summary saved: {{summary_path}}")


if __name__ == '__main__':
    if check_and_setup_environment():
        run_batch_processing()
'''
    script = (script
        .replace('__DATETIME_NOW__', _datetime_now)
        .replace('__SOURCE_PATH__', str(source_path))
        .replace('__OUTPUT_PATH__', str(output_path))
        .replace('__PROMPT_LINE__', prompt_line)
        .replace('__FILE_LIST_STR__', file_list_str)
    )
    return script


def _check_tesseract() -> bool:
    """Check if Tesseract OCR is installed and available on PATH."""
    import shutil
    return shutil.which('tesseract') is not None


def _extract_file(file_path: Path) -> dict:
    """Extract content from a PDF, Excel, or Word file and return structured dict."""
    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        return _extract_pdf(file_path)
    elif suffix in ('.xlsx', '.xls'):
        return _extract_excel(file_path)
    elif suffix == '.docx':
        return _extract_word(file_path)
    else:
        raise ValueError(f'Unsupported file type: {suffix}')


def _extract_pdf(file_path: Path) -> dict:
    """
    PDF extraction using pypdf (BSD 3-Clause — Apache 2.0 compatible).

    Per-page text is extracted and assembled into Markdown. Pages with very
    little text are flagged as likely scanned so callers can warn the user.
    """
    import datetime
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    meta = reader.metadata or {}

    page_texts = []
    page_info = []
    scanned_pages = []

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or '').strip()
        word_count = len(text.split())
        needs_ocr = word_count < 15
        if needs_ocr:
            scanned_pages.append(i + 1)
        page_info.append({'page': i + 1, 'words': word_count, 'needs_ocr': needs_ocr})
        if text:
            page_texts.append(f"\n\n---\n\n*Page {i + 1}*\n\n{text}")
        elif needs_ocr:
            page_texts.append(
                f"\n\n---\n\n*Page {i + 1} [SCANNED — no selectable text]*\n"
            )

    content = '\n'.join(page_texts)
    ocr_ratio = len(scanned_pages) / total_pages if total_pages > 0 else 0

    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'pdf',
        'pages': total_pages,
        'engine': 'pypdf',
        'scanned_pages': scanned_pages,
        'scanned_ratio': round(ocr_ratio, 2),
        'ocr_available': False,
        'ocr_hint': (
            'Some pages appear to be scanned images with no selectable text.'
            if scanned_pages else None
        ),
        'content': content,
        'page_stats': page_info,
        'metadata': {
            'title': meta.get('/Title', '') or '',
            'author': meta.get('/Author', '') or '',
            'subject': meta.get('/Subject', '') or '',
            'processed_at': datetime.datetime.now().isoformat(),
        },
        'processed_at': datetime.datetime.now().isoformat(),
    }


def _extract_excel(file_path: Path) -> dict:
    """Extract Excel to markdown tables using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            'openpyxl is required for Excel extraction. '
            'Run: pip install openpyxl'
        )

    import datetime
    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    sheets_content = []
    sheet_names = wb.sheetnames

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        # Filter out entirely empty rows
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue

        # Build markdown table
        md = f'## {sheet_name}\n\n'
        header = [str(c) if c is not None else '' for c in rows[0]]
        md += '| ' + ' | '.join(header) + ' |\n'
        md += '| ' + ' | '.join(['---'] * len(header)) + ' |\n'
        for row in rows[1:]:
            cells = [str(c) if c is not None else '' for c in row]
            md += '| ' + ' | '.join(cells) + ' |\n'
        sheets_content.append(md)

    wb.close()
    content = '\n\n'.join(sheets_content)

    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'excel',
        'sheets': sheet_names,
        'sheet_count': len(sheet_names),
        'content': content,
        'metadata': {},
        'processed_at': datetime.datetime.now().isoformat()
    }


def _extract_word(file_path: Path) -> dict:
    """Extract Word (.docx) content as markdown using python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
    except ImportError:
        raise ImportError(
            'python-docx is required for Word extraction. '
            'Run: pip install python-docx'
        )

    import datetime

    doc = Document(str(file_path))
    sections = []

    for block in doc.element.body:
        tag = block.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(
                n.text or '' for n in block.iter()
                if n.tag.split('}')[-1] in ('t', 'delText')
            )
            if text.strip():
                sections.append(text)
        elif tag == 'tbl':
            tbl = DocxTable(block, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            if rows:
                md = '| ' + ' | '.join(rows[0]) + ' |\n'
                md += '| ' + ' | '.join(['---'] * len(rows[0])) + ' |\n'
                for row in rows[1:]:
                    md += '| ' + ' | '.join(row) + ' |\n'
                sections.append(md)

    content = '\n\n'.join(sections)
    core = doc.core_properties
    return {
        'source': str(file_path),
        'filename': file_path.name,
        'type': 'word',
        'paragraphs': len(doc.paragraphs),
        'content': content,
        'metadata': {'title': core.title or '', 'author': core.author or ''},
        'processed_at': datetime.datetime.now().isoformat()
    }


# ── Web Scraper Integration ──────────────────────────────────────────────────

import sys
_scraper_path = str(PROJECT_ROOT / "servers" / "webscraper" / "mcp_project")
if _scraper_path not in sys.path:
    sys.path.insert(0, _scraper_path)

# Lazy-loaded scraper components (only import when first scrape route is hit)
_scrape_manager = None
_scrape_bs4 = None
_scrape_firecrawl = None

def _get_scrape_manager_only():
    """Load only JobManager so job list/CRUD works without scraper deps."""
    global _scrape_manager
    if _scrape_manager is None:
        from job_manager import JobManager
        scrape_storage = str(PROJECTS_DIR / ".scrape_jobs")
        _scrape_manager = JobManager(scrape_storage)
    return _scrape_manager


def _get_scrape_components():
    """Lazy-init scraper components to avoid import errors if deps missing."""
    global _scrape_manager, _scrape_bs4, _scrape_firecrawl
    if _scrape_manager is None or _scrape_bs4 is None or _scrape_firecrawl is None:
        try:
            from job_manager import JobManager
            from engines.bs4_engine import BS4Engine
            from engines.firecrawl_engine import FirecrawlEngine
            scrape_storage = str(PROJECTS_DIR / ".scrape_jobs")
            if _scrape_manager is None:
                _scrape_manager = JobManager(scrape_storage)
            if _scrape_bs4 is None:
                _scrape_bs4 = BS4Engine()
            if _scrape_firecrawl is None:
                _scrape_firecrawl = FirecrawlEngine()
            # Validate engines were created
            if _scrape_bs4 is None or _scrape_firecrawl is None:
                raise RuntimeError("Engine initialization returned None")
        except Exception as e:
            import traceback
            print(f"ERROR initializing scraper components: {e}")
            print(traceback.format_exc())
            raise
    return _scrape_manager, _scrape_bs4, _scrape_firecrawl


def _run_scrape_job_background(job_id: str, manager, engine):
    """Execute a scrape batch in a background thread with Socket.IO progress."""
    import asyncio as _aio
    from engines.base import ScrapeRequest

    loop = _aio.new_event_loop()
    _aio.set_event_loop(loop)

    job = manager.get_job(job_id)
    if not job:
        return

    if engine is None:
        error_msg = f"Scraper engine '{job.config.engine}' is unavailable"
        manager.mark_failed(job_id, error_msg)
        socketio.emit('scrape_progress', {
            'job_id': job_id,
            'status': 'error',
            'message': f'Job failed: {error_msg}'
        })
        return

    manager.mark_running(job_id)
    start_index = job.checkpoint.last_url_index + 1

    socketio.emit('scrape_progress', {
        'job_id': job_id,
        'status': 'started',
        'total': job.progress.total_urls,
        'completed': 0,
        'message': f'Starting scrape of {job.progress.total_urls} URL(s)...'
    })

    try:
        for i, url in enumerate(job.config.urls[start_index:], start=start_index):
            if manager.is_stopped(job_id):
                break
            # Pause support — block synchronously via loop.run_until_complete
            if manager.is_paused(job_id):
                socketio.emit('scrape_progress', {
                    'job_id': job_id, 'status': 'paused',
                    'total': job.progress.total_urls,
                    'completed': job.progress.completed,
                    'message': 'Job paused. Waiting to resume...'
                })
                loop.run_until_complete(manager.wait_if_paused(job_id))
                if manager.is_stopped(job_id):
                    break

            if url in job.checkpoint.completed_urls:
                job.progress.skipped += 1
                continue

            req = ScrapeRequest(url=url, schema=job.config.schema)
            result = loop.run_until_complete(engine.scrape(req))

            record = {
                'index': i,
                'url': result.url,
                'success': result.success,
                'title': result.title,
                'content_preview': (result.content or '')[:1000],
                'extracted_data': result.extracted_data,
                'status_code': result.status_code,
                'elapsed_ms': result.elapsed_ms,
                'error': result.error,
                'scraped_at': datetime.now().isoformat(),
            }
            manager.append_result(job_id, record)

            if result.success:
                job.progress.completed += 1
            else:
                job.progress.failed += 1
                job.progress.failed_urls.append({'url': url, 'error': result.error})

            manager.update_checkpoint(job_id, i, url)

            socketio.emit('scrape_progress', {
                'job_id': job_id,
                'status': 'progress',
                'total': job.progress.total_urls,
                'completed': job.progress.completed,
                'failed': job.progress.failed,
                'current_url': url,
                'current_result': record,
                'message': f'[{job.progress.completed}/{job.progress.total_urls}] {result.title or url}'
            })

        if not manager.is_stopped(job_id):
            manager.mark_complete(job_id)
            socketio.emit('scrape_progress', {
                'job_id': job_id,
                'status': 'complete',
                'total': job.progress.total_urls,
                'completed': job.progress.completed,
                'failed': job.progress.failed,
                'message': f'Done. {job.progress.completed} scraped, {job.progress.failed} failed.'
            })
        else:
            socketio.emit('scrape_progress', {
                'job_id': job_id,
                'status': 'stopped',
                'total': job.progress.total_urls,
                'completed': job.progress.completed,
                'message': 'Job stopped. Partial results preserved.'
            })

    except Exception as e:
        manager.mark_failed(job_id, str(e))
        socketio.emit('scrape_progress', {
            'job_id': job_id,
            'status': 'error',
            'message': f'Job failed: {e}'
        })
    finally:
        loop.close()


@app.route('/api/scrape/start', methods=['POST'])
def scrape_start():
    """Start a new scraping job."""
    data = request.json or {}
    urls = data.get('urls', [])
    engine_name = data.get('engine', 'beautifulsoup')
    schema = data.get('schema')
    job_name = data.get('name', '')
    project_path = data.get('project_path', '')

    if not urls:
        return jsonify({'error': 'urls list is required'}), 400

    # Ensure project scaffold + venv lifecycle before job starts.
    try:
        env_info = _ensure_scrape_project_environment(job_name, project_path)
        project_path = env_info['project_path']
    except subprocess.CalledProcessError as e:
        return jsonify({
            'error': 'Failed to prepare project virtual environment',
            'details': e.stderr.decode() if hasattr(e, 'stderr') and e.stderr else str(e)
        }), 500
    except Exception as e:
        return jsonify({'error': f'Failed to prepare project environment: {str(e)}'}), 500

    try:
        manager, bs4_eng, fc_eng = _get_scrape_components()
    except ModuleNotFoundError as e:
        missing_mod = str(e).replace("No module named ", "").replace("'", "").strip()
        install_hint = "pip install beautifulsoup4 lxml httpx firecrawl-py"
        return jsonify({
            'error': (
                f"Missing dependency in client runtime environment: {missing_mod}. "
                f"Install required scraper dependencies and restart the web app."
            ),
            'missing_module': missing_mod,
            'hint': install_hint,
            'project_name': env_info['project_name'],
            'project_path': env_info['project_path'],
            'venv_path': env_info['venv_path'],
            'requirements_path': env_info['requirements_path'],
        }), 500
    except Exception as e:
        import traceback
        error_msg = f"Failed to initialize scraper: {str(e)}"
        print(f"ERROR in scrape_start: {error_msg}")
        print(traceback.format_exc())
        return jsonify({
            'error': error_msg,
            'traceback': traceback.format_exc(),
            'project_name': env_info['project_name'],
            'project_path': env_info['project_path'],
            'venv_path': env_info['venv_path'],
            'requirements_path': env_info['requirements_path'],
        }), 500

    # Set Firecrawl API key if provided
    fc_key = data.get('firecrawl_api_key', '')
    if fc_key:
        fc_eng.set_api_key(fc_key)

    job = manager.create_job(
        urls=urls,
        name=job_name,
        engine=engine_name,
        schema=schema,
        project_path=project_path,
    )

    engine = fc_eng if engine_name == 'firecrawl' else bs4_eng
    if engine is None:
        manager.mark_failed(job.job_id, f"Scraper engine '{engine_name}' is unavailable")
        return jsonify({
            'error': f"Scraper engine '{engine_name}' is unavailable",
            'job_id': job.job_id,
        }), 500

    threading.Thread(
        target=_run_scrape_job_background,
        args=(job.job_id, manager, engine),
        daemon=True,
    ).start()

    return jsonify({
        'status': 'started',
        'job_id': job.job_id,
        'name': job.name,
        'total_urls': len(urls),
        'engine': engine_name,
        'project_id': env_info['project_id'],
        'project_name': env_info['project_name'],
        'project_path': env_info['project_path'],
        'venv_path': env_info['venv_path'],
        'requirements_path': env_info['requirements_path'],
    })


@app.route('/api/scrape/search', methods=['POST'])
def scrape_search():
    """Search the web for URLs matching a query."""
    data = request.json or {}
    query = data.get('query', '').strip()
    engine_name = data.get('engine', 'beautifulsoup')
    max_results = data.get('max_results', 10)

    if not query:
        return jsonify({'error': 'query is required'}), 400

    manager, bs4_eng, fc_eng = _get_scrape_components()
    engine = fc_eng if engine_name == 'firecrawl' else bs4_eng

    import asyncio as _aio
    loop = _aio.new_event_loop()
    try:
        results = loop.run_until_complete(engine.search(query, max_results=max_results))
    except Exception as e:
        results = [{'url': '', 'title': 'Error', 'snippet': str(e)}]
    finally:
        loop.close()

    return jsonify({'query': query, 'results': results, 'engine': engine_name})


@app.route('/api/scrape/jobs', methods=['GET'])
def scrape_list_jobs():
    """List all scraping jobs."""
    manager = _get_scrape_manager_only()
    project = request.args.get('project_path', None)
    return jsonify({'jobs': manager.list_jobs(project_path=project)})


@app.route('/api/scrape/jobs/<job_id>', methods=['GET'])
def scrape_get_job(job_id):
    """Get job details."""
    manager = _get_scrape_manager_only()
    job = manager.get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    return jsonify(job.to_dict())


@app.route('/api/scrape/jobs/<job_id>/pause', methods=['POST'])
def scrape_pause_job(job_id):
    """Pause a running job."""
    manager = _get_scrape_manager_only()
    if manager.pause_job(job_id):
        return jsonify({'status': 'paused', 'job_id': job_id})
    return jsonify({'error': 'Cannot pause (not running or not found)'}), 400


@app.route('/api/scrape/jobs/<job_id>/resume', methods=['POST'])
def scrape_resume_job(job_id):
    """Resume a paused or failed job."""
    manager = _get_scrape_manager_only()
    try:
        _, bs4_eng, fc_eng = _get_scrape_components()
    except Exception as e:
        return jsonify({'error': f'Cannot resume: scraper runtime unavailable ({str(e)})'}), 500
    
    # Validate engines are not None
    if bs4_eng is None or fc_eng is None:
        return jsonify({'error': 'Scraper engines failed to initialize'}), 500
    
    job = manager.get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    if not manager.resume_job(job_id):
        return jsonify({'error': 'Cannot resume (wrong status)'}), 400

    engine = fc_eng if job.config.engine == 'firecrawl' else bs4_eng
    
    # Double-check engine is valid
    if engine is None:
        return jsonify({'error': f'Engine {job.config.engine} is not available'}), 500
    
    threading.Thread(
        target=_run_scrape_job_background,
        args=(job.job_id, manager, engine),
        daemon=True,
    ).start()

    return jsonify({'status': 'resumed', 'job_id': job_id})


@app.route('/api/scrape/jobs/<job_id>/stop', methods=['POST'])
def scrape_stop_job(job_id):
    """Stop a running or paused job."""
    manager = _get_scrape_manager_only()
    if manager.stop_job(job_id):
        return jsonify({'status': 'stopped', 'job_id': job_id})
    return jsonify({'error': 'Cannot stop'}), 400


@app.route('/api/scrape/jobs/<job_id>', methods=['DELETE'])
def scrape_delete_job(job_id):
    """Delete a job and all its data."""
    manager = _get_scrape_manager_only()
    if manager.delete_job(job_id):
        return jsonify({'status': 'deleted'})
    return jsonify({'error': 'Job not found'}), 404


@app.route('/api/scrape/jobs/<job_id>/results', methods=['GET'])
def scrape_get_results(job_id):
    """Get paginated results for a job."""
    manager = _get_scrape_manager_only()
    page = request.args.get('page', 1, type=int)
    page_size = request.args.get('page_size', 50, type=int)
    search = request.args.get('search', None)
    data = manager.get_results(job_id, page=page, page_size=page_size, search=search)
    return jsonify(data)


@app.route('/api/scrape/jobs/<job_id>/results/<int:record_index>', methods=['PUT'])
def scrape_update_record(job_id, record_index):
    """Update a specific result record."""
    manager = _get_scrape_manager_only()
    updates = request.json or {}
    if manager.update_record(job_id, record_index, updates):
        return jsonify({'status': 'updated'})
    return jsonify({'error': 'Record not found'}), 404


@app.route('/api/scrape/jobs/<job_id>/results/<int:record_index>', methods=['DELETE'])
def scrape_delete_record(job_id, record_index):
    """Delete a specific result record."""
    manager = _get_scrape_manager_only()
    if manager.delete_record(job_id, record_index):
        return jsonify({'status': 'deleted'})
    return jsonify({'error': 'Record not found'}), 404


@app.route('/api/scrape/jobs/<job_id>/export', methods=['GET'])
def scrape_export_results(job_id):
    """Export job results to a file and return it for download."""
    from flask import send_file
    manager = _get_scrape_manager_only()
    job = manager.get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404

    fmt = request.args.get('format', 'json')
    all_data = manager.get_results(job_id, page=1, page_size=999999)
    results = all_data.get('results', [])
    if not results:
        return jsonify({'error': 'No results to export'}), 404

    # Determine export directory
    if job.project_path:
        export_dir = str(Path(job.project_path) / 'exports')
    else:
        export_dir = str(PROJECTS_DIR / '.scrape_jobs' / 'exports')

    from exporters import export_results as do_export, EXPORT_FORMATS
    if fmt not in EXPORT_FORMATS:
        return jsonify({'error': f'Unknown format. Available: {list(EXPORT_FORMATS.keys())}'}), 400

    output_path = do_export(results, format=fmt, output_dir=export_dir)
    return send_file(output_path, as_attachment=True)


@app.route('/api/scrape/configure', methods=['POST'])
def scrape_configure():
    """Configure scraping engine (e.g., Firecrawl API key)."""
    data = request.json or {}
    engine_name = data.get('engine', 'firecrawl')
    api_key = data.get('api_key', '').strip()

    _, _, fc_eng = _get_scrape_components()

    if engine_name == 'firecrawl' and api_key:
        fc_eng.set_api_key(api_key)
        # Persist to config
        config = load_config()
        config.setdefault('scraper', {})['firecrawl_api_key'] = api_key
        save_config(config)
        return jsonify({'status': 'configured', 'engine': 'firecrawl'})

    return jsonify({'error': 'Provide engine and api_key'}), 400


@app.route('/api/scrape/engine/status', methods=['GET'])
def scrape_engine_status():
    """Check engine availability."""
    import asyncio as _aio
    _, bs4_eng, fc_eng = _get_scrape_components()
    loop = _aio.new_event_loop()
    try:
        bs4_ok, bs4_msg = loop.run_until_complete(bs4_eng.check_available())
        fc_ok, fc_msg = loop.run_until_complete(fc_eng.check_available())
    finally:
        loop.close()

    return jsonify({
        'engines': {
            'beautifulsoup': {'available': bs4_ok, 'message': bs4_msg, 'premium': False},
            'firecrawl': {'available': fc_ok, 'message': fc_msg, 'premium': True},
        }
    })


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    print("🚀 Starting MCP Client UI...")
    print(f"📁 Project root: {PROJECT_ROOT}")
    print(f"🔧 Config file: {config_file}")
    print(f"\n🌐 Open http://localhost:{port} in your browser\n")
    debug = os.environ.get('FLASK_DEBUG', '1') == '1'
    socketio.run(app, debug=debug, host='0.0.0.0', port=port, allow_unsafe_werkzeug=True)
