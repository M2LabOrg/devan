"""
Document MCP Server — Unified Document Processing for LLM & RAG Pipelines

Combines Excel, PDF, and Word document extraction into a single MCP server.

Capabilities:
- Excel: Intelligent table detection, multi-sheet extraction, OpenSearch indexing
- PDF: Deep layout analysis, OCR, table/figure extraction, RAG chunking
- Word: Paragraph/table extraction, metadata, Markdown conversion

Extraction engines:
- Docling (IBM): Deep-learning layout analysis for complex layouts
- PyMuPDF4LLM: Ultra-fast Markdown conversion
- python-docx: Native Word document parsing
"""

import json
import os
import csv
import io
import re
import hashlib
import time
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass, asdict
from enum import Enum
from mcp.server.fastmcp import FastMCP, Context

# Smart extractor for intelligent Excel table detection
from smart_extractor import (
    SmartExcelExtractor,
    analysis_to_dict,
    table_to_csv_string,
    table_to_markdown,
    table_to_records,
)

# ─── Optional imports ─────────────────────────────────────────────────────────

try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from docling.chunking import HierarchicalChunker
    _chunker: Optional["HierarchicalChunker"] = None

    def get_chunker() -> "HierarchicalChunker":
        global _chunker
        if _chunker is None:
            _chunker = HierarchicalChunker()
        return _chunker

    CHUNKER_AVAILABLE = True
except ImportError:
    CHUNKER_AVAILABLE = False

try:
    from opensearchpy import OpenSearch
    OPENSEARCH_AVAILABLE = True
except ImportError:
    OPENSEARCH_AVAILABLE = False

try:
    import pyarrow as pa
    import pyarrow.parquet as pq
    import pyarrow.compute as pc
    PYARROW_AVAILABLE = True
except ImportError:
    PYARROW_AVAILABLE = False

# ─── Configuration ────────────────────────────────────────────────────────────

EXCEL_DIR = os.environ.get("EXCEL_DIR", "excel_files")
PDF_DIR = os.environ.get("PDF_DIR", "../pdf_files")
WORD_DIR = os.environ.get("WORD_DIR", "word_files")
OUTPUT_DIR = os.environ.get("DOC_OUTPUT_DIR", "../output")
CSV_DIR = os.environ.get("CSV_DIR", "csv_files")
TEXT_DIR = os.environ.get("TEXT_DIR", "text_files")
MAX_PAGES_PER_CHUNK = int(os.environ.get("PDF_MAX_PAGES_PER_CHUNK", "20"))

# ─── Initialize MCP Server ───────────────────────────────────────────────────

mcp = FastMCP("document-mcp")


# ─── Security: Path Validation ───────────────────────────────────────────────
# All configured directories that MCP tools are permitted to access.
# Prevents path traversal attacks and enforces MCP Roots-style access boundaries.

_ALLOWED_DIRS: list[str] = [EXCEL_DIR, PDF_DIR, WORD_DIR, OUTPUT_DIR, CSV_DIR, TEXT_DIR]

# Docker volume mount roots — any path under these is safe because the host
# filesystem is already mounted read-only by docker-compose.
_DOCKER_MOUNT_ROOTS: list[str] = ["/host/home", "/host/volumes"]


def _validate_path(path: str) -> str:
    """Resolve and verify a file path is within the server's allowed directories.

    Prevents path traversal (../../etc/passwd) and symlink escapes.
    Returns the resolved absolute path, or raises ValueError with a clear message.
    """
    resolved = os.path.realpath(os.path.abspath(path))

    # Allow any path under the Docker read-only volume mounts.
    for mount in _DOCKER_MOUNT_ROOTS:
        if resolved.startswith(mount + os.sep) or resolved == mount:
            return resolved

    for allowed in _ALLOWED_DIRS:
        allowed_resolved = os.path.realpath(os.path.abspath(allowed))
        if resolved.startswith(allowed_resolved + os.sep) or resolved == allowed_resolved:
            return resolved
    raise ValueError(
        f"Access denied: path is outside allowed directories.\n"
        f"  Requested: {resolved}\n"
        f"  Allowed:   {_ALLOWED_DIRS}\n"
        f"Tip: set EXCEL_DIR / PDF_DIR / WORD_DIR / DOC_OUTPUT_DIR / CSV_DIR / TEXT_DIR env vars to include your files."
    )


# ─── Shared Singletons ───────────────────────────────────────────────────────

_docling_converter: Optional["DocumentConverter"] = None
_opensearch_client: Optional["OpenSearch"] = None


def _get_docling_converter() -> "DocumentConverter":
    global _docling_converter
    if _docling_converter is None:
        _docling_converter = DocumentConverter()
    return _docling_converter


def _get_opensearch_client() -> Optional["OpenSearch"]:
    global _opensearch_client
    if _opensearch_client is None and OPENSEARCH_AVAILABLE:
        host = os.environ.get("OPENSEARCH_HOST", "localhost")
        port = int(os.environ.get("OPENSEARCH_PORT", "9200"))
        security_enabled = os.environ.get("OPENSEARCH_SECURITY_ENABLED", "false").lower() == "true"
        try:
            if security_enabled:
                auth = (
                    os.environ.get("OPENSEARCH_USER", "admin"),
                    os.environ.get("OPENSEARCH_PASSWORD", "admin"),
                )
                _opensearch_client = OpenSearch(
                    hosts=[{"host": host, "port": port}],
                    http_auth=auth,
                    use_ssl=True,
                    verify_certs=False,
                    ssl_assert_hostname=False,
                    ssl_show_warn=False,
                )
            else:
                _opensearch_client = OpenSearch(
                    hosts=[{"host": host, "port": port}],
                    http_compress=True,
                    use_ssl=False,
                    verify_certs=False,
                    ssl_assert_hostname=False,
                    ssl_show_warn=False,
                )
        except Exception:
            pass
    return _opensearch_client


# ─── PDF Data Models ──────────────────────────────────────────────────────────

@dataclass
class PageContent:
    page_number: int
    text: str
    tables: List[Dict[str, Any]]
    figures: List[Dict[str, Any]]
    word_count: int


@dataclass
class ExtractionResult:
    file_path: str
    file_name: str
    engine: str
    total_pages: int
    total_words: int
    total_tables: int
    total_figures: int
    extraction_time_seconds: float
    content_markdown: str
    pages: List[Dict[str, Any]]
    metadata: Dict[str, Any]


# ─── PDF Helpers ─────────────────────────────────────────────────────────────

def _parse_page_range_1indexed(page_range: str) -> set:
    """Parse '1-5,8,10-12' into a set of 1-indexed page numbers."""
    pages: set = set()
    for part in page_range.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            pages.update(range(int(a), int(b) + 1))
        else:
            pages.add(int(part))
    return pages


def _extract_pdf_with_docling(file_path: str) -> ExtractionResult:
    converter = _get_docling_converter()
    start = time.time()
    result = converter.convert(file_path)
    doc = result.document
    markdown_content = doc.export_to_markdown()
    text_content = doc.export_to_text()
    page_splits = re.split(r"\n---\n|\f", markdown_content)
    page_splits = [p.strip() for p in page_splits if p.strip()]
    page_contents = []
    for i, page_text in enumerate(page_splits):
        page_contents.append({
            "page_number": i + 1,
            "content": page_text,
            "character_count": len(page_text),
            "word_count": len(page_text.split()),
        })
    if not page_contents and markdown_content.strip():
        page_contents.append({
            "page_number": 1,
            "content": markdown_content.strip(),
            "character_count": len(markdown_content.strip()),
            "word_count": len(markdown_content.split()),
        })
    tables = []
    if hasattr(doc, "tables"):
        for i, table in enumerate(doc.tables):
            table_data = {"index": i, "markdown": table.export_to_markdown() if hasattr(table, "export_to_markdown") else str(table)}
            if hasattr(table, "num_rows"):
                table_data["rows"] = table.num_rows
            if hasattr(table, "num_cols"):
                table_data["cols"] = table.num_cols
            tables.append(table_data)
    figures = []
    if hasattr(doc, "pictures"):
        for i, pic in enumerate(doc.pictures):
            figures.append({"index": i, "caption": str(pic.caption) if hasattr(pic, "caption") and pic.caption else None})
    elapsed = time.time() - start
    page_count = len(result.pages) if hasattr(result, "pages") else 0
    return ExtractionResult(
        file_path=file_path,
        file_name=os.path.basename(file_path),
        engine="docling",
        total_pages=page_count if page_count else len(page_contents),
        total_words=len(text_content.split()),
        total_tables=len(tables),
        total_figures=len(figures),
        extraction_time_seconds=round(elapsed, 2),
        content_markdown=markdown_content,
        pages=page_contents,
        metadata={"tables": tables, "figures": figures},
    )




def _chunk_by_pages(markdown: str, pages_per_chunk: int = 5) -> List[Dict[str, Any]]:
    pages = re.split(r"\n---\n|\f", markdown)
    pages = [p.strip() for p in pages if p.strip()]
    chunks = []
    for i in range(0, len(pages), pages_per_chunk):
        chunk_pages = pages[i:i + pages_per_chunk]
        chunk_text = "\n\n".join(chunk_pages)
        chunks.append({
            "chunk_index": len(chunks),
            "page_start": i + 1,
            "page_end": min(i + pages_per_chunk, len(pages)),
            "text": chunk_text,
            "word_count": len(chunk_text.split()),
            "char_count": len(chunk_text),
        })
    return chunks


def _chunk_by_tokens(text: str, chunk_size: int = 512, overlap: int = 50) -> List[Dict[str, Any]]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunk_text = " ".join(chunk_words)
        chunks.append({
            "chunk_index": len(chunks),
            "word_start": start,
            "word_end": end,
            "text": chunk_text,
            "word_count": len(chunk_words),
            "char_count": len(chunk_text),
        })
        start += chunk_size - overlap
    return chunks


def _chunk_by_sections(markdown: str) -> List[Dict[str, Any]]:
    sections = re.split(r"\n(?=#{1,3}\s)", markdown)
    sections = [s.strip() for s in sections if s.strip()]
    chunks = []
    for section in sections:
        lines = section.split("\n")
        heading = lines[0].strip().lstrip("#").strip() if lines else "Untitled"
        chunks.append({
            "chunk_index": len(chunks),
            "heading": heading,
            "text": section,
            "word_count": len(section.split()),
            "char_count": len(section),
        })
    return chunks


def _parse_page_range(page_range: str) -> List[int]:
    pages = set()
    for part in page_range.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            for p in range(int(start), int(end) + 1):
                pages.add(p - 1)
        else:
            pages.add(int(part) - 1)
    return sorted(pages)


def _table_to_markdown(headers: List, rows: List[List]) -> str:
    if not headers:
        return ""
    clean_headers = [str(h) if h else "" for h in headers]
    md = "| " + " | ".join(clean_headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(clean_headers)) + " |\n"
    for row in rows:
        clean_row = [str(c) if c else "" for c in row]
        while len(clean_row) < len(clean_headers):
            clean_row.append("")
        md += "| " + " | ".join(clean_row[:len(clean_headers)]) + " |\n"
    return md


def _file_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()[:12]


def _pick_pdf_engine(file_path: str) -> str:
    return "docling" if DOCLING_AVAILABLE else "none"


# ─── Word Helpers ─────────────────────────────────────────────────────────────

def _word_table_to_markdown(table) -> str:
    """Convert a python-docx Table to a Markdown table string."""
    rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    headers = rows[0]
    md = "| " + " | ".join(headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for row in rows[1:]:
        while len(row) < len(headers):
            row.append("")
        md += "| " + " | ".join(row[:len(headers)]) + " |\n"
    return md


def _heading_level(paragraph) -> Optional[int]:
    """Return heading level (1-9) for a paragraph, or None if not a heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split(" ")[-1])
        except ValueError:
            pass
    return None


# ─── Excel MCP Tools ─────────────────────────────────────────────────────────

@mcp.tool()
def list_excel_files(directory: Optional[str] = None) -> str:
    """
    List all Excel files (.xlsx, .xls, .xlsm) in the specified directory.

    Args:
        directory: Path to directory. Defaults to EXCEL_DIR env var or 'excel_files'.

    Returns:
        JSON with list of Excel files and their metadata.
    """
    search_dir = directory or EXCEL_DIR
    if not os.path.exists(search_dir):
        return json.dumps({"error": f"Directory not found: {search_dir}", "hint": "Create the directory or set EXCEL_DIR environment variable"}, indent=2)
    excel_extensions = [".xlsx", ".xls", ".xlsm", ".xlsb"]
    files = []
    for item in os.listdir(search_dir):
        if any(item.lower().endswith(ext) for ext in excel_extensions):
            file_path = os.path.join(search_dir, item)
            try:
                stat = os.stat(file_path)
                files.append({"filename": item, "path": file_path, "size_bytes": stat.st_size, "modified": str(stat.st_mtime)})
            except OSError as e:
                files.append({"filename": item, "path": file_path, "error": str(e)})
    return json.dumps({"directory": search_dir, "file_count": len(files), "files": files}, indent=2)


@mcp.tool()
def extract_excel_content(file_path: str, sheet_name: Optional[str] = None) -> str:
    """
    Extract content from an Excel file using docling.

    Args:
        file_path: Path to the Excel file.
        sheet_name: Optional specific sheet. If omitted, extracts all sheets.

    Returns:
        JSON with extracted content structured as tables and text.
    """
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "docling is not installed. Install with: uv add docling"}, indent=2)
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        converter = _get_docling_converter()
        result = converter.convert(file_path)
        extracted_data = {"file_path": file_path, "file_name": os.path.basename(file_path), "document_type": "excel", "pages": []}
        for page in result.pages:
            page_data = {"page_number": page.page_no, "tables": [], "text_content": ""}
            for table in page.tables:
                table_data = {"headers": table.header if hasattr(table, "header") else [], "rows": []}
                for row in table.data:
                    table_data["rows"].append([str(cell) if cell else "" for cell in row])
                page_data["tables"].append(table_data)
            for item in page.items:
                if hasattr(item, "text"):
                    page_data["text_content"] += item.text + "\n"
            extracted_data["pages"].append(page_data)
        extracted_data["full_text"] = result.document.export_to_text()
        extracted_data["markdown"] = result.document.export_to_markdown()
        return json.dumps(extracted_data, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract content: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def query_excel_data(file_path: str, query: str) -> str:
    """
    Query specific information from an Excel file.

    Args:
        file_path: Path to the Excel file.
        query: Description of what data to look for (e.g., 'sales figures', 'Q1 revenue').

    Returns:
        JSON with relevant extracted data based on the query context.
    """
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "docling is not installed. Install with: uv add docling"}, indent=2)
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        converter = _get_docling_converter()
        result = converter.convert(file_path)
        structured_data = {"query": query, "file": os.path.basename(file_path), "sheets": [], "tables": [], "key_values": []}
        for page in result.pages:
            sheet_info = {"sheet_number": page.page_no, "tables_found": len(page.tables), "text_segments": []}
            for item in page.items:
                if hasattr(item, "text"):
                    sheet_info["text_segments"].append(item.text)
            structured_data["sheets"].append(sheet_info)
            for table in page.tables:
                table_rows = [[str(cell) if cell else "" for cell in row] for row in table.data]
                structured_data["tables"].append({"sheet": page.page_no, "row_count": len(table_rows), "data": table_rows[:50]})
        structured_data["full_content_preview"] = result.document.export_to_markdown()[:5000]
        return json.dumps(structured_data, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to query data: {str(e)}", "file_path": file_path, "query": query}, indent=2)


@mcp.tool()
def get_excel_summary(file_path: str) -> str:
    """
    Get a summary of an Excel file's structure and content.

    Args:
        file_path: Path to the Excel file.

    Returns:
        JSON with summary information about sheets, tables, and data.
    """
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "docling is not installed. Install with: uv add docling"}, indent=2)
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        converter = _get_docling_converter()
        result = converter.convert(file_path)
        summary = {"file_name": os.path.basename(file_path), "file_path": file_path, "total_sheets": len(result.pages), "sheets": []}
        for page in result.pages:
            sheet_summary = {"sheet_number": page.page_no, "table_count": len(page.tables), "has_text": len(page.items) > 0, "tables": []}
            for i, table in enumerate(page.tables):
                row_count = len(table.data) if hasattr(table, "data") else 0
                sheet_summary["tables"].append({"table_index": i, "row_count": row_count})
            summary["sheets"].append(sheet_summary)
        return json.dumps(summary, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to generate summary: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def convert_excel_to_csv(file_path: str, output_dir: Optional[str] = None) -> str:
    """
    Convert Excel file sheets to CSV format.

    Args:
        file_path: Path to the Excel file.
        output_dir: Optional directory to save CSV files. Defaults to same directory as the Excel file.

    Returns:
        JSON with paths to created CSV files.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        import pandas as pd
        if output_dir is None:
            output_dir = os.path.dirname(file_path) or "."
        os.makedirs(output_dir, exist_ok=True)
        excel_file = pd.ExcelFile(file_path)
        created_files = []
        base_name = Path(file_path).stem
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            safe_sheet = "".join(c if c.isalnum() else "_" for c in sheet_name)
            csv_path = os.path.join(output_dir, f"{base_name}_{safe_sheet}.csv")
            df.to_csv(csv_path, index=False)
            created_files.append({"sheet_name": sheet_name, "csv_path": csv_path, "rows": len(df), "columns": len(df.columns)})
        return json.dumps({"source_file": file_path, "output_directory": output_dir, "csv_files_created": len(created_files), "files": created_files}, indent=2)
    except ImportError:
        return json.dumps({"error": "pandas not installed. Install with: uv add pandas"}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to convert to CSV: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def convert_excel_to_json(file_path: str, output_dir: Optional[str] = None, include_metadata: bool = True) -> str:
    """
    Convert Excel file sheets to JSON format.

    Args:
        file_path: Path to the Excel file.
        output_dir: Optional directory to save JSON files.
        include_metadata: Whether to include sheet metadata (default: True).

    Returns:
        JSON with paths to created JSON files and inline data preview.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        import pandas as pd
        if output_dir is None:
            output_dir = os.path.dirname(file_path) or "."
        os.makedirs(output_dir, exist_ok=True)
        excel_file = pd.ExcelFile(file_path)
        created_files = []
        all_data = {}
        base_name = Path(file_path).stem
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            records = df.replace({pd.NaT: None, pd.NA: None}).to_dict(orient="records")
            clean_records = []
            for record in records:
                clean_record = {}
                for key, value in record.items():
                    if pd.isna(value):
                        clean_record[key] = None
                    elif isinstance(value, (pd.Timestamp, datetime)):
                        clean_record[key] = value.isoformat()
                    else:
                        clean_record[key] = value
                clean_records.append(clean_record)
            sheet_data = {"sheet_name": sheet_name, "row_count": len(df), "column_count": len(df.columns), "columns": list(df.columns), "data": clean_records}
            if include_metadata:
                sheet_data["metadata"] = {"source_file": os.path.basename(file_path), "sheet_index": excel_file.sheet_names.index(sheet_name), "dtypes": {col: str(df[col].dtype) for col in df.columns}}
            all_data[sheet_name] = sheet_data
            safe_sheet = "".join(c if c.isalnum() else "_" for c in sheet_name)
            json_path = os.path.join(output_dir, f"{base_name}_{safe_sheet}.json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(sheet_data, f, indent=2, ensure_ascii=False, default=str)
            created_files.append({"sheet_name": sheet_name, "json_path": json_path, "rows": len(df), "columns": len(df.columns)})
        combined_path = os.path.join(output_dir, f"{base_name}_all_sheets.json")
        combined_data = {"source_file": os.path.basename(file_path), "file_path": file_path, "total_sheets": len(excel_file.sheet_names), "sheet_names": excel_file.sheet_names, "sheets": all_data}
        with open(combined_path, "w", encoding="utf-8") as f:
            json.dump(combined_data, f, indent=2, ensure_ascii=False, default=str)
        return json.dumps({"source_file": file_path, "output_directory": output_dir, "json_files_created": len(created_files), "files": created_files, "combined_file": combined_path, "preview": {sheet: {"columns": data["columns"], "row_count": data["row_count"], "sample_rows": data["data"][:3]} for sheet, data in all_data.items()}}, indent=2, default=str)
    except ImportError:
        return json.dumps({"error": "pandas not installed. Install with: uv add pandas"}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to convert to JSON: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def extract_excel_rows(
    file_path: str,
    sheet_name: Optional[str] = None,
    start_row: int = 0,
    end_row: Optional[int] = None,
    max_rows: int = 1000,
) -> str:
    """
    Extract specific rows from an Excel file.

    Args:
        file_path: Path to the Excel file.
        sheet_name: Optional specific sheet (default: first sheet).
        start_row: Row number to start from (0-indexed, default: 0).
        end_row: Row number to end at (exclusive).
        max_rows: Maximum number of rows to return (default: 1000).

    Returns:
        JSON with selected rows and metadata.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        import pandas as pd
        excel_file = pd.ExcelFile(file_path)
        if sheet_name is None:
            sheet_name = excel_file.sheet_names[0]
        if end_row is None:
            end_row = start_row + max_rows
        skip_rows = list(range(1, start_row + 1)) if start_row > 0 else None
        n_rows = end_row - start_row
        df = pd.read_excel(file_path, sheet_name=sheet_name, skiprows=skip_rows, nrows=n_rows)
        records = df.replace({pd.NaT: None, pd.NA: None}).to_dict(orient="records")
        clean_records = []
        for record in records:
            clean_record = {}
            for key, value in record.items():
                if pd.isna(value):
                    clean_record[key] = None
                elif isinstance(value, (pd.Timestamp, datetime)):
                    clean_record[key] = value.isoformat()
                else:
                    clean_record[key] = value
            clean_records.append(clean_record)
        return json.dumps({"file_path": file_path, "sheet_name": sheet_name, "start_row": start_row, "end_row": end_row, "rows_returned": len(clean_records), "columns": list(df.columns), "data": clean_records}, indent=2, default=str)
    except ImportError:
        return json.dumps({"error": "pandas not installed. Install with: uv add pandas"}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract rows: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def chunk_excel_content(file_path: str, chunk_size: int = 1000) -> str:
    """
    Chunk Excel content using docling's HierarchicalChunker for LLM context windows or embeddings.

    Args:
        file_path: Path to the Excel file.
        chunk_size: Approximate number of tokens per chunk (default: 1000).

    Returns:
        JSON with chunks and metadata.
    """
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "docling is not installed. Install with: uv add docling"}, indent=2)
    if not CHUNKER_AVAILABLE:
        return json.dumps({"error": "docling chunking not available. Install with: uv add 'docling[chunking]'"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        converter = _get_docling_converter()
        chunker = get_chunker()
        result = converter.convert(file_path)
        chunks = list(chunker.chunk(result.document))
        chunk_data = []
        for i, chunk in enumerate(chunks):
            chunk_info = {
                "chunk_number": i + 1,
                "text": chunk.text,
                "metadata": {
                    "headings": chunk.meta.headings if hasattr(chunk.meta, "headings") else [],
                    "page_number": chunk.meta.page_number if hasattr(chunk.meta, "page_number") else None,
                    "doc_items": [str(ref) for ref in chunk.meta.doc_items] if hasattr(chunk.meta, "doc_items") else [],
                },
            }
            chunk_data.append(chunk_info)
        return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "total_chunks": len(chunk_data), "chunk_size_setting": chunk_size, "chunks": chunk_data}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to chunk content: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def index_excel_in_opensearch(file_path: str, index_name: Optional[str] = None, chunk_size: int = 1000) -> str:
    """
    Index an Excel file into OpenSearch for fast searching of large datasets.

    Args:
        file_path: Path to the Excel file.
        index_name: Optional custom OpenSearch index name (defaults to filename).
        chunk_size: Number of rows to process per batch (default: 1000).

    Returns:
        JSON with indexing status and statistics.
    """
    if not OPENSEARCH_AVAILABLE:
        return json.dumps({"error": "OpenSearch not installed. Install with: uv add opensearch-py"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        import pandas as pd
        client = _get_opensearch_client()
        if client is None:
            return json.dumps({"error": "Could not connect to OpenSearch"}, indent=2)
        if index_name is None:
            index_name = Path(file_path).stem.lower().replace(" ", "_")
        if not client.indices.exists(index=index_name):
            client.indices.create(index=index_name, body={"settings": {"number_of_shards": 1, "number_of_replicas": 0}, "mappings": {"properties": {"sheet_name": {"type": "keyword"}, "row_number": {"type": "integer"}, "data": {"type": "object"}, "full_text": {"type": "text"}, "source_file": {"type": "keyword"}}}})
        excel_file = pd.ExcelFile(file_path)
        total_docs = 0
        sheet_stats = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            records = df.replace({pd.NaT: None, pd.NA: None}).to_dict(orient="records")
            bulk_data = []
            for i, record in enumerate(records):
                clean_record = {}
                full_text_parts = []
                for key, value in record.items():
                    if pd.isna(value):
                        clean_record[key] = None
                    elif isinstance(value, (pd.Timestamp, datetime)):
                        clean_record[key] = value.isoformat()
                        full_text_parts.append(str(value))
                    else:
                        clean_record[key] = value
                        full_text_parts.append(str(value))
                bulk_data.append({"_index": index_name, "_source": {"sheet_name": sheet_name, "row_number": i + 1, "data": clean_record, "full_text": " ".join(full_text_parts), "source_file": os.path.basename(file_path)}})
            from opensearchpy.helpers import bulk
            indexed, errors = bulk(client, bulk_data, chunk_size=chunk_size, raise_on_error=False)
            sheet_stats.append({"sheet_name": sheet_name, "rows": len(records), "indexed": indexed, "errors": len(errors) if errors else 0})
            total_docs += indexed
        client.indices.refresh(index=index_name)
        return json.dumps({"status": "success", "index_name": index_name, "source_file": file_path, "total_documents_indexed": total_docs, "sheets": sheet_stats, "search_hint": f"Now use search_excel_opensearch with index_name='{index_name}'"}, indent=2)
    except ImportError:
        return json.dumps({"error": "pandas not installed. Install with: uv add pandas"}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to index in OpenSearch: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def search_excel_opensearch(query: str, index_name: Optional[str] = None, size: int = 10, file_path: Optional[str] = None) -> str:
    """
    Search indexed Excel data in OpenSearch using full-text search.

    Args:
        query: Search query keywords.
        index_name: OpenSearch index name. Derived from file_path if not provided.
        size: Maximum number of results to return (default: 10).
        file_path: Optional path to Excel file (used to derive index_name).

    Returns:
        JSON with search results.
    """
    if not OPENSEARCH_AVAILABLE:
        return json.dumps({"error": "OpenSearch not installed. Install with: uv add opensearch-py"}, indent=2)
    try:
        client = _get_opensearch_client()
        if client is None:
            return json.dumps({"error": "Could not connect to OpenSearch"}, indent=2)
        if index_name is None:
            if file_path:
                index_name = Path(file_path).stem.lower().replace(" ", "_")
            else:
                return json.dumps({"error": "Either index_name or file_path must be provided"}, indent=2)
        if not client.indices.exists(index=index_name):
            return json.dumps({"error": f"Index '{index_name}' not found. Index the file first using index_excel_in_opensearch."}, indent=2)
        search_body = {"size": size, "query": {"multi_match": {"query": query, "fields": ["full_text^2", "data.*"], "type": "best_fields", "fuzziness": "AUTO"}}, "highlight": {"fields": {"full_text": {}}}}
        response = client.search(index=index_name, body=search_body)
        hits = response["hits"]["hits"]
        results = []
        for hit in hits:
            source = hit["_source"]
            result = {"score": hit["_score"], "sheet_name": source.get("sheet_name"), "row_number": source.get("row_number"), "data": source.get("data"), "source_file": source.get("source_file")}
            if "highlight" in hit:
                result["highlights"] = hit["highlight"]
            results.append(result)
        return json.dumps({"query": query, "index_name": index_name, "total_hits": response["hits"]["total"]["value"], "returned": len(results), "results": results}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Search failed: {str(e)}", "query": query, "index_name": index_name}, indent=2)


@mcp.tool()
def list_opensearch_indices() -> str:
    """
    List all OpenSearch indices created for Excel files.

    Returns:
        JSON with list of indices and their stats.
    """
    if not OPENSEARCH_AVAILABLE:
        return json.dumps({"error": "OpenSearch not installed. Install with: uv add opensearch-py"}, indent=2)
    try:
        client = _get_opensearch_client()
        if client is None:
            return json.dumps({"error": "Could not connect to OpenSearch"}, indent=2)
        indices = client.cat.indices(format="json")
        excel_indices = [{"index_name": idx["index"], "doc_count": idx.get("docs.count", "0"), "size": idx.get("store.size", "0")} for idx in indices if not idx["index"].startswith(".")]
        return json.dumps({"indices": excel_indices, "total": len(excel_indices)}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to list indices: {str(e)}"}, indent=2)


@mcp.tool()
def smart_analyze_excel(file_path: str) -> str:
    """
    Intelligently analyze an Excel file, detecting all tables, text blocks,
    figures, merged cells, and named ranges across all sheets.

    Recommended first tool to call when working with an Excel file. Handles
    complex real-world spreadsheets with multiple tables per sheet.

    Args:
        file_path: Path to the Excel file.

    Returns:
        JSON with complete workbook analysis.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        result = analysis_to_dict(analysis)
        extractor.close()
        return json.dumps(result, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to analyze file: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def smart_extract_tables(file_path: str, sheet_name: Optional[str] = None, table_index: Optional[int] = None, output_format: str = "json") -> str:
    """
    Extract tables from an Excel file using intelligent table detection.

    Automatically finds multiple tables per sheet even when separated by
    empty rows/columns, instruction text, or figures.

    Args:
        file_path: Path to the Excel file.
        sheet_name: Optional sheet name. If omitted, extracts from all sheets.
        table_index: Optional specific table index (0-based) within a sheet.
        output_format: "json" (default), "csv", "markdown", or "records".

    Returns:
        JSON with extracted tables in the requested format.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    valid_formats = ["json", "csv", "markdown", "records"]
    if output_format not in valid_formats:
        return json.dumps({"error": f"Invalid output_format '{output_format}'. Must be one of: {valid_formats}"}, indent=2)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        extractor.close()
        tables_output = []
        for sheet in analysis.sheets:
            if sheet_name and sheet.sheet_name != sheet_name:
                continue
            for table in sheet.tables:
                if table_index is not None and table.table_index != table_index:
                    continue
                table_out = {"sheet_name": table.sheet_name, "table_index": table.table_index, "name": table.name, "headers": table.headers, "row_count": table.row_count, "col_count": table.col_count, "location": f"rows {table.start_row}-{table.end_row}, cols {table.start_col}-{table.end_col}"}
                if output_format == "csv":
                    table_out["data_csv"] = table_to_csv_string(table)
                elif output_format == "markdown":
                    table_out["data_markdown"] = table_to_markdown(table)
                elif output_format == "records":
                    table_out["data_records"] = table_to_records(table)
                else:
                    table_out["data"] = table.data
                tables_output.append(table_out)
        return json.dumps({"file_path": file_path, "output_format": output_format, "tables_found": len(tables_output), "tables": tables_output}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract tables: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def smart_extract_text(file_path: str, sheet_name: Optional[str] = None) -> str:
    """
    Extract standalone text blocks from an Excel file (instructions, notes, titles, labels)
    that are NOT part of data tables.

    Args:
        file_path: Path to the Excel file.
        sheet_name: Optional sheet name. If omitted, extracts from all sheets.

    Returns:
        JSON with classified text blocks (title, instruction, note, label).
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        extractor.close()
        text_output = []
        for sheet in analysis.sheets:
            if sheet_name and sheet.sheet_name != sheet_name:
                continue
            for text in sheet.text_blocks:
                text_output.append({"sheet_name": text.sheet_name, "row": text.row, "col": text.col, "text": text.text, "classification": text.classification, "is_bold": text.is_bold})
        return json.dumps({"file_path": file_path, "text_blocks_found": len(text_output), "text_blocks": text_output}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract text: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def smart_export_all(file_path: str, output_dir: Optional[str] = None, formats: Optional[str] = None) -> str:
    """
    Export all detected tables from an Excel file to multiple formats.

    Args:
        file_path: Path to the Excel file.
        output_dir: Directory to save exported files. Defaults to '<filename>_export/' next to the file.
        formats: Comma-separated list: "csv", "json", "markdown", "parquet". Defaults to "csv,json".

    Returns:
        JSON with paths to all created files and export summary.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if formats is None:
        format_list = ["csv", "json"]
    else:
        format_list = [f.strip().lower() for f in formats.split(",")]
    valid_formats = {"csv", "json", "markdown", "parquet"}
    invalid = set(format_list) - valid_formats
    if invalid:
        return json.dumps({"error": f"Invalid formats: {invalid}. Supported: {valid_formats}"}, indent=2)
    if output_dir is None:
        base = Path(file_path).stem
        output_dir = os.path.join(os.path.dirname(file_path) or ".", f"{base}_export")
    os.makedirs(output_dir, exist_ok=True)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        extractor.close()
        created_files = []
        base_name = Path(file_path).stem
        summary_path = os.path.join(output_dir, f"{base_name}_analysis.json")
        analysis_dict = analysis_to_dict(analysis)
        with open(summary_path, "w", encoding="utf-8") as f:
            json.dump(analysis_dict, f, indent=2, ensure_ascii=False, default=str)
        created_files.append({"type": "analysis_summary", "path": summary_path})
        for sheet in analysis.sheets:
            safe_sheet = "".join(c if c.isalnum() else "_" for c in sheet.sheet_name)
            for table in sheet.tables:
                table_label = table.name or f"table_{table.table_index}"
                safe_label = "".join(c if c.isalnum() else "_" for c in table_label)
                file_prefix = f"{base_name}_{safe_sheet}_{safe_label}"
                if "csv" in format_list:
                    csv_path = os.path.join(output_dir, f"{file_prefix}.csv")
                    with open(csv_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.writer(f)
                        if table.has_header:
                            writer.writerow(table.headers)
                        for row in table.data:
                            writer.writerow([str(v) if v is not None else "" for v in row])
                    created_files.append({"type": "csv", "sheet": sheet.sheet_name, "table": table_label, "path": csv_path, "rows": table.row_count})
                if "json" in format_list:
                    json_path = os.path.join(output_dir, f"{file_prefix}.json")
                    records = table_to_records(table)
                    with open(json_path, "w", encoding="utf-8") as f:
                        json.dump({"sheet_name": table.sheet_name, "table_name": table.name, "headers": table.headers, "row_count": table.row_count, "data": records}, f, indent=2, ensure_ascii=False, default=str)
                    created_files.append({"type": "json", "sheet": sheet.sheet_name, "table": table_label, "path": json_path, "rows": table.row_count})
                if "markdown" in format_list:
                    md_path = os.path.join(output_dir, f"{file_prefix}.md")
                    md_content = table_to_markdown(table)
                    with open(md_path, "w", encoding="utf-8") as f:
                        if table.name:
                            f.write(f"## {table.name}\n\n")
                        f.write(md_content + "\n")
                    created_files.append({"type": "markdown", "sheet": sheet.sheet_name, "table": table_label, "path": md_path, "rows": table.row_count})
                if "parquet" in format_list:
                    try:
                        import pandas as pd
                        pq_path = os.path.join(output_dir, f"{file_prefix}.parquet")
                        pd.DataFrame(table_to_records(table)).to_parquet(pq_path, index=False)
                        created_files.append({"type": "parquet", "sheet": sheet.sheet_name, "table": table_label, "path": pq_path, "rows": table.row_count})
                    except ImportError:
                        created_files.append({"type": "parquet", "error": "pandas or pyarrow not installed"})
        text_blocks = []
        for sheet in analysis.sheets:
            for text in sheet.text_blocks:
                text_blocks.append({"sheet_name": text.sheet_name, "row": text.row, "col": text.col, "text": text.text, "classification": text.classification})
        if text_blocks:
            text_path = os.path.join(output_dir, f"{base_name}_text_blocks.json")
            with open(text_path, "w", encoding="utf-8") as f:
                json.dump(text_blocks, f, indent=2, ensure_ascii=False, default=str)
            created_files.append({"type": "text_blocks", "path": text_path, "count": len(text_blocks)})
        return json.dumps({"status": "success", "source_file": file_path, "output_directory": output_dir, "total_tables_exported": analysis.total_tables, "total_files_created": len(created_files), "formats": format_list, "files": created_files}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to export: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def smart_get_sheet_overview(file_path: str, sheet_name: str) -> str:
    """
    Get a detailed spatial overview of a specific Excel sheet.

    Args:
        file_path: Path to the Excel file.
        sheet_name: Name of the sheet to analyze.

    Returns:
        JSON with spatial layout of all content (tables, text, figures, merged regions).
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        extractor.close()
        target_sheet = next((s for s in analysis.sheets if s.sheet_name == sheet_name), None)
        if target_sheet is None:
            return json.dumps({"error": f"Sheet '{sheet_name}' not found", "available_sheets": analysis.sheet_names}, indent=2)
        content_map = []
        for table in target_sheet.tables:
            content_map.append({"type": "table", "name": table.name, "table_index": table.table_index, "location": {"start_row": table.start_row, "end_row": table.end_row, "start_col": table.start_col, "end_col": table.end_col}, "size": f"{table.row_count} rows x {table.col_count} cols", "headers": table.headers, "sample_data": table.data[:3] if table.data else []})
        for text in target_sheet.text_blocks:
            content_map.append({"type": "text", "classification": text.classification, "location": {"row": text.row, "col": text.col}, "text": text.text[:200] + ("..." if len(text.text) > 200 else ""), "is_bold": text.is_bold})
        for fig in target_sheet.figures:
            content_map.append({"type": "figure", "chart_type": fig.chart_type, "title": fig.title, "series_count": fig.series_count})
        for merged in target_sheet.merged_cells:
            content_map.append({"type": "merged_region", "range": merged["range"], "value": merged["value"]})

        def sort_key(item):
            loc = item.get("location", {})
            if isinstance(loc, dict):
                return (loc.get("start_row", loc.get("row", 999)), loc.get("start_col", loc.get("col", 999)))
            return (999, 999)

        content_map.sort(key=sort_key)
        return json.dumps({"file_path": file_path, "sheet_name": sheet_name, "dimensions": f"{target_sheet.total_rows} rows x {target_sheet.total_cols} cols", "summary": {"tables": len(target_sheet.tables), "text_blocks": len(target_sheet.text_blocks), "figures": len(target_sheet.figures), "merged_regions": len(target_sheet.merged_cells)}, "content_map": content_map}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to get sheet overview: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def smart_table_to_format(file_path: str, sheet_name: str, table_index: int, output_format: str = "csv") -> str:
    """
    Convert a specific detected Excel table to a chosen format inline (no file written).

    Args:
        file_path: Path to the Excel file.
        sheet_name: Name of the sheet containing the table.
        table_index: Index of the table within the sheet (0-based).
        output_format: "csv", "json", "markdown", or "records".

    Returns:
        The table content in the requested format as a string.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    valid_formats = ["csv", "json", "markdown", "records"]
    if output_format not in valid_formats:
        return json.dumps({"error": f"Invalid format '{output_format}'. Use one of: {valid_formats}"}, indent=2)
    try:
        extractor = SmartExcelExtractor(file_path)
        analysis = extractor.analyze()
        extractor.close()
        target_table = None
        for sheet in analysis.sheets:
            if sheet.sheet_name == sheet_name:
                for table in sheet.tables:
                    if table.table_index == table_index:
                        target_table = table
                        break
                break
        if target_table is None:
            available = [{"sheet": s.sheet_name, "table_index": t.table_index, "name": t.name, "rows": t.row_count} for s in analysis.sheets for t in s.tables]
            return json.dumps({"error": f"Table index {table_index} not found in sheet '{sheet_name}'", "available_tables": available}, indent=2)
        if output_format == "csv":
            return table_to_csv_string(target_table)
        elif output_format == "markdown":
            return table_to_markdown(target_table)
        elif output_format == "records":
            return json.dumps(table_to_records(target_table), indent=2, default=str)
        else:
            return json.dumps({"headers": target_table.headers, "data": target_table.data, "row_count": target_table.row_count}, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to convert table: {str(e)}", "file_path": file_path}, indent=2)


# ─── PDF MCP Tools ────────────────────────────────────────────────────────────

@mcp.tool()
def list_pdf_files(directory: Optional[str] = None) -> str:
    """
    List all PDF files in the specified directory with metadata.

    Args:
        directory: Path to directory. Defaults to PDF_DIR env var or '../pdf_files'.

    Returns:
        JSON with file list including size, page count, and recommended engine.
    """
    search_dir = directory or PDF_DIR
    if not os.path.exists(search_dir):
        return json.dumps({"error": f"Directory not found: {search_dir}", "hint": "Create the directory or set PDF_DIR environment variable"}, indent=2)
    files = []
    for root, dirs, filenames in os.walk(search_dir):
        dirs.sort()
        for item in sorted(filenames):
            if not item.lower().endswith(".pdf"):
                continue
            file_path = os.path.join(root, item)
            rel_path = os.path.relpath(file_path, search_dir)
            try:
                stat = os.stat(file_path)
                size_mb = round(stat.st_size / (1024 * 1024), 2)
                page_count = None
                if PYPDF_AVAILABLE:
                    try:
                        page_count = len(PdfReader(file_path).pages)
                    except Exception:
                        pass
                files.append({"filename": item, "relative_path": rel_path, "path": file_path, "size_mb": size_mb, "pages": page_count, "hash": _file_hash(file_path)})
            except OSError as e:
                files.append({"filename": item, "relative_path": rel_path, "error": str(e)})
    return json.dumps({"directory": os.path.abspath(search_dir), "file_count": len(files), "total_size_mb": round(sum(f.get("size_mb", 0) for f in files), 2), "total_pages": sum(f.get("pages", 0) or 0 for f in files), "engine": "docling", "files": files}, indent=2)


@mcp.tool()
def extract_pdf(file_path: str, engine: str = "auto", page_range: Optional[str] = None, use_ai_ocr: bool = False) -> str:
    """
    Extract content from a PDF file as clean Markdown with per-page metadata.

    Handles regular text, tables, and scanned/OCR pages. Figures are detected
    and catalogued; use describe_pdf_figures for AI-powered figure descriptions.

    Args:
        file_path: Path to the PDF file.
        engine: Kept for API compatibility; Docling is always used.
        page_range: Optional page range like '1-5' or '1-10,15,20-25'.
        use_ai_ocr: If True and ANTHROPIC_API_KEY is set, uses Claude vision for scanned pages.

    Returns:
        JSON with extracted Markdown content, per-page metadata, tables, figures, and OCR info.
    """
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not DOCLING_AVAILABLE:
        # Lightweight pypdf fallback — text only, no tables/OCR
        try:
            reader = PdfReader(file_path)
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = (page.extract_text() or "").strip()
                if text:
                    pages_text.append(f"<!-- Page {i+1} -->\n{text}")
            markdown = "\n\n".join(pages_text)
            return json.dumps({
                "file_path": file_path,
                "content_markdown": markdown,
                "total_pages": len(reader.pages),
                "total_words": len(markdown.split()),
                "total_tables": 0,
                "total_figures": 0,
                "engine": "pypdf",
            }, indent=2)
        except Exception as e:
            return json.dumps({"error": f"pypdf extraction failed: {e}"}, indent=2)
    try:
        result = _extract_pdf_with_docling(file_path)
        return json.dumps(asdict(result), indent=2)
    except Exception as e:
        return json.dumps({"error": f"Extraction failed: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def extract_pdf_tables(file_path: str, page_range: Optional[str] = None) -> str:
    """
    Extract only tables from a PDF file as Markdown tables.

    Args:
        file_path: Path to the PDF file.
        page_range: Optional page range like '1-5'.

    Returns:
        JSON with list of tables as Markdown, including page numbers and dimensions.
    """
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "Docling not installed. Run: uv add docling"}, indent=2)
    try:
        pages_filter = _parse_page_range_1indexed(page_range) if page_range else None
        conv = _get_docling_converter()
        result = conv.convert(file_path)
        doc = result.document
        page_count = len(result.pages) if hasattr(result, "pages") else 0
        tables = []
        for i, tbl in enumerate(doc.tables):
            page_no = tbl.prov[0].page_no if tbl.prov else None
            if pages_filter and page_no not in pages_filter:
                continue
            entry: Dict[str, Any] = {"table_index": i}
            if page_no is not None:
                entry["page"] = page_no
            try:
                entry["markdown"] = tbl.export_to_markdown()
            except Exception:
                entry["markdown"] = ""
            if hasattr(tbl, "num_rows"):
                entry["rows"] = tbl.num_rows
            if hasattr(tbl, "num_cols"):
                entry["cols"] = tbl.num_cols
            tables.append(entry)
        return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "total_pages": page_count, "pages_filter": page_range, "tables_found": len(tables), "tables": tables}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Table extraction failed: {str(e)}"}, indent=2)


@mcp.tool()
def extract_pdf_figures(file_path: str, page_range: Optional[str] = None) -> str:
    """
    Extract figure/image metadata from a PDF file.

    Returns metadata about embedded images (dimensions, colorspace, page).
    Does not extract actual image data. Use describe_pdf_figures for AI descriptions.

    Args:
        file_path: Path to the PDF file.
        page_range: Optional page range like '1-5'.

    Returns:
        JSON with list of figures and their metadata.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "Docling not installed. Run: uv add docling"}, indent=2)
    try:
        pages_filter = _parse_page_range_1indexed(page_range) if page_range else None
        conv = _get_docling_converter()
        result = conv.convert(file_path)
        doc = result.document
        page_count = len(result.pages) if hasattr(result, "pages") else 0
        figures = []
        for i, pic in enumerate(doc.pictures):
            page_no = pic.prov[0].page_no if pic.prov else None
            if pages_filter and page_no not in pages_filter:
                continue
            entry: Dict[str, Any] = {"figure_index": i}
            if page_no is not None:
                entry["page"] = page_no
            if hasattr(pic, "caption") and pic.caption:
                entry["caption"] = str(pic.caption)
            if hasattr(pic, "image") and pic.image and hasattr(pic.image, "size"):
                w, h = pic.image.size
                entry["width_px"] = w
                entry["height_px"] = h
            figures.append(entry)
        return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "total_pages": page_count, "pages_filter": page_range, "figures_found": len(figures), "figures": figures}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Figure extraction failed: {str(e)}"}, indent=2)


@mcp.tool()
def chunk_pdf_for_rag(
    file_path: str,
    strategy: str = "tokens",
    chunk_size: int = 512,
    overlap: int = 50,
    engine: str = "auto",
    page_range: Optional[str] = None,
) -> str:
    """
    Extract and chunk a PDF into pieces ready for a RAG vector database.

    Args:
        file_path: Path to the PDF file.
        strategy: 'tokens' (word-based), 'pages', or 'sections' (by headings).
        chunk_size: Words per chunk for 'tokens' (default 512). Pages per chunk for 'pages'.
        overlap: Word overlap between chunks for 'tokens' strategy (default 50).
        engine: Kept for API compatibility; Docling is always used.
        page_range: Optional page range like '1-50'.

    Returns:
        JSON with chunks array, each containing text, metadata, and position info.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    extraction_json = extract_pdf(file_path, engine=engine, page_range=page_range)
    extraction = json.loads(extraction_json)
    if "error" in extraction:
        return extraction_json
    markdown = extraction.get("content_markdown", "")
    if not markdown:
        return json.dumps({"error": "No content extracted from PDF"}, indent=2)
    if strategy == "tokens":
        chunks = _chunk_by_tokens(markdown, chunk_size=chunk_size, overlap=overlap)
    elif strategy == "pages":
        chunks = _chunk_by_pages(markdown, pages_per_chunk=chunk_size if chunk_size < 50 else 5)
    elif strategy == "sections":
        chunks = _chunk_by_sections(markdown)
    else:
        return json.dumps({"error": f"Unknown strategy: {strategy}. Use 'tokens', 'pages', or 'sections'."}, indent=2)
    doc_id = _file_hash(file_path)
    for chunk in chunks:
        chunk["doc_id"] = doc_id
        chunk["doc_name"] = os.path.basename(file_path)
        chunk["source"] = file_path
    return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "doc_id": doc_id, "engine": extraction.get("engine"), "strategy": strategy, "total_pages": extraction.get("total_pages"), "total_words": extraction.get("total_words"), "total_tables": extraction.get("total_tables"), "total_figures": extraction.get("total_figures"), "chunk_count": len(chunks), "chunks": chunks}, indent=2)


@mcp.tool()
def batch_extract_pdfs(directory: Optional[str] = None, engine: str = "auto", output_format: str = "markdown") -> str:
    """
    Batch extract all PDFs in a directory.

    Args:
        directory: Directory containing PDFs. Defaults to PDF_DIR.
        engine: Kept for API compatibility; Docling is always used.
        output_format: 'markdown' or 'json'.

    Returns:
        JSON summary with extraction results per file.
    """
    search_dir = directory or PDF_DIR
    if not os.path.exists(search_dir):
        return json.dumps({"error": f"Directory not found: {search_dir}"}, indent=2)
    pdf_files = []
    for root, dirs, filenames in os.walk(search_dir):
        dirs.sort()
        for f in sorted(filenames):
            if f.lower().endswith(".pdf"):
                pdf_files.append(os.path.join(root, f))
    if not pdf_files:
        return json.dumps({"error": "No PDF files found", "directory": search_dir}, indent=2)
    out_dir = OUTPUT_DIR
    os.makedirs(out_dir, exist_ok=True)
    results = []
    total_start = time.time()
    for file_path in pdf_files:
        try:
            extraction_json = extract_pdf(file_path, engine=engine)
            extraction = json.loads(extraction_json)
            if "error" in extraction:
                results.append({"file": os.path.basename(file_path), "status": "error", "error": extraction["error"]})
                continue
            base_name = Path(file_path).stem
            if output_format == "markdown":
                out_path = os.path.join(out_dir, f"{base_name}.md")
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(extraction.get("content_markdown", ""))
            else:
                out_path = os.path.join(out_dir, f"{base_name}.json")
                with open(out_path, "w", encoding="utf-8") as f:
                    json.dump(extraction, f, indent=2)
            results.append({"file": os.path.basename(file_path), "status": "success", "engine": extraction.get("engine"), "pages": extraction.get("total_pages"), "words": extraction.get("total_words"), "tables": extraction.get("total_tables"), "figures": extraction.get("total_figures"), "time_seconds": extraction.get("extraction_time_seconds"), "output": out_path})
        except Exception as e:
            results.append({"file": os.path.basename(file_path), "status": "error", "error": str(e)})
    total_elapsed = round(time.time() - total_start, 2)
    succeeded = [r for r in results if r["status"] == "success"]
    failed = [r for r in results if r["status"] == "error"]
    return json.dumps({"directory": os.path.abspath(search_dir), "output_directory": os.path.abspath(out_dir), "total_files": len(pdf_files), "succeeded": len(succeeded), "failed": len(failed), "total_pages": sum(r.get("pages", 0) or 0 for r in succeeded), "total_words": sum(r.get("words", 0) or 0 for r in succeeded), "total_time_seconds": total_elapsed, "results": results}, indent=2)


@mcp.tool()
def analyze_pdf_structure(file_path: str) -> str:
    """
    Analyze a PDF's structure without full extraction.

    Quick scan reporting page count, tables/figures per page, complexity,
    and recommended extraction engine.

    Args:
        file_path: Path to the PDF file.

    Returns:
        JSON with structural analysis and recommendations.
    """
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "Docling not installed. Run: uv add docling"}, indent=2)
    try:
        stat = os.stat(file_path)
        extraction = _extract_pdf_with_docling(file_path)
        complexity = (
            "high" if (extraction.total_tables > 10 or extraction.total_figures > 20)
            else "medium" if (extraction.total_tables > 3 or extraction.total_figures > 5)
            else "low"
        )
        return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "size_mb": round(stat.st_size / (1024 * 1024), 2), "total_pages": extraction.total_pages, "total_words": extraction.total_words, "total_tables": extraction.total_tables, "total_figures": extraction.total_figures, "complexity": complexity, "engine": "docling", "extraction_time_seconds": extraction.extraction_time_seconds}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Analysis failed: {str(e)}"}, indent=2)


@mcp.tool()
def describe_pdf_figures(file_path: str, page_range: Optional[str] = None, max_figures: int = 10) -> str:
    """
    Extract figures/charts from a PDF and describe their content using Claude vision AI.

    Args:
        file_path: Path to the PDF file.
        page_range: Optional page range like '1-5'.
        max_figures: Maximum number of figures to describe (default 10).

    Returns:
        JSON with list of figures, each with page number, dimensions, and AI description.
    """
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not DOCLING_AVAILABLE:
        return json.dumps({"error": "Docling not installed. Run: uv add docling"}, indent=2)
    if not ANTHROPIC_AVAILABLE:
        return json.dumps({"error": "Anthropic package not available. Install: uv add anthropic", "fallback": "Use extract_pdf_figures for metadata-only listing."}, indent=2)
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return json.dumps({"error": "ANTHROPIC_API_KEY not set.", "hint": "Set ANTHROPIC_API_KEY to enable AI figure descriptions."}, indent=2)
    try:
        import base64, io as _io
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import PdfFormatOption
        opts = PdfPipelineOptions()
        opts.generate_picture_images = True
        fig_converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        )
        result = fig_converter.convert(file_path)
        doc = result.document
        page_count = len(result.pages) if hasattr(result, "pages") else 0
        pages_filter = _parse_page_range_1indexed(page_range) if page_range else None
        client = anthropic.Anthropic(api_key=api_key)
        described_figures = []
        figure_count = 0
        describe_prompt = ("This image is from a PDF document. Describe what it shows concisely in 1-3 sentences. "
                           "If it is a chart or graph, mention the type, subject, and key values or trends visible. "
                           "If it is a diagram or photo, describe its content briefly.")
        for i, pic in enumerate(doc.pictures):
            if figure_count >= max_figures:
                break
            page_no = pic.prov[0].page_no if pic.prov else None
            if pages_filter and page_no not in pages_filter:
                continue
            pil_img = pic.image.pil_image if (hasattr(pic, "image") and pic.image and hasattr(pic.image, "pil_image")) else None
            if pil_img is None:
                continue
            w, h = pil_img.size
            if w < 50 or h < 50:
                continue
            buf = _io.BytesIO()
            pil_img.save(buf, format="PNG")
            img_b64 = base64.standard_b64encode(buf.getvalue()).decode("utf-8")
            try:
                response = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=300,
                    messages=[{"role": "user", "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img_b64}},
                        {"type": "text", "text": describe_prompt},
                    ]}],
                )
                description = response.content[0].text.strip() if response.content else ""
                described_figures.append({"page": page_no, "figure_index": i, "width_px": w, "height_px": h, "description": description})
                figure_count += 1
            except Exception as exc:
                described_figures.append({"page": page_no, "figure_index": i, "error": str(exc)})
        return json.dumps({"file_path": file_path, "file_name": os.path.basename(file_path), "total_pages": page_count, "figures_described": len(described_figures), "max_figures_limit": max_figures, "truncated": figure_count >= max_figures, "model": "claude-haiku-4-5-20251001", "figures": described_figures}, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Figure description failed: {str(e)}"}, indent=2)


# ─── Word MCP Tools ───────────────────────────────────────────────────────────

@mcp.tool()
def list_word_files(directory: Optional[str] = None) -> str:
    """
    List all Word files (.docx, .doc) in the specified directory.

    Args:
        directory: Path to directory. Defaults to WORD_DIR env var or 'word_files'.

    Returns:
        JSON with list of Word files and their metadata.
    """
    search_dir = directory or WORD_DIR
    if not os.path.exists(search_dir):
        return json.dumps({"error": f"Directory not found: {search_dir}", "hint": "Create the directory or set WORD_DIR environment variable"}, indent=2)
    word_extensions = [".docx", ".doc", ".docm"]
    files = []
    for item in sorted(os.listdir(search_dir)):
        if any(item.lower().endswith(ext) for ext in word_extensions):
            file_path = os.path.join(search_dir, item)
            try:
                stat = os.stat(file_path)
                files.append({"filename": item, "path": file_path, "size_bytes": stat.st_size, "modified": str(stat.st_mtime), "docx_supported": item.lower().endswith(".docx")})
            except OSError as e:
                files.append({"filename": item, "path": file_path, "error": str(e)})
    return json.dumps({"directory": search_dir, "file_count": len(files), "note": ".doc files require conversion to .docx for full extraction", "files": files}, indent=2)


@mcp.tool()
def extract_word_content(file_path: str) -> str:
    """
    Extract full content from a Word (.docx) file including paragraphs, tables,
    headings, and document metadata.

    Args:
        file_path: Path to the .docx file.

    Returns:
        JSON with structured document content: paragraphs, tables, headings, metadata.
    """
    if not DOCX_AVAILABLE:
        return json.dumps({"error": "python-docx not installed. Install with: uv add python-docx"}, indent=2)
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not file_path.lower().endswith(".docx"):
        return json.dumps({"error": "Only .docx files are supported. Convert .doc files to .docx first."}, indent=2)
    try:
        doc = DocxDocument(file_path)
        # Core properties / metadata
        props = doc.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision,
            "keywords": props.keywords or "",
        }
        # Paragraphs with heading classification
        paragraphs = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            level = _heading_level(para)
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "heading_level": level,
                "is_bold": any(run.bold for run in para.runs),
                "is_italic": any(run.italic for run in para.runs),
            })
        # Tables
        tables = []
        for i, table in enumerate(doc.tables):
            rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]
            headers = rows[0] if rows else []
            tables.append({
                "table_index": i,
                "row_count": len(rows),
                "col_count": len(headers),
                "headers": headers,
                "data": rows[1:] if len(rows) > 1 else [],
                "markdown": _word_table_to_markdown(table),
            })
        # Section info
        sections = []
        for i, section in enumerate(doc.sections):
            sections.append({
                "index": i,
                "page_width_inches": round(section.page_width.inches, 2) if section.page_width else None,
                "page_height_inches": round(section.page_height.inches, 2) if section.page_height else None,
                "orientation": "landscape" if (section.page_width and section.page_height and section.page_width > section.page_height) else "portrait",
            })
        return json.dumps({
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "metadata": metadata,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "section_count": len(sections),
            "paragraphs": paragraphs,
            "tables": tables,
            "sections": sections,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract Word content: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def get_word_summary(file_path: str) -> str:
    """
    Get a structural summary of a Word (.docx) file without returning full content.

    Args:
        file_path: Path to the .docx file.

    Returns:
        JSON with document structure: heading outline, table dimensions, word/paragraph counts.
    """
    if not DOCX_AVAILABLE:
        return json.dumps({"error": "python-docx not installed. Install with: uv add python-docx"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not file_path.lower().endswith(".docx"):
        return json.dumps({"error": "Only .docx files are supported."}, indent=2)
    try:
        doc = DocxDocument(file_path)
        props = doc.core_properties
        total_words = 0
        heading_outline = []
        para_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            para_count += 1
            total_words += len(text.split())
            level = _heading_level(para)
            if level:
                heading_outline.append({"level": level, "text": text[:120]})
        tables_summary = []
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            first_row = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            tables_summary.append({"table_index": i, "rows": rows, "cols": cols, "headers": first_row})
        stat = os.stat(file_path)
        return json.dumps({
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "size_bytes": stat.st_size,
            "title": props.title or "",
            "author": props.author or "",
            "modified": str(props.modified) if props.modified else "",
            "paragraph_count": para_count,
            "estimated_word_count": total_words,
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "heading_count": len(heading_outline),
            "heading_outline": heading_outline,
            "tables": tables_summary,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to summarize Word document: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def word_to_markdown(file_path: str) -> str:
    """
    Convert a Word (.docx) file to Markdown format.

    Preserves headings, bold/italic text, bullet lists, numbered lists, and tables.

    Args:
        file_path: Path to the .docx file.

    Returns:
        JSON with the Markdown string and conversion statistics.
    """
    if not DOCX_AVAILABLE:
        return json.dumps({"error": "python-docx not installed. Install with: uv add python-docx"}, indent=2)
    try:
        file_path = _validate_path(file_path)
    except ValueError as e:
        return json.dumps({"error": str(e)}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not file_path.lower().endswith(".docx"):
        return json.dumps({"error": "Only .docx files are supported."}, indent=2)
    try:
        doc = DocxDocument(file_path)
        md_lines = []
        tables_in_doc = {id(t): t for t in doc.tables}
        processed_tables = set()

        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

            if tag == "p":
                # Find matching paragraph object
                para_text = "".join(r.text for r in block.iterchildren("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r") if r is not None)
                style_name = ""
                pPr = block.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
                if pPr is not None:
                    pStyle = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
                    if pStyle is not None:
                        style_name = pStyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")

                if not para_text.strip():
                    md_lines.append("")
                    continue

                # Determine heading level from style name
                heading_prefix = ""
                if style_name.lower().startswith("heading"):
                    try:
                        lvl = int(style_name[-1])
                        heading_prefix = "#" * min(lvl, 6) + " "
                    except (ValueError, IndexError):
                        pass

                md_lines.append(heading_prefix + para_text.strip())

            elif tag == "tbl":
                # Find matching table
                for tbl_id, tbl_obj in tables_in_doc.items():
                    if tbl_obj._element is block and tbl_id not in processed_tables:
                        processed_tables.add(tbl_id)
                        md_lines.append("")
                        md_lines.append(_word_table_to_markdown(tbl_obj))
                        md_lines.append("")
                        break

        markdown_text = "\n".join(md_lines).strip()
        word_count = len(markdown_text.split())
        table_count = len(processed_tables)
        return json.dumps({
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "word_count": word_count,
            "table_count": table_count,
            "markdown": markdown_text,
        }, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to convert Word to Markdown: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def extract_word_tables(file_path: str) -> str:
    """
    Extract all tables from a Word (.docx) file.

    Args:
        file_path: Path to the .docx file.

    Returns:
        JSON with all tables as Markdown and structured data.
    """
    if not DOCX_AVAILABLE:
        return json.dumps({"error": "python-docx not installed. Install with: uv add python-docx"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not file_path.lower().endswith(".docx"):
        return json.dumps({"error": "Only .docx files are supported."}, indent=2)
    try:
        doc = DocxDocument(file_path)
        tables = []
        for i, table in enumerate(doc.tables):
            rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]
            headers = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []
            records = [dict(zip(headers, row)) for row in data] if headers else [{"row": row} for row in data]
            tables.append({
                "table_index": i,
                "row_count": len(rows),
                "col_count": len(headers),
                "headers": headers,
                "data": data,
                "records": records,
                "markdown": _word_table_to_markdown(table),
            })
        return json.dumps({
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "tables_found": len(tables),
            "tables": tables,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract Word tables: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def word_to_text(file_path: str) -> str:
    """
    Extract plain text from a Word (.docx) file, preserving paragraph structure.

    Args:
        file_path: Path to the .docx file.

    Returns:
        JSON with plain text content and basic statistics.
    """
    if not DOCX_AVAILABLE:
        return json.dumps({"error": "python-docx not installed. Install with: uv add python-docx"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    if not file_path.lower().endswith(".docx"):
        return json.dumps({"error": "Only .docx files are supported."}, indent=2)
    try:
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
        word_count = len(full_text.split())
        char_count = len(full_text)
        return json.dumps({
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "paragraph_count": len(paragraphs),
            "word_count": word_count,
            "character_count": char_count,
            "text": full_text,
        }, indent=2)
    except Exception as e:
        return json.dumps({"error": f"Failed to extract Word text: {str(e)}", "file_path": file_path}, indent=2)


# ─── Arrow / Parquet Compute Tools ────────────────────────────────────────────


@mcp.tool()
def read_parquet(file_path: str, columns: Optional[str] = None, row_limit: Optional[int] = None) -> str:
    """
    Read a Parquet file and return its contents as JSON.

    Args:
        file_path: Path to the Parquet file.
        columns: Comma-separated column names to read. Reads all if omitted.
        row_limit: Maximum number of rows to return. Returns all if omitted.

    Returns:
        JSON with schema, row count, and data rows.
    """
    if not PYARROW_AVAILABLE:
        return json.dumps({"error": "pyarrow is not installed. Run: uv add pyarrow"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        col_list = [c.strip() for c in columns.split(",")] if columns else None
        table = pq.read_table(file_path, columns=col_list)
        total_rows = table.num_rows
        if row_limit and row_limit < total_rows:
            table = table.slice(0, row_limit)
        schema_info = [{"name": field.name, "type": str(field.type)} for field in table.schema]
        rows = table.to_pydict()
        # Convert column-oriented dict to row-oriented list
        row_count = table.num_rows
        data = []
        for i in range(row_count):
            data.append({col: rows[col][i] for col in rows})
        return json.dumps({
            "file_path": file_path,
            "total_rows": total_rows,
            "returned_rows": row_count,
            "num_columns": len(table.schema),
            "schema": schema_info,
            "data": data,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to read Parquet: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def parquet_schema(file_path: str) -> str:
    """
    Get the schema (column names and types) of a Parquet file without reading data.

    Args:
        file_path: Path to the Parquet file.

    Returns:
        JSON with column names, types, row count, and file metadata.
    """
    if not PYARROW_AVAILABLE:
        return json.dumps({"error": "pyarrow is not installed. Run: uv add pyarrow"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        pf = pq.ParquetFile(file_path)
        schema = pf.schema_arrow
        metadata = pf.metadata
        schema_info = [{"name": field.name, "type": str(field.type), "nullable": field.nullable} for field in schema]
        return json.dumps({
            "file_path": file_path,
            "num_columns": len(schema),
            "num_rows": metadata.num_rows,
            "num_row_groups": metadata.num_row_groups,
            "created_by": metadata.created_by,
            "schema": schema_info,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Failed to read schema: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def query_parquet(file_path: str, columns: Optional[str] = None, filter_expr: Optional[str] = None,
                  sort_by: Optional[str] = None, sort_descending: bool = False, row_limit: Optional[int] = None) -> str:
    """
    Query a Parquet file with column selection, filtering, and sorting using Arrow compute.

    Args:
        file_path: Path to the Parquet file.
        columns: Comma-separated column names to include in results. All if omitted.
        filter_expr: Filter expression in the form 'column operator value', e.g.
                     'age > 30', 'status == active', 'price <= 100.5'.
                     Supported operators: ==, !=, >, <, >=, <=.
        sort_by: Column name to sort results by.
        sort_descending: Sort in descending order (default ascending).
        row_limit: Maximum rows to return.

    Returns:
        JSON with matching rows — column names, types, and data.
    """
    if not PYARROW_AVAILABLE:
        return json.dumps({"error": "pyarrow is not installed. Run: uv add pyarrow"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        table = pq.read_table(file_path)
        total_rows = table.num_rows

        # Apply filter
        if filter_expr:
            table = _apply_arrow_filter(table, filter_expr)

        # Select columns
        if columns:
            col_list = [c.strip() for c in columns.split(",")]
            missing = [c for c in col_list if c not in table.column_names]
            if missing:
                return json.dumps({"error": f"Columns not found: {missing}", "available": table.column_names}, indent=2)
            table = table.select(col_list)

        # Sort
        if sort_by:
            if sort_by not in table.column_names:
                return json.dumps({"error": f"Sort column '{sort_by}' not found", "available": table.column_names}, indent=2)
            indices = pc.sort_indices(table, sort_keys=[(sort_by, "descending" if sort_descending else "ascending")])
            table = table.take(indices)

        filtered_rows = table.num_rows
        if row_limit and row_limit < filtered_rows:
            table = table.slice(0, row_limit)

        rows = table.to_pydict()
        row_count = table.num_rows
        data = []
        for i in range(row_count):
            data.append({col: rows[col][i] for col in rows})

        return json.dumps({
            "file_path": file_path,
            "total_rows_in_file": total_rows,
            "rows_after_filter": filtered_rows,
            "returned_rows": row_count,
            "data": data,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Query failed: {str(e)}", "file_path": file_path}, indent=2)


@mcp.tool()
def aggregate_parquet(file_path: str, group_by: Optional[str] = None,
                      aggregations: str = "count") -> str:
    """
    Compute aggregations on a Parquet file using Arrow compute.

    Args:
        file_path: Path to the Parquet file.
        group_by: Column name to group by. If omitted, aggregates over the entire table.
        aggregations: Comma-separated list of 'column:function' pairs.
                      Functions: count, sum, mean, min, max.
                      Use just 'count' for row count. Examples:
                      'count', 'price:sum,price:mean', 'age:min,age:max,salary:sum'.

    Returns:
        JSON with aggregation results.
    """
    if not PYARROW_AVAILABLE:
        return json.dumps({"error": "pyarrow is not installed. Run: uv add pyarrow"}, indent=2)
    if not os.path.exists(file_path):
        return json.dumps({"error": f"File not found: {file_path}"}, indent=2)
    try:
        table = pq.read_table(file_path)
        agg_specs = _parse_aggregations(aggregations, table.column_names)

        if group_by:
            if group_by not in table.column_names:
                return json.dumps({"error": f"Group column '{group_by}' not found", "available": table.column_names}, indent=2)
            result = table.group_by(group_by).aggregate(agg_specs)
        else:
            # Whole-table aggregation
            results = {}
            for col, func in agg_specs:
                label = f"{col}_{func}" if col else func
                if func == "hash_count_all":
                    results["count"] = table.num_rows
                elif col:
                    column = table.column(col)
                    value = _compute_agg(column, func)
                    results[label] = value
            return json.dumps({
                "file_path": file_path,
                "total_rows": table.num_rows,
                "aggregations": results,
            }, indent=2, default=str)

        # Convert grouped result
        rows = result.to_pydict()
        row_count = result.num_rows
        data = []
        for i in range(row_count):
            data.append({col: rows[col][i] for col in rows})
        return json.dumps({
            "file_path": file_path,
            "total_rows": table.num_rows,
            "group_by": group_by,
            "groups": row_count,
            "data": data,
        }, indent=2, default=str)
    except Exception as e:
        return json.dumps({"error": f"Aggregation failed: {str(e)}", "file_path": file_path}, indent=2)


def _apply_arrow_filter(table: "pa.Table", expr: str) -> "pa.Table":
    """Parse a simple filter expression and apply it to an Arrow table."""
    op_map = {
        ">=": pc.greater_equal, "<=": pc.less_equal,
        "!=": pc.not_equal,     "==": pc.equal,
        ">":  pc.greater,       "<":  pc.less,
    }
    for symbol in (">=", "<=", "!=", "==", ">", "<"):
        if symbol in expr:
            parts = expr.split(symbol, 1)
            col_name = parts[0].strip()
            raw_value = parts[1].strip()
            if col_name not in table.column_names:
                raise ValueError(f"Filter column '{col_name}' not found. Available: {table.column_names}")
            column = table.column(col_name)
            # Infer type from column
            if pa.types.is_integer(column.type):
                value = int(raw_value)
            elif pa.types.is_floating(column.type):
                value = float(raw_value)
            else:
                value = raw_value.strip("'\"")
            mask = op_map[symbol](column, pa.scalar(value))
            return table.filter(mask)
    raise ValueError(f"Could not parse filter expression: '{expr}'. Use 'column operator value' with ==, !=, >, <, >=, <=")


def _parse_aggregations(spec: str, column_names: list) -> list:
    """Parse aggregation spec like 'count', 'price:sum,qty:mean' into Arrow group_by format."""
    result = []
    for part in spec.split(","):
        part = part.strip()
        if part == "count":
            result.append(("", "hash_count_all"))
            continue
        if ":" not in part:
            raise ValueError(f"Invalid aggregation '{part}'. Use 'column:function' or 'count'.")
        col, func = part.split(":", 1)
        col, func = col.strip(), func.strip()
        if col not in column_names:
            raise ValueError(f"Aggregation column '{col}' not found. Available: {column_names}")
        func_map = {"sum": "hash_sum", "mean": "hash_mean", "min": "hash_min_max", "max": "hash_min_max", "count": "hash_count"}
        if func not in func_map and func not in func_map.values():
            raise ValueError(f"Unknown aggregation function '{func}'. Supported: sum, mean, min, max, count.")
        result.append((col, func_map.get(func, func)))
    return result


def _compute_agg(column: "pa.ChunkedArray", func: str):
    """Compute a scalar aggregation on a single column."""
    func_map = {
        "hash_sum": pc.sum, "hash_mean": pc.mean,
        "hash_min_max": pc.min_max, "hash_count": pc.count,
        "hash_count_all": lambda _: None,
    }
    compute_fn = func_map.get(func)
    if compute_fn is None:
        raise ValueError(f"Unknown compute function: {func}")
    result = compute_fn(column)
    if hasattr(result, "as_py"):
        return result.as_py()
    return result


# ─── CSV Tools ────────────────────────────────────────────────────────────────


@mcp.tool()
def list_csv_files(directory: Optional[str] = None) -> str:
    """List all CSV files in the configured CSV directory (or a custom path).

    Args:
        directory: Optional override directory. Defaults to CSV_DIR env var.

    Returns:
        JSON with list of CSV file paths, sizes, and row estimates.
    """
    search_dir = directory or CSV_DIR
    try:
        resolved = os.path.realpath(os.path.abspath(search_dir))
        if not os.path.isdir(resolved):
            return json.dumps({"error": f"Directory not found: {search_dir}"})
        files = []
        for p in sorted(Path(resolved).rglob("*.csv")):
            size = p.stat().st_size
            files.append({
                "path": str(p),
                "name": p.name,
                "size_bytes": size,
                "size_kb": round(size / 1024, 1),
            })
        return json.dumps({"directory": str(resolved), "files": files, "total": len(files)}, indent=2)
    except Exception as exc:
        return json.dumps({"error": str(exc)})


@mcp.tool()
def read_csv(
    file_path: str,
    row_limit: Optional[int] = None,
    delimiter: str = ",",
) -> str:
    """Read a CSV file and return its contents as structured JSON.

    Args:
        file_path: Path to the CSV file.
        row_limit: Maximum number of data rows to return (None = all).
        delimiter: Field delimiter character (default: comma).

    Returns:
        JSON with columns, rows (list of dicts), row count, and file stats.
    """
    try:
        resolved = _validate_path(file_path)
    except ValueError as exc:
        return json.dumps({"error": str(exc)})

    try:
        with open(resolved, newline="", encoding="utf-8-sig") as fh:
            reader = csv.DictReader(fh, delimiter=delimiter)
            rows = list(reader)

        total_rows = len(rows)
        columns = list(rows[0].keys()) if rows else []

        if row_limit is not None:
            rows = rows[:row_limit]

        return json.dumps({
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "columns": columns,
            "column_count": len(columns),
            "total_rows": total_rows,
            "returned_rows": len(rows),
            "rows": rows,
        }, indent=2, ensure_ascii=False)
    except Exception as exc:
        return json.dumps({"error": str(exc), "file_path": file_path})


@mcp.tool()
def chunk_csv_for_rag(
    file_path: str,
    rows_per_chunk: int = 50,
    include_header_in_each_chunk: bool = True,
    delimiter: str = ",",
) -> str:
    """Split a CSV file into text chunks ready for the indexer.

    Each chunk is a plain-text block (header + rows) with a source_ref
    indicating the row range, so citations point back to exact rows.

    Args:
        file_path: Path to the CSV file.
        rows_per_chunk: Number of data rows per chunk (default 50).
        include_header_in_each_chunk: Prepend column names to every chunk.
        delimiter: Field delimiter (default: comma).

    Returns:
        JSON list of chunks, each with content and source_ref — ready to
        pass directly to the indexer's index_chunks tool.
    """
    try:
        resolved = _validate_path(file_path)
    except ValueError as exc:
        return json.dumps({"error": str(exc)})

    try:
        with open(resolved, newline="", encoding="utf-8-sig") as fh:
            reader = csv.DictReader(fh, delimiter=delimiter)
            all_rows = list(reader)

        if not all_rows:
            return json.dumps({"chunks": [], "total_chunks": 0, "file_path": file_path})

        columns = list(all_rows[0].keys())
        header_line = delimiter.join(columns)

        chunks = []
        for start in range(0, len(all_rows), rows_per_chunk):
            batch = all_rows[start: start + rows_per_chunk]
            end = start + len(batch)

            lines = []
            if include_header_in_each_chunk:
                lines.append(header_line)
            for row in batch:
                lines.append(delimiter.join(str(row.get(c, "")) for c in columns))

            content = "\n".join(lines)
            source_ref = f"rows {start + 1}–{end}"

            chunks.append({
                "content": content,
                "source_ref": source_ref,
                "metadata": {
                    "row_start": start + 1,
                    "row_end": end,
                    "columns": columns,
                },
            })

        return json.dumps({
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "total_rows": len(all_rows),
            "total_chunks": len(chunks),
            "rows_per_chunk": rows_per_chunk,
            "chunks": chunks,
        }, indent=2, ensure_ascii=False)

    except Exception as exc:
        return json.dumps({"error": str(exc), "file_path": file_path})


# ─── Plain-text / Markdown Tools ──────────────────────────────────────────────


@mcp.tool()
def list_text_files(directory: Optional[str] = None) -> str:
    """List .txt and .md files in the configured text directory (or a custom path).

    Args:
        directory: Optional override directory. Defaults to TEXT_DIR env var.

    Returns:
        JSON with list of text file paths, sizes, and extensions.
    """
    search_dir = directory or TEXT_DIR
    try:
        resolved = os.path.realpath(os.path.abspath(search_dir))
        if not os.path.isdir(resolved):
            return json.dumps({"error": f"Directory not found: {search_dir}"})
        files = []
        for ext in ("*.txt", "*.md"):
            for p in sorted(Path(resolved).rglob(ext)):
                size = p.stat().st_size
                files.append({
                    "path": str(p),
                    "name": p.name,
                    "extension": p.suffix,
                    "size_bytes": size,
                    "size_kb": round(size / 1024, 1),
                })
        files.sort(key=lambda f: f["path"])
        return json.dumps({"directory": str(resolved), "files": files, "total": len(files)}, indent=2)
    except Exception as exc:
        return json.dumps({"error": str(exc)})


@mcp.tool()
def read_text_file(file_path: str) -> str:
    """Read a plain-text or Markdown file and return its content.

    Args:
        file_path: Path to a .txt or .md file.

    Returns:
        JSON with full content, line count, word count, and char count.
    """
    try:
        resolved = _validate_path(file_path)
    except ValueError as exc:
        return json.dumps({"error": str(exc)})

    try:
        content = Path(resolved).read_text(encoding="utf-8")
        lines = content.splitlines()
        return json.dumps({
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "content": content,
            "line_count": len(lines),
            "word_count": len(content.split()),
            "char_count": len(content),
        }, indent=2, ensure_ascii=False)
    except Exception as exc:
        return json.dumps({"error": str(exc), "file_path": file_path})


@mcp.tool()
def chunk_text_for_rag(
    file_path: str,
    chunk_size: int = 512,
    overlap: int = 50,
    split_on_headings: bool = True,
) -> str:
    """Split a .txt or .md file into overlapping word-count chunks for RAG.

    When split_on_headings is True (default), Markdown headings (## …) act as
    hard boundaries before the fixed-size chunking is applied, so section
    structure is preserved.  source_ref is either the nearest heading or the
    word-offset range, enabling precise citations in the indexer.

    Args:
        file_path: Path to a .txt or .md file.
        chunk_size: Target words per chunk (default 512).
        overlap: Words of overlap between consecutive chunks (default 50).
        split_on_headings: Use Markdown headings as hard split points first.

    Returns:
        JSON with chunks list (content + source_ref) ready for index_chunks.
    """
    try:
        resolved = _validate_path(file_path)
    except ValueError as exc:
        return json.dumps({"error": str(exc)})

    try:
        content = Path(resolved).read_text(encoding="utf-8")
    except Exception as exc:
        return json.dumps({"error": str(exc), "file_path": file_path})

    # ── Section splitting ───────────────────────────────────────────────────
    heading_re = re.compile(r"^#{1,6}\s+.+", re.MULTILINE)

    def _split_sections(text: str) -> list[tuple[str, str]]:
        """Return [(heading_label, section_text), …]."""
        matches = list(heading_re.finditer(text))
        if not matches:
            return [("", text)]
        sections = []
        for i, m in enumerate(matches):
            label = m.group(0).lstrip("#").strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            sections.append((label, text[start:end].strip()))
        return sections

    sections: list[tuple[str, str]] = (
        _split_sections(content) if split_on_headings else [("", content)]
    )

    # ── Word-window chunking within each section ────────────────────────────
    chunks: list[dict] = []
    global_word_offset = 0

    def _word_chunks(text: str, heading: str) -> list[dict]:
        nonlocal global_word_offset
        words = text.split()
        result = []
        step = max(1, chunk_size - overlap)
        for start in range(0, max(1, len(words)), step):
            batch = words[start: start + chunk_size]
            if not batch:
                break
            word_start = global_word_offset + start + 1
            word_end = global_word_offset + start + len(batch)
            ref = heading if heading else f"words {word_start}–{word_end}"
            if heading and start > 0:
                ref = f"{heading} (cont.)"
            result.append({
                "content": " ".join(batch),
                "source_ref": ref,
                "metadata": {
                    "word_start": word_start,
                    "word_end": word_end,
                    "section": heading,
                },
            })
        global_word_offset += len(words)
        return result

    for heading, section_text in sections:
        if section_text.strip():
            chunks.extend(_word_chunks(section_text, heading))

    return json.dumps({
        "file_path": file_path,
        "file_name": Path(file_path).name,
        "total_words": len(content.split()),
        "total_chunks": len(chunks),
        "chunk_size": chunk_size,
        "overlap": overlap,
        "split_on_headings": split_on_headings,
        "chunks": chunks,
    }, indent=2, ensure_ascii=False)


@mcp.tool()
def get_document_status() -> str:
    """
    Get the status of all available document processing engines and capabilities.

    Returns:
        JSON with engine availability, configuration, and tool listing.
    """
    return json.dumps({
        "server": "document-mcp",
        "engines": {
            "docling": {
                "available": DOCLING_AVAILABLE,
                "use_cases": "Excel extraction, PDF complex layouts, tables, figures",
                "install": "uv add docling",
            },

            "python_docx": {
                "available": DOCX_AVAILABLE,
                "use_cases": "Word (.docx) parsing, table extraction, Markdown conversion",
                "install": "uv add python-docx",
            },
            "anthropic_vision": {
                "available": ANTHROPIC_AVAILABLE,
                "use_cases": "AI OCR for scanned PDFs, figure/chart descriptions",
                "install": "uv add anthropic",
                "note": "Requires ANTHROPIC_API_KEY environment variable",
            },
            "opensearch": {
                "available": OPENSEARCH_AVAILABLE,
                "use_cases": "Full-text search indexing for large Excel datasets",
                "install": "uv add opensearch-py",
            },
            "pyarrow": {
                "available": PYARROW_AVAILABLE,
                "use_cases": "Parquet read/query/aggregate, Arrow columnar compute",
                "install": "uv add pyarrow",
            },
        },
        "configuration": {
            "excel_dir": os.path.abspath(EXCEL_DIR),
            "pdf_dir": os.path.abspath(PDF_DIR),
            "word_dir": os.path.abspath(WORD_DIR),
            "output_dir": os.path.abspath(OUTPUT_DIR),
            "csv_dir": os.path.abspath(CSV_DIR),
            "text_dir": os.path.abspath(TEXT_DIR),
        },
        "tools": {
            "excel": [
                "list_excel_files",
                "extract_excel_content",
                "query_excel_data",
                "get_excel_summary",
                "convert_excel_to_csv",
                "convert_excel_to_json",
                "extract_excel_rows",
                "chunk_excel_content",
                "index_excel_in_opensearch",
                "search_excel_opensearch",
                "list_opensearch_indices",
                "smart_analyze_excel",
                "smart_extract_tables",
                "smart_extract_text",
                "smart_export_all",
                "smart_get_sheet_overview",
                "smart_table_to_format",
            ],
            "pdf": [
                "list_pdf_files",
                "extract_pdf",
                "extract_pdf_tables",
                "extract_pdf_figures",
                "chunk_pdf_for_rag",
                "batch_extract_pdfs",
                "analyze_pdf_structure",
                "describe_pdf_figures",
            ],
            "word": [
                "list_word_files",
                "extract_word_content",
                "get_word_summary",
                "word_to_markdown",
                "extract_word_tables",
                "word_to_text",
            ],
            "parquet_compute": [
                "read_parquet",
                "parquet_schema",
                "query_parquet",
                "aggregate_parquet",
            ],
            "csv": [
                "list_csv_files",
                "read_csv",
                "chunk_csv_for_rag",
            ],
            "text": [
                "list_text_files",
                "read_text_file",
                "chunk_text_for_rag",
            ],
        },
    }, indent=2)


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    mcp.run(transport="stdio")
